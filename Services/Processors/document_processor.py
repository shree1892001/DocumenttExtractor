import os
import re
import json
import logging
import pytesseract
from PIL import Image
import cv2
import numpy as np
from pdf2image import convert_from_path
from docx import Document
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Union, Any
import traceback
from Common.document_constants import (
    DOCUMENT_CATEGORIES,
    DOCUMENT_TYPE_MAPPING,
    DOCUMENT_INDICATORS,
    MIN_CONFIDENCE_THRESHOLD,
    HIGH_CONFIDENCE_THRESHOLD,
    MIN_GENUINENESS_SCORE,
    VERIFICATION_THRESHOLD,
    SUPPORTED_EXTENSIONS,
    FIELD_PATTERNS,
    DOCUMENT_SEPARATORS,
    OCR_ERROR_PATTERNS,
    CHARACTER_PATTERNS,
    CONTENT_INDICATORS,
    DOCUMENT_INDICATOR_KEYWORDS,
    NON_GENUINE_INDICATORS,
    DOCUMENT_PROMPTS,
    DOCUMENT_FIELD_TEMPLATES
)
from Services.Processors.text_processor import TextProcessor
from Services.Classifiers.document_classifier import DocumentClassifier
from Services.Verifiers.document_verifier import DocumentVerifier
from Services.Classifiers.template_matcher import TemplateMatcher

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, api_key: str, templates_dir: str = "D:\\imageextractor\\identites\\Templates"):
        self.api_key = api_key
        self.templates_dir = templates_dir
        self.classifier = DocumentClassifier(api_key)
        self.text_processor = TextProcessor(api_key)
        self.verifier = DocumentVerifier(api_key)
        self.template_matcher = TemplateMatcher(api_key, templates_dir)

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Process a document and extract information"""
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in SUPPORTED_EXTENSIONS:
                raise ValueError(f"Unsupported file extension: {file_ext}")
            
            # Extract text based on file type
            if file_ext in {'.jpg', '.jpeg', '.png', '.tiff', '.bmp'}:
                text = self._extract_text_from_image(file_path)
            elif file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_path)
            elif file_ext == '.docx':
                text = self._extract_text_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Process the text content
            result = self._process_text_content(text, file_path, MIN_CONFIDENCE_THRESHOLD)
            
            if result and result.get("status") == "success":
                return result
            else:
                raise ValueError(f"Document processing failed: {result.get('rejection_reason', 'Unknown error')}")
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            logger.error(traceback.format_exc())
            raise

    def _extract_text_from_image(self, file_path: str) -> str:
        """Extract text from an image file"""
        try:
            img = Image.open(file_path)
            ocr_text = pytesseract.image_to_string(img)
            
            if not self._is_good_ocr_result(ocr_text):
                logger.info("Tesseract OCR yielded poor results, trying Gemini Vision")
                response = self.text_processor.process_text(
                    "Extract all text from this document image. Return only the raw text without any formatting."
                )
                ocr_text = response.strip()
            
            return ocr_text
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            raise

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file"""
        try:
            text = ""
            images = convert_from_path(file_path)
            
            for img in images:
                ocr_text = pytesseract.image_to_string(img)
                if not self._is_good_ocr_result(ocr_text):
                    logger.info("Tesseract OCR yielded poor results, trying Gemini Vision")
                    response = self.text_processor.process_text(
                        "Extract all text from this document image. Return only the raw text without any formatting."
                    )
                    ocr_text = response.strip()
                text += ocr_text + "\n\n"
            
            return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text.append(cell.text.strip())
            
            return "\n".join(text)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    def _is_good_ocr_result(self, text: str) -> bool:
        """Check if OCR result is of good quality"""
        try:
            if not text.strip():
                return False
            
            if len(text.strip()) < 10:
                return False
            
            if not self._has_meaningful_content(text):
                return False
            
            for pattern in OCR_ERROR_PATTERNS:
                if re.search(pattern, text):
                    return False
            
            return True
        except Exception as e:
            logger.error(f"Error checking OCR result quality: {str(e)}")
            return False

    def _has_meaningful_content(self, text: str) -> bool:
        """Check if text contains meaningful content"""
        try:
            indicator_count = 0
            for pattern in CONTENT_INDICATORS:
                if re.search(pattern, text):
                    indicator_count += 1
            
            return indicator_count >= 2
        except Exception as e:
            logger.error(f"Error checking meaningful content: {str(e)}")
            return False

    def _process_text_content(self, text: str, source_file: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process text content directly without creating temporary files"""
        try:
            # Use detection prompt from constants
            detection_prompt = DOCUMENT_PROMPTS["detection"].format(text=text)
            response = self.text_processor.process_text(text, detection_prompt)
            detection_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())

            doc_type = detection_result.get("document_type", "").lower()
            confidence = detection_result.get("confidence", 0.0)

            doc_type = self._standardize_document_type(doc_type)

            if doc_type == "unknown" or confidence < min_confidence:
                logger.warning(f"Could not confidently determine document type (confidence: {confidence})")
                return {
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": source_file,
                    "rejection_reason": f"Document type could not be determined with sufficient confidence (confidence: {confidence})",
                    "confidence": confidence
                }

            logger.info(f"Detected document type: {doc_type} with confidence {confidence}")

            # Verify document genuineness
            verification_result = self.verifier.verify_document_genuineness(text)
            if not verification_result["is_genuine"]:
                return {
                    "status": "rejected",
                    "document_type": doc_type,
                    "source_file": source_file,
                    "rejection_reason": verification_result["rejection_reason"],
                    "verification_result": verification_result
                }

            # Extract fields
            fields_template = DOCUMENT_FIELD_TEMPLATES.get(doc_type, "Extract all visible text and information from the document.")
            extraction_prompt = DOCUMENT_PROMPTS["extraction"].format(
                doc_type=doc_type,
                text=text,
                fields=fields_template
            )

            response = self.text_processor.process_text(text, extraction_prompt)
            extracted_data = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
            logger.info(f"Extracted data: {json.dumps(extracted_data, indent=2)}")

            # Verify extracted data
            verification_result = self.verifier.verify_document(extracted_data, doc_type)
            if not verification_result["is_genuine"]:
                return {
                    "status": "rejected",
                    "document_type": doc_type,
                    "source_file": source_file,
                    "rejection_reason": verification_result["rejection_reason"],
                    "verification_result": verification_result
                }

            return {
                "extracted_data": extracted_data,
                "status": "success",
                "confidence": confidence,
                "document_type": doc_type,
                "source_file": source_file,
                "validation_level": "strict",
                "verification_result": verification_result
            }

        except Exception as e:
            logger.exception(f"Error processing text content: {str(e)}")
            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": source_file,
                "rejection_reason": f"Error processing document: {str(e)}"
            }

    def _standardize_document_type(self, doc_type: str) -> str:
        """Standardize document type using pattern matching"""
        try:
            doc_type = doc_type.lower().strip()
            return DOCUMENT_TYPE_MAPPING.get(doc_type, doc_type)
        except Exception as e:
            logger.error(f"Error standardizing document type: {str(e)}")
            return doc_type 