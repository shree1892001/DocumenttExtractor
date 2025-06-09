from abc import ABC, abstractmethod

import fitz
import tempfile
from typing import List, Dict, Any, Optional, Tuple
import os
import json
import re
from docx import Document
from Common.constants import *
from PIL import Image
import pytesseract
import traceback
import sys
import logging
from dataclasses import dataclass
import google.generativeai as genai
from Extractor.ImageExtractor import ImageTextExtractor
from Extractor.Paddle import flatten_json
from Factories.DocumentFactory import (
    DocumentExtractorFactory,
    DocumentExtractor,
    TextExtractorFactory,
    BaseTextExtractor,
    DocxTextExtractor,
    PdfTextExtractor,
    ImageTextExtractor,
    TextExtractor
)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Constants
MIN_TEXT_LENGTH = 50
MIN_CONFIDENCE_THRESHOLD = 0.4
HIGH_CONFIDENCE_THRESHOLD = 0.8
MIN_GENUINENESS_SCORE = 0.6
VERIFICATION_THRESHOLD = 0.5
SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.jpg', '.jpeg', '.png', '.tiff', '.bmp'}


@dataclass
class DocumentInfo:
    document_type: str
    confidence_score: float
    extracted_data: Dict[str, Any]
    matched_fields: Dict[str, Any]


class BaseTextExtractor(ABC):
    def __init__(self, api_key: str):
        self.api_key = API_KEY_3
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        pass


class TextProcessor:
    def __init__(self, api_key: str):
        self.api_key = API_KEY_3
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    def process_text(self, text: str, prompt: str) -> str:
        """Process text using Gemini without image handling"""
        try:
            response = self.model.generate_content([prompt, text])
            return response.text
        except Exception as e:
            logger.error(f"Error processing text with Gemini: {str(e)}")
            raise


class DocumentProcessor:
    # Constants
    MIN_CONFIDENCE_THRESHOLD = 0.4
    HIGH_CONFIDENCE_THRESHOLD = 0.8
    MIN_GENUINENESS_SCORE = 0.6
    VERIFICATION_THRESHOLD = 0.5
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.jpg', '.jpeg', '.png', '.tiff', '.bmp'}

    def __init__(self, api_key: str, templates_dir: str = "D:\\imageextractor\\identites\\Templates"):
        self.api_key = API_KEY_3
        self.templates_dir = templates_dir

        # Initialize text processing components
        self.text_processor = TextProcessor(api_key)
        self.text_extractor = TextExtractor(api_key)

        # Initialize Gemini models
        try:
            genai.configure(api_key=api_key)
            self.gemini_vision_model = genai.GenerativeModel(
                '                                                                                                                                                                                                                                                     h')
            self.gemini_text_model = genai.GenerativeModel('gemini-1.5-flash')
        except Exception as e:
            logger.error(f"Error initializing Gemini models: {str(e)}")
            raise RuntimeError(f"Gemini initialization failed: {str(e)}")

        # Initialize Pytesseract
        try:
            version = pytesseract.get_tesseract_version()
            logger.info(f"Tesseract version: {version}")
        except Exception as e:
            logger.error(f"Error initializing Pytesseract: {str(e)}")
            raise RuntimeError(f"Pytesseract initialization failed: {str(e)}")

        # Initialize document categories
        self.document_categories = {
            'identity': [
                'passport', 'national_id', 'drivers_license', 'residence_permit',
                'citizenship_card', 'voter_id', 'aadhaar_card', 'pan_card'
            ],
            'legal': [
                'contract', 'agreement', 'deed', 'power_of_attorney',
                'court_order', 'legal_notice', 'affidavit', 'will'
            ],
            'financial': [
                'bank_statement', 'tax_return', 'invoice', 'receipt',
                'credit_card_statement', 'loan_document', 'insurance_policy'
            ],
            'educational': [
                'degree_certificate', 'diploma', 'transcript', 'report_card',
                'scholarship_document', 'academic_certificate'
            ],
            'medical': [
                'medical_report', 'prescription', 'health_insurance',
                'vaccination_record', 'medical_certificate'
            ],
            'business': [
                'business_license', 'incorporation_document', 'tax_registration',
                'commercial_invoice', 'shipping_document'
            ],
            'other': []  # For uncategorized document types
        }

        # Initialize document patterns
        self.document_patterns = self._initialize_document_patterns()
        self.compiled_patterns = {
            doc_type: [re.compile(pattern) for pattern in patterns]
            for doc_type, patterns in self.document_patterns.items()
        }

    def _create_text_processor(self):
        """Create and configure the text processor"""
        try:
            return self.gemini_text_model
                except Exception as e:
            logger.error(f"Error creating text processor: {str(e)}")
            raise RuntimeError(f"Text processor creation failed: {str(e)}")

    def _split_into_chunks(self, text: str) -> List[str]:
        """Split text into chunks based on document boundaries"""
        try:
            # Common document separators
            separators = [
                r'\n\s*\n\s*\n',  # Multiple blank lines
                r'[-=]{3,}',  # Lines of dashes or equals
                r'_{3,}',  # Lines of underscores
                r'\*{3,}',  # Lines of asterisks
                r'Page \d+',  # Page numbers
                r'Document \d+',  # Document numbers
                r'Copy \d+',  # Copy numbers
                r'Original',  # Original document marker
                r'Duplicate',  # Duplicate document marker
                r'COPY',  # COPY marker
                r'ORIGINAL'  # ORIGINAL marker
            ]

            # Combine separators into a single pattern
            separator_pattern = '|'.join(f'({sep})' for sep in separators)

            # Split text into chunks
            chunks = re.split(separator_pattern, text)

            # Filter out empty chunks and separators
            chunks = [chunk.strip() for chunk in chunks if chunk and not any(sep in chunk for sep in separators)]

            # If no chunks were found, return the whole text as one chunk
            if not chunks:
                return [text]

            return chunks

        except Exception as e:
            logger.error(f"Error splitting text into chunks: {str(e)}")
            return [text]

    def process_file(self, file_path: str, min_confidence: float = 0.0) -> List[Dict[str, Any]]:
        """Process a file that may contain multiple documents"""
        if file_path.lower().endswith('.pdf'):
            return self._process_pdf(file_path, min_confidence)
        elif file_path.lower().endswith('.docx'):
            return self._process_docx(file_path, min_confidence)
        else:
            result = self._process_single_image(file_path, min_confidence)
            return [result] if result else []

    def _process_docx(self, file_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a DOCX file and extract text from both content and embedded images"""
        processed_images = set()  # Track processed images to avoid duplicates
        consolidated_results = []

        try:
            # First, try to extract and process images
            logger.info("Attempting to extract and process images from DOCX")
            document_segments = self._extract_text_from_docx_images(file_path)

            if document_segments:
                for segment in document_segments:
                    # Skip if we've already processed this image
                    if segment["image_path"] in processed_images:
                    continue

                    processed_images.add(segment["image_path"])

                    if segment["text"].strip():
                        logger.info(f"Processing image {segment['image_number']}")
                        # Try to process as multiple documents first
                        multi_doc_results = self._process_multiple_documents(segment["text"], file_path, min_confidence)
                        if multi_doc_results:
                            for result in multi_doc_results:
                                result["image_number"] = segment["image_number"]
                                consolidated_results.append(result)
                        else:
                            # If no multiple documents found, process as single document
                            result = self._process_text_content(segment["text"], file_path, min_confidence)
                            if result:
                                result["image_number"] = segment["image_number"]
                                consolidated_results.append(result)

            # Then process the document text
            doc = Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            # Process the combined text content
            if text_content:
                combined_text = "\n".join(text_content)
                # Try to process as multiple documents first
                multi_doc_results = self._process_multiple_documents(combined_text, file_path, min_confidence)
                if multi_doc_results:
                    consolidated_results.extend(multi_doc_results)
            else:
                    # If no multiple documents found, process as single document
                    result = self._process_text_content(combined_text, file_path, min_confidence)
                    if result:
                        consolidated_results.append(result)

            # If no valid results were found
            if not consolidated_results:
                logger.warning(f"No valid content found in DOCX file: {file_path}")
                return [{
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": file_path,
                    "rejection_reason": "No valid content could be extracted"
                }]

            # Log the results
            for result in consolidated_results:
                logger.info(f"Processed document type: {result.get('document_type', 'unknown')}, "
                            f"Status: {result.get('status', 'unknown')}, "
                            f"Image number: {result.get('image_number', 'N/A')}")

            return consolidated_results

        except Exception as e:
            logger.error(f"Error processing DOCX file: {str(e)}")
            return [{
                "status": "error",
                "document_type": "unknown",
                "source_file": file_path,
                "rejection_reason": f"Error processing document: {str(e)}",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "error_traceback": traceback.format_exc()
                }
            }]

    def _consolidate_results(self, results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Consolidate results to remove duplicates and combine related documents"""
        try:
            consolidated = []
            seen_documents = set()

            for result in results:
                if result.get("status") != "success":
                    continue

                # Create a unique identifier for the document
                doc_type = result.get("document_type", "unknown")
                extracted_data = result.get("extracted_data", {})

                # Create a unique key based on document type and key fields
                if doc_type == "aadhaar_card":
                    key = f"aadhaar_{extracted_data.get('Aadhaar Number', '')}"
                elif doc_type == "license":
                    key = f"license_{extracted_data.get('Name', '')}_{extracted_data.get('Date of Birth', '')}"
                else:
                    key = f"{doc_type}_{json.dumps(extracted_data)}"

                if key not in seen_documents:
                    seen_documents.add(key)
                    consolidated.append(result)

            return consolidated

        except Exception as e:
            logger.error(f"Error consolidating results: {str(e)}")
            return results

    def _process_text_content(self, text: str, source_file: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process text content and extract document information"""
        try:
            if not text or not isinstance(text, str):
                raise ValueError("Invalid text input: text must be a non-empty string")

            if not source_file or not isinstance(source_file, str):
                raise ValueError("Invalid source file: source_file must be a non-empty string")

            if not isinstance(min_confidence, (int, float)) or not 0 <= min_confidence <= 1:
                raise ValueError("Invalid confidence threshold: min_confidence must be between 0 and 1")

            # Create detection prompt
            detection_prompt = f"""
            Analyze this text and determine what type of document it is. Look for any official document indicators.
            Consider these aspects:
            1. Document Type Indicators:
               - Official headers and titles
               - Government or organizational logos
               - Document numbers or identifiers
               - Official stamps or seals
               - Security features mentioned
               - Issuing authority information

            2. Document Categories:
               - Identity documents (ID cards, passports, etc.)
               - Legal documents (certificates, licenses, etc.)
               - Financial documents (statements, receipts, etc.)
               - Educational documents (degrees, certificates, etc.)
               - Business documents (contracts, agreements, etc.)
               - Government documents (permits, registrations, etc.)
               - Medical documents (reports, prescriptions, etc.)
               - Any other official document type

            Text:
            {text}

            Return a JSON response with:
            {{
                "document_type": "detected_type",
                "document_category": "category",
                "confidence": 0.0-1.0,
                "reasoning": "explanation",
                "key_indicators": ["list of key indicators found"],
                "issuing_authority": "detected authority if any"
            }}
            """

            # Process text with Gemini
            response = self.text_processor.process_text(text, detection_prompt)
            if not response:
                raise ValueError("Empty response from text processor")

            try:
                detection_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse detection result: {str(e)}")
                raise ValueError(f"Invalid JSON response from text processor: {str(e)}")

            doc_type = detection_result.get("document_type", "").lower()
            confidence = detection_result.get("confidence", 0.0)

            if not doc_type or doc_type == "unknown" or confidence < min_confidence:
                logger.warning(f"Could not confidently determine document type (confidence: {confidence})")
                return {
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": source_file,
                    "rejection_reason": f"Document type could not be determined with sufficient confidence (confidence: {confidence})",
                    "confidence": confidence,
                    "extracted_data": {}
                }

            logger.info(f"Detected document type: {doc_type} with confidence {confidence}")

            # Create extraction prompt
            extraction_prompt = f"""
            Extract all relevant information from this document.

            Document Text:
            {text}

            Document Type: {doc_type}
            Document Category: {detection_result.get('document_category', 'unknown')}
            Issuing Authority: {detection_result.get('issuing_authority', 'unknown')}

            Return the extracted information in JSON format with this structure:
            {{
                "data": {{
                    // All extracted fields
                }},
                "confidence": 0.0-1.0,
                "additional_info": "Any additional relevant information",
                "document_metadata": {{
                    "type": "{doc_type}",
                    "category": "{detection_result.get('document_category', 'unknown')}",
                    "issuing_authority": "{detection_result.get('issuing_authority', 'unknown')}",
                    "key_indicators": {json.dumps(detection_result.get('key_indicators', []))}
                }}
            }}
            """

            # Extract information
            response = self.text_processor.process_text(text, extraction_prompt)
            if not response:
                raise ValueError("Empty response from text processor during extraction")

            try:
                extracted_data = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
                logger.info(f"Extracted data: {json.dumps(extracted_data, indent=2)}")
            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse extraction result: {str(e)}")
                raise ValueError(f"Invalid JSON response from text processor during extraction: {str(e)}")

            if not extracted_data:
                raise ValueError("Invalid extraction result: missing required fields")

            # Verify document
            verification_result = self.verify_document(extracted_data, doc_type)

            if not verification_result["is_genuine"]:
                rejection_reason = verification_result.get("rejection_reason",
                                                           "Document failed authenticity verification")
                logger.warning(f"Document rejected - Not genuine: {rejection_reason}")
                return {
                    "status": "rejected",
                    "document_type": doc_type,
                    "source_file": source_file,
                    "rejection_reason": rejection_reason,
                    "verification_result": verification_result,
                    "extracted_data": extracted_data
                }

            return {
                "status": "success",
                "document_type": doc_type,
                "confidence": confidence,
                "extracted_data": extracted_data,
                "source_file": source_file,
                "validation_level": "strict",
                "verification_result": verification_result,
                "document_metadata": extracted_data.get("document_metadata", {})
            }

        except Exception as e:
            logger.error(f"Error processing text content: {str(e)}")
            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": source_file,
                "rejection_reason": f"Error processing document: {str(e)}",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "error_traceback": traceback.format_exc()
                },
                "extracted_data": {}
            }

    def verify_document(self, extracted_data: Dict[str, Any], doc_type: str) -> Dict[str, Any]:
        """Verify document authenticity and data validity"""
        try:
            verification_prompt = f"""
            You are a document verification expert. Analyze this document data and determine if it is genuine.
            Provide a detailed verification report with specific checks and findings.

            Document Data:
            {json.dumps(extracted_data, indent=2)}

            Document Type: {doc_type}
            Document Category: {extracted_data.get('document_metadata', {}).get('category', 'unknown')}
            Issuing Authority: {extracted_data.get('document_metadata', {}).get('issuing_authority', 'unknown')}

            Perform the following checks:

            1. Document Authenticity:
               - Verify presence of required fields for this document type
               - Check for official text and formatting
               - Validate document structure and layout
               - Check for security features
               - Verify document quality indicators
               - Look for official headers and footers
               - Check for any official stamps or marks

            2. Security Feature Analysis:
               - Look for official seals and stamps
               - Check for watermarks or holograms
               - Verify presence of security patterns
               - Look for QR codes or barcodes
               - Check for any digital signatures
               - Verify any security numbers or codes
               - Look for any anti-counterfeit measures

            3. Data Validation:
               - Verify all identification numbers and their formats
               - Check date formats and their logical consistency
               - Validate name and address formatting
               - Look for any data inconsistencies
               - Verify field relationships and dependencies
               - Check for logical data patterns
               - Validate any reference numbers

            4. Document Quality:
               - Check printing quality indicators
               - Verify text alignment
               - Check for any smudges or marks
               - Verify color consistency
               - Check for any signs of poor quality
               - Verify professional formatting
               - Check for any signs of manipulation

            Return a JSON response with:
            {{
                "is_genuine": true/false,
                "confidence_score": 0.0-1.0,
                "rejection_reason": "reason if not genuine",
                "verification_checks": {{
                    "authenticity": {{
                        "passed": true/false,
                        "details": "explanation",
                        "confidence": 0.0-1.0
                    }},
                    "security_features": {{
                        "passed": true/false,
                        "details": "explanation",
                        "confidence": 0.0-1.0
                    }},
                    "data_validation": {{
                        "passed": true/false,
                        "details": "explanation",
                        "confidence": 0.0-1.0
                    }},
                    "quality": {{
                        "passed": true/false,
                        "details": "explanation",
                        "confidence": 0.0-1.0
                    }}
                }},
                "security_features_found": ["list of security features"],
                "verification_summary": "Overall verification summary",
                "recommendations": ["list of recommendations for improvement"]
            }}

            Important:
            - Be thorough but fair in your analysis
            - Consider document-specific requirements
            - Account for variations in official document formats
            - Consider both digital and physical security features
            - Look for logical consistency in the data
            - Consider the document's intended use
            - Be lenient with minor formatting variations
            - Focus on key security features and data validity
            - Consider official document standards
            - Account for regional variations in document formats
            """

            response = self.text_processor.process_text(json.dumps(extracted_data), verification_prompt)
            verification_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())

            checks = verification_result.get("verification_checks", {})
            confidence_scores = [
                checks.get("authenticity", {}).get("confidence", 0),
                checks.get("security_features", {}).get("confidence", 0),
                checks.get("data_validation", {}).get("confidence", 0),
                checks.get("quality", {}).get("confidence", 0)
            ]
            overall_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0

            verification_result["confidence_score"] = overall_confidence
            verification_result["is_genuine"] = overall_confidence >= self.VERIFICATION_THRESHOLD

            logger.info(f"Document verification results: {json.dumps(verification_result, indent=2)}")

            return verification_result

                    except Exception as e:
            logger.exception(f"Error verifying document genuineness: {str(e)}")
            return {
                "is_genuine": False,
                "confidence_score": 0.0,
                "rejection_reason": f"Error during verification: {str(e)}",
                "verification_checks": {
                    "authenticity": {"passed": False, "details": f"Error during verification: {str(e)}",
                                     "confidence": 0.0},
                    "security_features": {"passed": False, "details": "Verification failed", "confidence": 0.0},
                    "data_validation": {"passed": False, "details": "Verification failed", "confidence": 0.0},
                    "quality": {"passed": False, "details": "Verification failed", "confidence": 0.0}
                },
                "security_features_found": [],
                "verification_summary": f"Document verification failed due to error: {str(e)}",
                "recommendations": ["Verification process failed due to technical error"]
            }

    def _process_pdf(self, pdf_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a PDF file that may contain multiple documents"""
        try:
            if not pdf_path or not os.path.exists(pdf_path):
                raise ValueError(f"Invalid PDF path: {pdf_path}")

            if not isinstance(min_confidence, (int, float)) or not 0 <= min_confidence <= 1:
                raise ValueError("Invalid confidence threshold: min_confidence must be between 0 and 1")

            results = []
            pdf_document = None

            try:
                pdf_document = fitz.open(pdf_path)
        except Exception as e:
                logger.error(f"Error opening PDF file: {str(e)}")
                raise ValueError(f"Failed to open PDF file: {str(e)}")

            try:
                # First try to process the entire PDF as one document
                all_text = ""
                for page_num in range(pdf_document.page_count):
                    try:
                        page = pdf_document[page_num]
                        text = page.get_text()
                        if text.strip():
                            all_text += text + "\n\n"
        except Exception as e:
                        logger.warning(f"Error processing page {page_num + 1}: {str(e)}")
                        continue

                if all_text.strip():
                    try:
                        # Try to process as multiple documents first
                        multi_doc_results = self._process_multiple_documents(all_text, pdf_path, min_confidence)
                        if multi_doc_results:
                            results.extend(multi_doc_results)
                            return results
        except Exception as e:
                        logger.warning(f"Error processing PDF as multiple documents: {str(e)}")

                # If multiple document processing didn't yield results, process page by page
                for page_num in range(pdf_document.page_count):
        try:
                        logger.info(f"Processing page {page_num + 1} of {pdf_document.page_count}")

                page = pdf_document[page_num]
                text = page.get_text()

                        needs_ocr = self._needs_ocr(text, page)

                        if needs_ocr:
                            logger.info(f"Page {page_num + 1} requires OCR processing")
                            try:
                                images = page.get_images(full=True)
        except Exception as e:
                                logger.error(f"Error extracting images from page {page_num + 1}: {str(e)}")
                                continue

                            if not images:
                                logger.warning(f"No images found on page {page_num + 1} for OCR")
                                continue

                            for img_index, img_info in enumerate(images):
                                try:
                                    with tempfile.TemporaryDirectory() as temp_dir:
                                        try:
                                            img_data = pdf_document.extract_image(img_info[0])
                                        except Exception as e:
                                            logger.error(f"Error extracting image data: {str(e)}")
                                            continue

                                        temp_image_path = os.path.join(temp_dir, f"page_{page_num}_img_{img_index}.png")

                                        try:
                                            with open(temp_image_path, "wb") as img_file:
                                                img_file.write(img_data["image"])
        except Exception as e:
                                            logger.error(f"Error saving image: {str(e)}")
                                            continue

        try:
                                            ocr_text = self._perform_ocr(temp_image_path)
        except Exception as e:
                                            logger.error(f"Error performing OCR: {str(e)}")
                                            continue

                                        if ocr_text:
                                            try:
                                                # Try to process OCR text as multiple documents
                                                multi_doc_results = self._process_multiple_documents(ocr_text, pdf_path,
                                                                                                     min_confidence)
                                                if multi_doc_results:
                                                    for result in multi_doc_results:
                                                        result["page_number"] = page_num + 1
                                                        result["image_index"] = img_index + 1
                                                        result["processing_method"] = "ocr"
                                                    results.extend(multi_doc_results)
                                                else:
                                                    # Process as single document if multiple document detection failed
                                                    result = self._process_text_content(ocr_text, pdf_path,
                                                                                        min_confidence)
                                                    if result:
                                                        result["page_number"] = page_num + 1
                                                        result["image_index"] = img_index + 1
                                                        result["processing_method"] = "ocr"
                                                        results.append(result)
                                            except Exception as e:
                                                logger.error(f"Error processing OCR text: {str(e)}")
                                                continue
                                except Exception as e:
                                    logger.error(
                                        f"Error processing image {img_index + 1} on page {page_num + 1}: {str(e)}")
                                    continue
                        else:
                            if text.strip():
                                try:
                                    # Try to process text as multiple documents
                                    multi_doc_results = self._process_multiple_documents(text, pdf_path, min_confidence)
                                    if multi_doc_results:
                                        for result in multi_doc_results:
                                            result["page_number"] = page_num + 1
                                            result["processing_method"] = "direct_text"
                                        results.extend(multi_doc_results)
                                    else:
                                        # Process as single document if multiple document detection failed
                                        result = self._process_text_content(text, pdf_path, min_confidence)
                                        if result:
                                            result["page_number"] = page_num + 1
                                            result["processing_method"] = "direct_text"
                                            results.append(result)
                                except Exception as e:
                                    logger.error(f"Error processing text content: {str(e)}")
                                    continue
                    except Exception as e:
                        logger.error(f"Error processing page {page_num + 1}: {str(e)}")
                        continue

                if not results:
                    logger.warning(f"No valid content found in PDF file: {pdf_path}")
                    return [{
                        "status": "rejected",
                        "document_type": "unknown",
                        "source_file": pdf_path,
                        "rejection_reason": "No valid content found in document"
                    }]

            return results

        except Exception as e:
                logger.error(f"Error processing PDF file: {str(e)}")
                return [{
                    "status": "error",
                    "document_type": "unknown",
                    "source_file": pdf_path,
                    "rejection_reason": f"Error processing document: {str(e)}",
                    "error_details": {
                        "error_type": type(e).__name__,
                        "error_message": str(e),
                        "error_traceback": traceback.format_exc()
                    }
                }]

            finally:
                if pdf_document:
                    try:
                        pdf_document.close()
                    except Exception as e:
                        logger.error(f"Error closing PDF file: {str(e)}")

        except Exception as e:
            logger.error(f"Error in PDF processing: {str(e)}")
            return [{
                "status": "error",
                "document_type": "unknown",
                "source_file": pdf_path,
                "rejection_reason": f"Error processing document: {str(e)}",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "error_traceback": traceback.format_exc()
                }
            }]

    def _needs_ocr(self, text: str, page) -> bool:
        """Determine if OCR is needed for the page"""
        try:

            if not text.strip():
                return True

            if self._is_character_by_character(text):
                return True

            images = page.get_images(full=True)
            if images:

                if not self._has_meaningful_content(text):
                    return True

            if self._has_document_indicators(page):
                return True

            return False

        except Exception as e:
            logger.error(f"Error determining OCR need: {str(e)}")
            return True

    def _is_character_by_character(self, text: str) -> bool:
        """Check if text appears to be extracted character by character"""
        try:

            patterns = [
                r'[A-Z]\s+[A-Z]\s+[A-Z]',
                r'[a-z]\s+[a-z]\s+[a-z]',
                r'[0-9]\s+[0-9]\s+[0-9]',
                r'[A-Za-z]\s+[A-Za-z]\s+[A-Za-z]'
            ]

            for pattern in patterns:
                if re.search(pattern, text):
                    return True

            if len(text) > 0:
                space_ratio = text.count(' ') / len(text)
                if space_ratio > 0.3:
                    return True

            return False

        except Exception as e:
            logger.error(f"Error checking character-by-character extraction: {str(e)}")
            return False

    def _has_meaningful_content(self, text: str) -> bool:
        """Check if text contains meaningful content"""
        try:

            indicators = [
                r'[A-Z]{2,}',
                r'\d{4,}',
                r'[A-Za-z]+\s+[A-Za-z]+',
                r'[A-Za-z]+\s+\d+',
                r'\d+\s+[A-Za-z]+'
            ]

            indicator_count = 0
            for pattern in indicators:
                if re.search(pattern, text):
                    indicator_count += 1

            return indicator_count >= 2

        except Exception as e:
            logger.error(f"Error checking meaningful content: {str(e)}")
            return False

    def _has_document_indicators(self, page) -> bool:
        """Check if page has indicators of being a document that might need OCR"""
        try:

            images = page.get_images(full=True)
            if not images:
                return False

            blocks = page.get_text("blocks")
            if not blocks:
                return True

            for block in blocks:

                if block[4].strip():
                    if any(indicator in block[4].lower() for indicator in [
                        "name", "date", "address", "signature", "photo",
                        "passport", "license", "id", "number", "issued"
                    ]):
                        return True

            return False

        except Exception as e:
            logger.error(f"Error checking document indicators: {str(e)}")
            return True

    def _perform_ocr(self, image_path: str) -> str:
        """Perform OCR on an image"""
        try:
            if not image_path or not os.path.exists(image_path):
                raise ValueError(f"Invalid image path: {image_path}")

            try:
                img = Image.open(image_path)
        except Exception as e:
                logger.error(f"Error opening image: {str(e)}")
                raise ValueError(f"Failed to open image: {str(e)}")

            try:
                ocr_text = pytesseract.image_to_string(img)
            except Exception as e:
                logger.error(f"Tesseract OCR failed: {str(e)}")
                raise ValueError(f"OCR processing failed: {str(e)}")

            if not self._is_good_ocr_result(ocr_text):
                logger.info("Tesseract OCR yielded poor results, trying Gemini Vision")
                try:
                    response = self.text_extractor.process_with_gemini(
                        image_path,
                        "Extract all text from this document image. Return only the raw text without any formatting."
                    )
                    if not response:
                        raise ValueError("Empty response from Gemini Vision")
                    ocr_text = response.strip()
                except Exception as e:
                    logger.error(f"Gemini Vision OCR failed: {str(e)}")
                    raise ValueError(f"Alternative OCR processing failed: {str(e)}")

            return ocr_text

        except Exception as e:
            logger.error(f"Error performing OCR: {str(e)}")
            raise ValueError(f"OCR processing failed: {str(e)}")

    def _is_good_ocr_result(self, text: str) -> bool:
        """Check if OCR result is of good quality"""
        try:
            if not text.strip():
                return False

            if len(text.strip()) < 10:
                return False

            if not self._has_meaningful_content(text):
                return False

            error_patterns = [
                r'[|]{2,}',
                r'[l1]{3,}',
                r'[o0]{3,}',
                r'[rn]{3,}',
                r'\s{3,}'
            ]

            for pattern in error_patterns:
                if re.search(pattern, text):
                    return False

            return True

        except Exception as e:
            logger.error(f"Error checking OCR result quality: {str(e)}")
            return False

    def _process_single_image(self, image_path: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process a single image file"""
        try:
            if not image_path or not os.path.exists(image_path):
                raise ValueError(f"Invalid image path: {image_path}")

            if not isinstance(min_confidence, (int, float)) or not 0 <= min_confidence <= 1:
                raise ValueError("Invalid confidence threshold: min_confidence must be between 0 and 1")

            logger.info(f"Processing single image: {image_path}")

            try:
                img = Image.open(image_path)
            except Exception as e:
                logger.error(f"Error opening image: {str(e)}")
                raise ValueError(f"Failed to open image: {str(e)}")

            try:
                ocr_text = pytesseract.image_to_string(img)
        except Exception as e:
                logger.error(f"Tesseract OCR failed: {str(e)}")
                raise ValueError(f"OCR processing failed: {str(e)}")

            if not self._is_good_ocr_result(ocr_text):
                logger.info("Tesseract OCR yielded poor results, trying Gemini Vision")
                try:
                    response = self.text_extractor.process_with_gemini(
                        image_path,
                        "Extract all text from this document image. Return only the raw text without any formatting."
                    )
                    if not response:
                        raise ValueError("Empty response from Gemini Vision")
                    ocr_text = response.strip()
                except Exception as e:
                    logger.error(f"Gemini Vision OCR failed: {str(e)}")
                    raise ValueError(f"Alternative OCR processing failed: {str(e)}")

            if not ocr_text.strip():
                logger.warning("No text could be extracted from the image")
                return {
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": image_path,
                    "rejection_reason": "No text could be extracted from the image"
                }

            try:
                result = self._process_text_content(ocr_text, image_path, min_confidence)
            except Exception as e:
                logger.error(f"Error processing extracted text: {str(e)}")
                raise ValueError(f"Text processing failed: {str(e)}")

            if result:
                result["processing_method"] = "ocr"
                result["source_file"] = image_path
                return result

            return None

        except Exception as e:
            logger.error(f"Error processing single image: {str(e)}")
            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": image_path,
                "rejection_reason": f"Error processing image: {str(e)}",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "error_traceback": traceback.format_exc()
                }
            }

    def _initialize_document_patterns(self) -> Dict[str, List[str]]:
        """Initialize document patterns dynamically based on categories"""
        patterns = {}

        # Add patterns for each category
        for category, doc_types in self.document_categories.items():
            for doc_type in doc_types:
                patterns[doc_type] = self._generate_patterns_for_type(doc_type)

        return patterns

    def _generate_patterns_for_type(self, doc_type: str) -> List[str]:
        """Generate patterns for a specific document type"""
        base_patterns = []

        # Convert document type to searchable terms
        terms = doc_type.replace('_', ' ').split()

        # Generate basic patterns
        for term in terms:
            base_patterns.extend([
                f'(?i){term}',  # Case insensitive match
                f'(?i){term}s',  # Plural form
                f'(?i){term}ing',  # Gerund form
                f'(?i){term}ed'  # Past tense
            ])

        # Add document-specific patterns based on category
        if doc_type in self.document_categories['identity']:
            base_patterns.extend([
                r'(?i)id\s*card',
                r'(?i)identification',
                r'(?i)identity\s*document',
                r'(?i)official\s*id',
                r'(?i)government\s*issued'
            ])
        elif doc_type in self.document_categories['legal']:
            base_patterns.extend([
                r'(?i)legal\s*document',
                r'(?i)official\s*document',
                r'(?i)notarized',
                r'(?i)certified',
                r'(?i)authorized'
            ])
        elif doc_type in self.document_categories['financial']:
            base_patterns.extend([
                r'(?i)financial\s*document',
                r'(?i)monetary',
                r'(?i)payment',
                r'(?i)transaction',
                r'(?i)account'
            ])
        elif doc_type in self.document_categories['educational']:
            base_patterns.extend([
                r'(?i)educational\s*document',
                r'(?i)academic',
                r'(?i)school',
                r'(?i)university',
                r'(?i)institution'
            ])
        elif doc_type in self.document_categories['medical']:
            base_patterns.extend([
                r'(?i)medical\s*document',
                r'(?i)health',
                r'(?i)patient',
                r'(?i)clinical',
                r'(?i)diagnostic'
            ])
        elif doc_type in self.document_categories['business']:
            base_patterns.extend([
                r'(?i)business\s*document',
                r'(?i)commercial',
                r'(?i)corporate',
                r'(?i)company',
                r'(?i)enterprise'
            ])

        return base_patterns

    def _process_multiple_documents(self, text: str, source_file: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process text that may contain multiple documents"""
        try:
            results = []

            # Split text into potential document chunks
            chunks = self._split_into_chunks(text)
            logger.info(f"Split text into {len(chunks)} potential document chunks")

            for chunk_index, chunk in enumerate(chunks):
                if not chunk.strip():
                    continue

                logger.info(f"Processing chunk {chunk_index + 1} of {len(chunks)}")

                # Process each chunk as a separate document
                result = self._process_text_content(chunk, source_file, min_confidence)

                if result:
                    # Add chunk information to the result
                    result["chunk_index"] = chunk_index + 1
                    result["total_chunks"] = len(chunks)
                    results.append(result)

            return results

    except Exception as e:
            logger.error(f"Error processing multiple documents: {str(e)}")
            return []

# def main():
#     try:
#         input_path ="D:\\imageextractor\\identites\\OIP.pdf"
#         api_key = API_KEY_1
#
#         if not input_path or not os.path.exists(input_path):
#             raise ValueError(f"Invalid input path: {input_path}")
#
#         if not api_key:
#             raise ValueError("API key is required")
#
#         processor = DocumentProcessor(api_key=api_key)
#
#         try:
#             results = processor.process_file(file_path=input_path)
#         except Exception as e:
#             logger.error(f"Error processing file: {str(e)}")
#             raise
#
#         if results:
#             try:
#                 grouped_results = {}
#                 for result in results:
#                     doc_type = result.get("document_type", "unknown")
#                     if doc_type not in grouped_results:
#                         grouped_results[doc_type] = []
#                     grouped_results[doc_type].append(result)
#
#                 try:
#                     with open("results.json", "w", encoding="utf-8") as f:
#                         json.dump({
#                             "total_documents": len(results),
#                             "document_types": list(grouped_results.keys()),
#                             "results": grouped_results
#                         }, f, indent=2, ensure_ascii=False)
#                     print(f"✅ Processed {len(results)} documents and saved to results.json")
#                 except Exception as e:
#                     logger.error(f"Error saving results to file: {str(e)}")
#                     raise
#             except Exception as e:
#                 logger.error(f"Error grouping results: {str(e)}")
#                 raise
#         else:
#             print("⚠️ No valid documents detected")
#
#     except Exception as e:
#         logger.error(f"Error in main function: {str(e)}")
#         print(f"❌ Error: {str(e)}")
#         sys.exit(1)
#
# if __name__ == "__main__":
#     main()