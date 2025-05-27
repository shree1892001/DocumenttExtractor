import fitz
import tempfile
from typing import List, Dict, Any, Optional
import os
import json
import re
from docx import Document
from Common.constants import *
from dataclasses import dataclass
from Extractor.ImageExtractor import ImageTextExtractor
from Extractor.Paddle import flatten_json
from Factories.DocumentFactory import (
    DocumentExtractorFactory, 
    DocumentExtractor,
    TextExtractorFactory,
    BaseTextExtractor,
    DocxTextExtractor,
    PdfTextExtractor,
    ImageTextExtractor,
    TextExtractor
)
import logging
import google.generativeai as genai
from abc import ABC, abstractmethod
import pytesseract
from PIL import Image

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentInfo:
    document_type: str
    confidence_score: float
    extracted_data: Dict[str, Any]
    matched_fields: Dict[str, Any]

class BaseTextExtractor(ABC):
    def __init__(self, api_key: str):
        self.api_key = api_key
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        pass

class TextProcessor:
    def __init__(self, api_key: str):
        self.api_key = api_key
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    def process_text(self, text: str, prompt: str) -> str:
        """Process text using Gemini without image handling"""
        try:
            response = self.model.generate_content([prompt, text])
            return response.text
        except Exception as e:
            logger.error(f"Error processing text with Gemini: {str(e)}")
            raise

class TemplateMatcher:
    # Define document type mapping as a class variable
    document_type_mapping = {
        "aadhaar": "aadhaar_card",
        "aadhaarcard": "aadhaar_card",
        "aadhar": "aadhaar_card",
        "aadharcard": "aadhaar_card",
        "pan": "pan_card",
        "pancard": "pan_card",
        "license": "license",
        "driving license": "license",
        "dl": "license"
    }

    # Define supported file extensions
    SUPPORTED_EXTENSIONS = {'.docx', '.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.bmp'}

    # Define confidence thresholds
    MIN_CONFIDENCE_THRESHOLD = 0.4
    HIGH_CONFIDENCE_THRESHOLD = 0.8

    def __init__(self, api_key: str, templates_dir: str = "D:\\imageextractor\\identites\\Templates"):
        self.templates_dir = templates_dir
        self.api_key = api_key
        self.text_processor = TextProcessor(api_key)
        
        # Validate templates directory
        if not os.path.exists(self.templates_dir):
            logger.error(f"Templates directory does not exist: {self.templates_dir}")
            raise ValueError(f"Templates directory not found: {self.templates_dir}")
            
        # Load templates
        self.templates = self._load_templates()
        if not self.templates:
            logger.error(f"No valid templates found in directory: {self.templates_dir}")
            raise ValueError(f"No valid templates found in directory: {self.templates_dir}")
        
    def _load_templates(self) -> Dict[str, Dict[str, Any]]:
        """Load all templates from the templates directory"""
        templates = {}
        try:
            files = os.listdir(self.templates_dir)
            if not files:
                logger.error(f"Templates directory is empty: {self.templates_dir}")
                return templates
                
            logger.info(f"Found {len(files)} files in templates directory")
            
            for filename in files:
                file_path = os.path.join(self.templates_dir, filename)
                file_ext = os.path.splitext(filename)[1].lower()
                
                # Skip unsupported file formats
                if file_ext not in self.SUPPORTED_EXTENSIONS:
                    logger.debug(f"Skipping unsupported file format: {filename}")
                    continue
                
                if os.path.isfile(file_path):
                    try:
                        # Extract document type from filename
                        doc_type = filename.replace('sample_', '').replace('.docx', '').replace('.pdf', '').replace('.jpg', '').replace('.png', '')
                        doc_type = self._standardize_document_type(doc_type)
                        
                        logger.info(f"Processing template: {filename} as {doc_type}")
                        
                        # Use TextExtractorFactory to get appropriate extractor
                        text_extractor = TextExtractorFactory.create_extractor(file_path, self.api_key)
                        template_text = text_extractor.extract_text(file_path)
                        
                        if template_text:
                            templates[doc_type] = {
                                'content': template_text,
                                'fields': self._extract_fields_from_text(template_text),
                                'structure': template_text
                            }
                            logger.info(f"Successfully loaded template: {filename}")
                        else:
                            logger.warning(f"No text extracted from template: {filename}")
                    except Exception as e:
                        logger.warning(f"Error processing template file {filename}: {str(e)}")
                        continue
            
            if not templates:
                logger.error(f"No valid templates found in directory: {self.templates_dir}")
            else:
                logger.info(f"Successfully loaded {len(templates)} templates")
            
            return templates
            
        except Exception as e:
            logger.error(f"Error loading templates from directory {self.templates_dir}: {str(e)}")
            raise
    
    def _extract_fields_from_text(self, text: str) -> set:
        """Extract potential field names from text"""
        fields = set()
        # Look for text between brackets or after colons
        field_patterns = [
            r'\{([^}]+)\}',  # {field_name}
            r'([^:]+):',     # field_name:
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'  # Capitalized words
        ]
        
        for pattern in field_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    field = match[0] if match[0] else match[1]
                else:
                    field = match
                fields.add(field.strip())
        
        return fields
    
    def _standardize_document_type(self, doc_type: str) -> str:
        """Standardize document type to match factory mappings"""
        doc_type = doc_type.lower().strip()
        mapping = {
            "aadhaar": "aadhaar_card",
            "aadhaarcard": "aadhaar_card",
            "aadhar": "aadhaar_card",
            "aadharcard": "aadhaar_card",
            "pan": "pan_card",
            "pancard": "pan_card",
            "license": "license",
            "driving license": "license",
            "dl": "license",
            "indian_license": "license",
            "florida_license": "license",
            "passport": "passport"
        }
        return mapping.get(doc_type, doc_type)
    
    def _create_template_matching_prompt(self, input_text: str, template_content: str, doc_type: str) -> str:
        """Create a prompt for Gemini to match document against template"""
        return f"""
        You are a document analysis expert specializing in identity document verification. Your task is to analyze if the given document matches the {doc_type} template structure.

        Document Text:
        {input_text}

        Template Structure:
        {template_content}

        Please analyze the document and return a JSON response with:
        1. Whether this document matches the {doc_type} template (true/false)
        2. Confidence score (0-1) based on:
           - Key field presence (0.3 weight): Check if essential fields like name, ID number, etc. are present
           - Content similarity (0.3 weight): Compare the actual content with template
           - Structure match (0.2 weight): Layout and organization similarity
           - Format consistency (0.2 weight): Consistent formatting and style

        For {doc_type} specifically, look for these key indicators:
        {self._get_document_specific_indicators(doc_type)}

        3. Extracted fields based on the template structure
        4. Any additional relevant information found

        Format the response as:
        {{
            "matches_template": true/false,
            "confidence_score": 0.0-1.0,
            "extracted_fields": {{
                "field1": "value1",
                "field2": "value2"
            }},
            "additional_info": "any additional relevant information"
        }}

        Important:
        - Consider partial matches if key fields are present
        - Account for variations in formatting and layout
        - Be lenient with confidence scoring if key identifiers are present
        - Consider a match if confidence score is above 0.4 and key fields are present
        """

    def _get_document_specific_indicators(self, doc_type: str) -> str:
        """Get document-specific indicators for matching"""
        indicators = {
            "aadhaar_card": """
            - 12-digit Aadhaar number
            - Name of the holder
            - Date of birth
            - Gender
            - Address
            - UIDAI logo or text
            - QR code or barcode
            """,
            "license": """
            - License number
            - Name of the holder
            - Date of birth
            - Address
            - License type/class
            - Issue and expiry dates
            - Issuing authority
            """,
            "pan_card": """
            - 10-character PAN number
            - Name of the holder
            - Father's name
            - Date of birth
            - PAN card number format
            - Income Tax Department text
            """
        }
        return indicators.get(doc_type, "Look for standard identity document fields and structure.")

    def match_document(self, file_path: str) -> Dict[str, Any]:
        """Match document against templates using Gemini"""
        if not os.path.exists(file_path):
            raise ValueError(f"Document file not found: {file_path}")
            
        if not self.templates:
            raise ValueError("No templates available for matching")
            
        best_match = {
            'document_type': None,
            'confidence': 0.0,
            'matched_fields': {},
            'additional_info': None
        }

        try:
            text_extractor = TextExtractorFactory.create_extractor(file_path, self.api_key)
            
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    input_text = text_extractor.extract_text(file_path)
                    if input_text:
                        break
                except Exception as e:
                    if attempt == max_retries - 1:
                        raise
                    logger.warning(f"Attempt {attempt + 1} failed: {str(e)}")
                    continue
            
            if not input_text:
                raise ValueError(f"Could not extract text from document: {file_path}")
            
            logger.info(f"Extracted text from document: {len(input_text)} characters")
            
            # Try to match against each template
            for doc_type, template in self.templates.items():
                try:
                    prompt = self._create_template_matching_prompt(input_text, template['content'], doc_type)
                    response = self.text_processor.process_text(input_text, prompt)
                    match_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
                    
                    confidence = match_result.get('confidence_score', 0)
                    logger.info(f"Template match result for {doc_type}: {confidence}")
                    
                    # Update best match if this is better
                    if confidence > best_match['confidence']:
                        best_match = {
                            'document_type': doc_type,
                            'confidence': confidence,
                            'matched_fields': match_result.get('extracted_fields', {}),
                            'additional_info': match_result.get('additional_info')
                        }
                        logger.info(f"Found better match: {doc_type} with confidence {confidence}")
                except Exception as e:
                    logger.warning(f"Error matching template for {doc_type}: {str(e)}")
                    continue
            
            if not best_match['document_type']:
                logger.error("No matching template found for the document")
                raise ValueError("No matching template found for the document")
            
            logger.info(f"Best match found: {best_match['document_type']} with confidence {best_match['confidence']}")
            return best_match
            
        except Exception as e:
            logger.error(f"Error matching document: {str(e)}")
            raise

class DocumentClassifier:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.template_matcher = TemplateMatcher(api_key)

    def identify_document_type(self, file_path: str) -> DocumentInfo:
        try:
            # Match against templates using Gemini
            template_match = self.template_matcher.match_document(file_path)
            
            # Use the best match even if below threshold
            if template_match['document_type']:
                return DocumentInfo(
                    document_type=template_match['document_type'],
                    confidence_score=template_match['confidence'],
                    extracted_data=template_match['matched_fields'],
                    matched_fields=template_match['matched_fields']
                )
            else:
                raise ValueError("Document type not identified")
                
        except Exception as e:
            logger.exception("Failed to classify document")
            raise ValueError(f"Failed to classify document: {str(e)}")

class DocumentProcessor:
    def __init__(self, api_key: str, templates_dir: str = "D:\\imageextractor\\identites\\Templates"):
        self.api_key = api_key
        self.templates_dir = templates_dir
        self.classifier = DocumentClassifier(api_key)
        self.text_extractor = TextExtractor(api_key)
        self.text_processor = TextProcessor(api_key)
        
        # Define pattern-based mapping rules
        self.document_patterns = {
            'license': [
                r'(?i)license|dl|driving|permit|rto|dmv',
                r'(?i)vehicle|motor|transport',
                r'(?i)driver|driving'
            ],
            'aadhaar_card': [
                r'(?i)aadhaar|aadhar|uidai|unique\s*id',
                r'(?i)आधार|यूआईडीएआई'
            ],
            'pan_card': [
                r'(?i)pan|permanent\s*account|income\s*tax',
                r'(?i)tax\s*id|tax\s*number'
            ],
            'passport': [
                r'(?i)passport|travel\s*doc|nationality',
                r'(?i)immigration|border|customs'
            ]
        }
        
        # Compile regex patterns for better performance
        self.compiled_patterns = {
            doc_type: [re.compile(pattern) for pattern in patterns]
            for doc_type, patterns in self.document_patterns.items()
        }

    def _standardize_document_type(self, doc_type: str) -> str:
        """Standardize document type using pattern matching"""
        try:
            doc_type = doc_type.lower().strip()
            
            # First check for exact matches in common mappings
            common_mappings = {
                "aadhaar": "aadhaar_card",
                "aadhaarcard": "aadhaar_card",
                "aadhar": "aadhaar_card",
                "aadharcard": "aadhaar_card",
                "pan": "pan_card",
                "pancard": "pan_card",
                "license": "license",
                "driving license": "license",
                "dl": "license",
                "passport": "passport"
            }
            
            if doc_type in common_mappings:
                return common_mappings[doc_type]
            
            # If no exact match, use pattern matching
            for standardized_type, patterns in self.compiled_patterns.items():
                for pattern in patterns:
                    if pattern.search(doc_type):
                        logger.debug(f"Matched {doc_type} to {standardized_type} using pattern {pattern.pattern}")
                        return standardized_type
            
            # If no pattern matches, try to extract the base type from the filename
            base_type = self._extract_base_type(doc_type)
            if base_type:
                return base_type
            
            # If all else fails, return the original type
            logger.warning(f"No pattern match found for document type: {doc_type}")
            return doc_type
            
        except Exception as e:
            logger.error(f"Error standardizing document type: {str(e)}")
            return doc_type

    def _extract_base_type(self, doc_type: str) -> Optional[str]:
        """Extract base document type from filename or text"""
        try:
            # Remove common prefixes and suffixes
            doc_type = re.sub(r'(?i)sample_|template_|example_', '', doc_type)
            doc_type = re.sub(r'\.(docx|pdf|jpg|png|jpeg)$', '', doc_type)
            
            # Remove location/country prefixes
            doc_type = re.sub(r'^(indian|florida|california|texas|new_york|uk|us|canada)_', '', doc_type)
            
            # Remove version numbers
            doc_type = re.sub(r'_v\d+', '', doc_type)
            
            # Remove dates
            doc_type = re.sub(r'_\d{4}(_\d{2}){0,2}', '', doc_type)
            
            # Clean up any remaining special characters
            doc_type = re.sub(r'[^a-z0-9]', '_', doc_type.lower())
            doc_type = re.sub(r'_+', '_', doc_type)  # Replace multiple underscores with single
            doc_type = doc_type.strip('_')
            
            # Check if the cleaned type matches any of our patterns
            for standardized_type, patterns in self.compiled_patterns.items():
                for pattern in patterns:
                    if pattern.search(doc_type):
                        return standardized_type
            
            return None
            
        except Exception as e:
            logger.error(f"Error extracting base type: {str(e)}")
            return None

    def add_document_pattern(self, doc_type: str, pattern: str) -> None:
        """Add a new pattern for document type detection"""
        try:
            if doc_type not in self.document_patterns:
                self.document_patterns[doc_type] = []
                self.compiled_patterns[doc_type] = []
            
            self.document_patterns[doc_type].append(pattern)
            self.compiled_patterns[doc_type].append(re.compile(pattern))
            logger.info(f"Added new pattern for {doc_type}: {pattern}")
            
        except Exception as e:
            logger.error(f"Error adding document pattern: {str(e)}")

    def get_document_patterns(self) -> Dict[str, List[str]]:
        """Get current document patterns"""
        return {
            doc_type: [pattern.pattern for pattern in patterns]
            for doc_type, patterns in self.compiled_patterns.items()
        }

    def _extract_text_from_docx_images(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from images in a DOCX file and return list of document segments"""
        try:
            doc = Document(file_path)
            document_segments = []
            
            # Process each paragraph for images
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.xpath('.//w:drawing'):
                        # Found an image, extract it
                        with tempfile.TemporaryDirectory() as temp_dir:
                            image_path = os.path.join(temp_dir, f"image_{len(document_segments)}.png")
                            try:
                                # Extract image data
                                image_data = run._element.xpath('.//a:blip/@r:embed')[0]
                                image_part = doc.part.related_parts[image_data]
                                
                                # Save image
                                with open(image_path, 'wb') as f:
                                    f.write(image_part.blob)
                                
                                # Extract text from image using OCR
                                img = Image.open(image_path)
                                ocr_text = pytesseract.image_to_string(img)
                                
                                if not ocr_text.strip():
                                    # If OCR fails, use Gemini Vision
                                    response = self.text_extractor.process_with_gemini(
                                        image_path,
                                        "Extract all text from this document image. Return only the raw text without any formatting."
                                    )
                                    ocr_text = response.strip()
                                
                                if ocr_text.strip():
                                    document_segments.append({
                                        "text": ocr_text.strip(),
                                        "image_path": image_path,
                                        "segment_index": len(document_segments)
                                    })
                                    logger.info(f"Extracted text from image: {len(ocr_text)} characters")
                            except Exception as e:
                                logger.warning(f"Error processing image in DOCX: {str(e)}")
                                continue
            
            if not document_segments:
                logger.warning("No text could be extracted from images in DOCX")
                return []
            
            return document_segments
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX images: {str(e)}")
            return []

    def process_file(self, file_path: str, min_confidence: float = 0.0) -> List[Dict[str, Any]]:
        """Process a file that may contain multiple documents"""
        if file_path.lower().endswith('.pdf'):
            return self._process_pdf(file_path, min_confidence)
        elif file_path.lower().endswith('.docx'):
            return self._process_docx(file_path, min_confidence)
        else:
            result = self._process_single_image(file_path, min_confidence)
            return [result] if result else []

    def _process_docx(self, file_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a DOCX file that may contain multiple documents"""
        results = []
        
        try:
            # First try to extract text directly
            doc = Document(file_path)
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            
            if text_content:
                # Process the text content directly
                text = '\n'.join(text_content)
                result = self._process_text_content(text, file_path, min_confidence)
                if result:
                    results.append(result)
            
            # Extract and process images
            document_segments = self._extract_text_from_docx_images(file_path)
            
            for segment in document_segments:
                # Process the segment text directly
                result = self._process_text_content(segment["text"], file_path, min_confidence)
                if result:
                    result["segment_index"] = segment["segment_index"]
                    results.append(result)
            
            return results
            
        except Exception as e:
            logger.error(f"Error processing DOCX file: {str(e)}")
            return []

    def _process_text_content(self, text: str, source_file: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process text content directly without creating temporary files"""
        try:
            # Create a prompt for document type detection
            detection_prompt = f"""
            Analyze this text and determine what type of document it is (license, aadhaar_card, pan_card, or passport).
            Look for these indicators:
            - Aadhaar: "Aadhaar", "UIDAI", "आधार", "Unique Identification"
            - PAN: "PAN", "Permanent Account Number", "Income Tax"
            - License: "License", "Driving License", "DL", "RTO"
            - Passport: "Passport", "Passport Number", "Nationality"

            Text:
            {text}

            Return a JSON response with:
            {{
                "document_type": "detected_type",
                "confidence": 0.0-1.0,
                "reasoning": "explanation"
            }}
            """
            
            # Use Gemini to detect document type
            response = self.text_processor.process_text(text, detection_prompt)
            detection_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
            
            doc_type = detection_result.get("document_type")
            confidence = detection_result.get("confidence", 0.0)
            
            if not doc_type or confidence < min_confidence:
                logger.warning(f"Could not confidently determine document type (confidence: {confidence})")
                return None
            
            logger.info(f"Detected document type: {doc_type} with confidence {confidence}")
            
            # Get the appropriate extractor
            extractor = DocumentExtractorFactory.get_extractor(doc_type, self.api_key)
            
            # Create extraction prompt
            extraction_prompt = f"""
            You are a document analysis expert. Extract all relevant information from this {doc_type} document.
            
            Document Text:
            {text}
            
            For a {doc_type}, look for and extract these specific fields:
            {self._get_document_specific_fields(doc_type)}
            
            Important:
            1. Extract ALL visible information from the document
            2. If a field is not found, use "None" as the value
            3. For dates, use YYYY-MM-DD format
            4. For numbers, extract them exactly as shown
            5. For names and addresses, preserve the exact spelling and formatting
            6. Be lenient in matching fields - look for variations in field names
            7. If you're not sure about a field, include it anyway with low confidence
            8. Pay special attention to any text that looks like a license number, name, or date
            
            Return the extracted information in JSON format with this structure:
            {{
                "data": {{
                    // All extracted fields
                }},
                "confidence": 0.0-1.0,
                "additional_info": "Any additional relevant information"
            }}
            """
            
            # Process with Gemini
            response = self.text_processor.process_text(text, extraction_prompt)
            extracted_data = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
            logger.info(f"Extracted data: {json.dumps(extracted_data, indent=2)}")
            
            # For low confidence matches, be more lenient with validation
            if confidence < 0.5:
                logger.info("Low confidence match - using lenient validation")
                if self._validate_with_leniency(extracted_data, doc_type):
                    flattened_data = flatten_json(extracted_data)
                    return {
                        "extracted_data": flattened_data,
                        "status": "success",
                        "confidence": confidence,
                        "document_type": doc_type,
                        "source_file": source_file,
                        "validation_level": "lenient"
                    }
            else:
                if extractor.validate_fields(extracted_data):
                    flattened_data = flatten_json(extracted_data)
                    return {
                        "extracted_data": flattened_data,
                        "status": "success",
                        "confidence": confidence,
                        "document_type": doc_type,
                        "source_file": source_file,
                        "validation_level": "strict"
                    }
            
            # Even if validation fails, return the extracted data with a warning
            logger.warning(f"Validation failed for {source_file}, but returning extracted data anyway")
            flattened_data = flatten_json(extracted_data)
            return {
                "extracted_data": flattened_data,
                "status": "partial_success",
                "confidence": confidence,
                "document_type": doc_type,
                "source_file": source_file,
                "validation_level": "failed",
                "warning": "Document validation failed but data was extracted"
            }

        except Exception as e:
            logger.exception(f"Error processing text content: {str(e)}")
            return None

    def _process_pdf(self, pdf_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a PDF file that may contain multiple documents"""
        results = []
        pdf_document = fitz.open(pdf_path)

        try:
            # Process each page separately
            for page_num in range(pdf_document.page_count):
                logger.info(f"Processing page {page_num + 1} of {pdf_document.page_count}")
                
                # Extract images from the page
                page = pdf_document[page_num]
                images = page.get_images(full=True)
                
                if not images:
                    # If no images found, try to extract text directly
                    text = page.get_text()
                    if text.strip():
                        result = self._process_text_content(text, pdf_path, min_confidence)
                        if result:
                            result["page_number"] = page_num + 1
                            results.append(result)
                    continue

                # Process each image on the page
                for img_index, img_info in enumerate(images):
                    with tempfile.TemporaryDirectory() as temp_dir:
                        img_data = pdf_document.extract_image(img_info[0])
                        temp_image_path = os.path.join(temp_dir, f"page_{page_num}_img_{img_index}.png")

                        with open(temp_image_path, "wb") as img_file:
                            img_file.write(img_data["image"])

                        result = self._process_single_image(temp_image_path, min_confidence)
                        if result:
                            result["page_number"] = page_num + 1
                            result["image_index"] = img_index + 1
                            results.append(result)

            return results
        finally:
            pdf_document.close()

    def _process_single_image(self, file_path: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process a single image or document with document type detection"""
        try:
            # First identify the document type
            doc_info = self.classifier.identify_document_type(file_path)
            logger.info(f"Document type identified as: {doc_info.document_type} with confidence {doc_info.confidence_score}")

            # Standardize the document type
            standardized_type = self._standardize_document_type(doc_info.document_type)
            logger.info(f"Standardized document type: {standardized_type}")

            # Get the appropriate extractor
            extractor = DocumentExtractorFactory.get_extractor(standardized_type, self.api_key)
            
            # Extract text from the input document
            input_text = self.text_extractor.extract_text(file_path)
            logger.info(f"Extracted text length: {len(input_text)} characters")
            
            if not input_text.strip():
                logger.error("No text could be extracted from the document")
                return None
            
            logger.debug(f"Extracted text: {input_text}")
            
            # Create a specific extraction prompt for the document type
            extraction_prompt = f"""
            You are a document analysis expert. Extract all relevant information from this {standardized_type} document.
            
            Document Text:
            {input_text}
            
            For a {standardized_type}, look for and extract these specific fields:
            {self._get_document_specific_fields(standardized_type)}
            
            Important:
            1. Extract ALL visible information from the document
            2. If a field is not found, use "None" as the value
            3. For dates, use YYYY-MM-DD format
            4. For numbers, extract them exactly as shown
            5. For names and addresses, preserve the exact spelling and formatting
            6. Be lenient in matching fields - look for variations in field names
            7. If you're not sure about a field, include it anyway with low confidence
            8. Pay special attention to any text that looks like a license number, name, or date
            
            Return the extracted information in JSON format with this structure:
            {{
                "data": {{
                    // All extracted fields
                }},
                "confidence": 0.0-1.0,
                "additional_info": "Any additional relevant information"
            }}
            """
            
            # Process with Gemini
            response = self.text_processor.process_text(input_text, extraction_prompt)
            extracted_data = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
            logger.info(f"Extracted data: {json.dumps(extracted_data, indent=2)}")
            
            # For low confidence matches, be more lenient with validation
            if doc_info.confidence_score < 0.5:
                logger.info("Low confidence match - using lenient validation")
                if self._validate_with_leniency(extracted_data, standardized_type):
                    flattened_data = flatten_json(extracted_data)
                    return {
                        "extracted_data": flattened_data,
                        "status": "success",
                        "confidence": doc_info.confidence_score,
                        "document_type": standardized_type,
                        "source_file": file_path,
                        "validation_level": "lenient"
                    }
            else:
                if extractor.validate_fields(extracted_data):
                    flattened_data = flatten_json(extracted_data)
                    return {
                        "extracted_data": flattened_data,
                        "status": "success",
                        "confidence": doc_info.confidence_score,
                        "document_type": standardized_type,
                        "source_file": file_path,
                        "validation_level": "strict"
                    }
            
            # Even if validation fails, return the extracted data with a warning
            logger.warning(f"Validation failed for {file_path}, but returning extracted data anyway")
            flattened_data = flatten_json(extracted_data)
            return {
                "extracted_data": flattened_data,
                "status": "partial_success",
                "confidence": doc_info.confidence_score,
                "document_type": standardized_type,
                "source_file": file_path,
                "validation_level": "failed",
                "warning": "Document validation failed but data was extracted"
            }

        except Exception as e:
            logger.exception(f"Error processing file {file_path}")
            return None

    def _validate_with_leniency(self, extracted_data: Dict[str, Any], doc_type: str) -> bool:
        """Validate extracted data with lenient rules for low confidence matches"""
        try:
            data = extracted_data.get("data", {})
            
            # Check if we have at least some basic information
            if not data:
                return False
                
            # Document type specific lenient validation
            if doc_type == "license":
                # For license, require at least one of these fields
                required_fields = ["license_number", "name", "date_of_birth"]
                return any(field in data for field in required_fields)
            elif doc_type == "aadhaar_card":
                # For Aadhaar, require at least one of these fields
                required_fields = ["aadhaar_number", "name", "date_of_birth"]
                return any(field in data for field in required_fields)
            elif doc_type == "pan_card":
                # For PAN, require at least one of these fields
                required_fields = ["pan_number", "name", "date_of_birth"]
                return any(field in data for field in required_fields)
            else:
                # For other document types, require at least one field
                return len(data) > 0
                
        except Exception as e:
            logger.error(f"Error in lenient validation: {str(e)}")
            return False

    def _get_document_specific_fields(self, doc_type: str) -> str:
        """Get document-specific fields to extract"""
        fields = {
            "license": """
            - License Number (look for DL number, license number, etc.)
            - Name (full name of the license holder)
            - Date of Birth (DOB, birth date)
            - Address (current address)
            - Valid From (issue date, date of issue)
            - Valid Until (expiry date, valid till)
            - Vehicle Categories (classes of vehicles allowed)
            - Issuing Authority (RTO, DMV, etc.)
            """,
            "aadhaar_card": """
            - Aadhaar Number (12-digit number)
            - Name (full name)
            - Gender
            - Date of Birth
            - Address
            - Father's Name
            - Photo
            """,
            "pan_card": """
            - PAN Number (10-character alphanumeric)
            - Name
            - Father's Name
            - Date of Birth
            - Photo
            """,
            "passport": """
            - Passport Number
            - Surname
            - Given Names
            - Nationality
            - Date of Birth
            - Gender
            - Date of Issue
            - Date of Expiry
            - Place of Issue
            """
        }
        return fields.get(doc_type, "Extract all visible text and information from the document.")

def main():
    input_path = "D:\\imageextractor\\identites\\aadhar_card.jpg"
    api_key = API_KEY

    processor = DocumentProcessor(api_key=api_key)

    results = processor.process_file(file_path=input_path)

    if results:
        # Group results by document type
        grouped_results = {}
        for result in results:
            doc_type = result["document_type"]
            if doc_type not in grouped_results:
                grouped_results[doc_type] = []
            grouped_results[doc_type].append(result)

        # Save results to JSON
        with open("results.json", "w", encoding="utf-8") as f:
            json.dump({
                "total_documents": len(results),
                "document_types": list(grouped_results.keys()),
                "results": grouped_results
            }, f, indent=2, ensure_ascii=False)
        print(f"✅ Processed {len(results)} documents and saved to results.json")
    else:
        print("⚠️ No valid documents detected")

if __name__ == "__main__":
    main()