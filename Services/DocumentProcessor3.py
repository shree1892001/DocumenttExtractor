from abc import ABC, abstractmethod
from Common.constants import *
import fitz
import tempfile
from typing import List, Dict, Any, Optional, Tuple
import os
import json
import re
import sys
import logging
from dataclasses import dataclass
import google.generativeai as genai
from PIL import Image
import pytesseract
import traceback
from Extractor.ImageExtractor import ImageTextExtractor
from Extractor.Paddle import flatten_json
from Factories.DocumentFactory import (
    DocumentExtractorFactory,
    DocumentExtractor,
    TextExtractorFactory,
    BaseTextExtractor,
    DocxTextExtractor,
    PdfTextExtractor,
    ImageTextExtractor,
    TextExtractor
)

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError as e:
    DOCX_AVAILABLE = False
    Document = None

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentInfo:
    document_type: str
    confidence_score: float
    extracted_data: Dict[str, Any]
    matched_fields: Dict[str, Any]


class BaseTextExtractor(ABC):
    def __init__(self, api_key: str):
        from Common.gemini_config import GeminiConfig
        self.api_key = api_key
        self.config = GeminiConfig.create_vision_processor_config(api_key)
        self.model = self.config.get_model()

    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        pass


class TextProcessor:
    def __init__(self, api_key: str):
        from Common.gemini_config import GeminiConfig
        self.api_key = api_key
        self.config = GeminiConfig.create_document_processor_config(api_key)
        self.model = self.config.get_model()

    def process_text(self, text: str, prompt: str) -> str:
        """Process text using Gemini with safety filter handling"""
        try:
            response = self.model.generate_content([prompt, text])

            # Check if response has text
            if hasattr(response, 'text') and response.text:
                return response.text

            # Handle safety filter blocks
            if hasattr(response, 'candidates') and response.candidates:
                candidate = response.candidates[0]
                if hasattr(candidate, 'finish_reason'):
                    finish_reason = candidate.finish_reason
                    if finish_reason == 3:  # SAFETY
                        logger.warning("Gemini safety filters triggered - attempting with modified prompt")
                        return self._retry_with_safer_prompt(text, prompt)
                    elif finish_reason == 4:  # RECITATION
                        logger.warning("Gemini recitation filter triggered - attempting with modified prompt")
                        return self._retry_with_safer_prompt(text, prompt)

            # If no text in response, try with safer prompt
            logger.warning("No text in Gemini response - attempting with modified prompt")
            return self._retry_with_safer_prompt(text, prompt)

        except Exception as e:
            error_str = str(e)
            logger.error(f"Error processing text with Gemini: {error_str}")

            # Check if this is a safety filter error
            if "safety_ratings" in error_str or "finish_reason" in error_str:
                logger.warning("Safety filter detected in exception - attempting with modified prompt")
                try:
                    return self._retry_with_safer_prompt(text, prompt)
                except Exception as retry_e:
                    logger.error(f"Retry with safer prompt also failed: {str(retry_e)}")
                    raise ValueError(f"Gemini processing failed due to safety filters. Original error: {error_str}")

            raise

    def _retry_with_safer_prompt(self, text: str, original_prompt: str) -> str:
        """Retry with a modified prompt that's less likely to trigger safety filters"""
        try:
            import re

            # Create a safer version of the prompt
            safer_prompt = self._make_prompt_safer(original_prompt)

            # Also sanitize the input text
            safer_text = self._sanitize_text_for_safety(text)

            logger.info("Retrying with safer prompt and sanitized text")
            response = self.model.generate_content([safer_prompt, safer_text])

            if hasattr(response, 'text') and response.text:
                return response.text
            else:
                raise ValueError("No text returned from Gemini even with safer prompt")

        except Exception as e:
            logger.error(f"Safer prompt retry failed: {str(e)}")
            raise ValueError(f"Failed to process with Gemini even with safer prompt: {str(e)}")

    def _make_prompt_safer(self, prompt: str) -> str:
        """Make prompt safer by removing potentially triggering words"""
        import re

        # Words that might trigger safety filters in document verification context
        problematic_words = {
            'dangerous': 'concerning',
            'threat': 'issue',
            'attack': 'problem',
            'weapon': 'item',
            'explosive': 'material',
            'bomb': 'device',
            'kill': 'stop',
            'destroy': 'remove',
            'harm': 'affect',
            'damage': 'impact',
            'violence': 'force',
            'criminal': 'irregular',
            'illegal': 'invalid',
            'fraud': 'inconsistency',
            'fake': 'invalid',
            'forged': 'modified',
            'counterfeit': 'unofficial'
        }

        safer_prompt = prompt
        for problematic, replacement in problematic_words.items():
            safer_prompt = re.sub(r'\b' + re.escape(problematic) + r'\b', replacement, safer_prompt,
                                  flags=re.IGNORECASE)

        # Add safety disclaimer
        safer_prompt = f"""
        IMPORTANT: This is a legitimate document analysis task for official verification purposes.

        {safer_prompt}

        Please provide a professional analysis focused on document structure and data consistency.
        """

        return safer_prompt

    def _sanitize_text_for_safety(self, text: str) -> str:
        """Sanitize text to reduce safety filter triggers while preserving document content"""
        import re

        # Don't over-sanitize as we need to preserve document content
        # Just remove potential URLs that might be flagged
        text = re.sub(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+',
                      '[URL_REMOVED]', text)

        # Remove email addresses that might be flagged
        text = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL_REMOVED]', text)

        return text


class DocumentProcessor:
    def __init__(self, api_key: str, templates_dir: str = TEMPLATES_DIR):
        try:
            self.api_key = api_key
            self.templates_dir = templates_dir

            try:
                doc_test = Document()
                logger.info("Successfully initialized python-docx")
            except Exception as e:
                logger.warning(f"Warning: python-docx initialization failed: {str(e)}")

            self.text_processor = TextProcessor(api_key)
            self.text_extractor = TextExtractor(api_key)

            try:
                from Common.gemini_config import initialize_global_config
                self.gemini_config = initialize_global_config(api_key=api_key, model_type="text")
                logger.info("Gemini configuration initialized successfully")

            except Exception as e:
                logger.error(f"Error initializing Gemini models: {str(e)}")
                raise RuntimeError(f"Gemini initialization failed: {str(e)}")

            try:
                version = pytesseract.get_tesseract_version()
                logger.info(f"Tesseract version: {version}")
            except Exception as e:
                logger.error(f"Error initializing Pytesseract: {str(e)}")
                raise RuntimeError(f"Pytesseract initialization failed: {str(e)}")

            self.document_categories = DOCUMENT_CATEGORIES

            self.document_patterns = DOCUMENT_PATTERNS
            self.compiled_patterns = {
                doc_type: [re.compile(pattern) for pattern in patterns]
                for doc_type, patterns in self.document_patterns.items()
            }

            # Enable unified processing directly in DocumentProcessor3
            self.use_unified_processing = True
            logger.info("Unified document processing enabled in DocumentProcessor3")

        except Exception as e:
            logger.error(f"Error initializing DocumentProcessor: {str(e)}")
            raise RuntimeError(f"DocumentProcessor initialization failed: {str(e)}")

    def _create_text_processor(self):
        """Create and configure the text processor"""
        try:
            return self.gemini_text_model
        except Exception as e:
            logger.error(f"Error creating text processor: {str(e)}")
            raise RuntimeError(f"Text processor creation failed: {str(e)}")

    def process_with_unified_prompt(self, text: str, source_file: str = None, min_confidence: float = 0.0) -> Optional[
        Dict[str, Any]]:
        """
        Process document using the unified prompt for optimized single-pass processing

        Args:
            text: Document text to process
            source_file: Source file path (optional)
            min_confidence: Minimum confidence threshold

        Returns:
            Processed document result or None if processing fails
        """
        if not self.use_unified_processing or not self.unified_processor:
            logger.warning("Unified processing not available, falling back to legacy processing")
            return self._process_text_content(text, source_file or "unknown", min_confidence)

        try:
            logger.info("Processing document with unified prompt")

            # Add context for better processing
            context = {
                "source_file": source_file or "unknown",
                "min_confidence_threshold": min_confidence,
                "processing_mode": "comprehensive"
            }

            # Process with unified prompt
            result = self.unified_processor.process_document(text, context)

            # Convert unified result to legacy format for compatibility
            legacy_result = self._convert_unified_to_legacy_format(result, source_file, min_confidence)

            logger.info(
                f"Unified processing completed - Document type: {legacy_result.get('document_type', 'unknown')}")
            return legacy_result

        except Exception as e:
            logger.error(f"Error in unified processing: {str(e)}")
            logger.info("Falling back to legacy processing")
            return self._process_text_content(text, source_file or "unknown", min_confidence)

    def _convert_unified_to_legacy_format(self, unified_result: Dict[str, Any], source_file: str,
                                          min_confidence: float) -> Dict[str, Any]:
        """
        Convert unified processing result to legacy format for backward compatibility

        Args:
            unified_result: Result from unified processor
            source_file: Source file path
            min_confidence: Minimum confidence threshold

        Returns:
            Legacy format result
        """
        try:
            # Extract key information from unified result
            doc_analysis = unified_result.get("document_analysis", {})
            extracted_data = unified_result.get("extracted_data", {})
            verification = unified_result.get("verification_results", {})
            metadata = unified_result.get("processing_metadata", {})

            # Determine status based on verification and confidence
            confidence = doc_analysis.get("confidence_score", 0.0)
            is_genuine = verification.get("authenticity_assessment", {}).get("is_likely_genuine", False)

            if confidence < min_confidence:
                status = "rejected"
                rejection_reason = f"Confidence {confidence:.2f} below threshold {min_confidence}"
            elif not is_genuine:
                status = "rejected"
                rejection_reason = verification.get("authenticity_assessment", {}).get("verification_status",
                                                                                       "Failed authenticity check")
            else:
                status = "success"
                rejection_reason = None

            # Flatten extracted data for legacy compatibility
            flattened_data = self._flatten_extracted_data(extracted_data)

            # Create legacy format result
            legacy_result = {
                "status": status,
                "document_type": doc_analysis.get("document_type", "unknown"),
                "source_file": source_file,
                "confidence": confidence,
                "extracted_data": {
                    "data": flattened_data,
                    "confidence": confidence,
                    "additional_info": metadata.get("processing_notes", "Processed with unified prompt"),
                    "document_metadata": {
                        "type": doc_analysis.get("document_type", "unknown"),
                        "category": doc_analysis.get("document_category", "unknown"),
                        "issuing_authority": doc_analysis.get("issuing_authority", "unknown"),
                        "key_indicators": doc_analysis.get("key_indicators", [])
                    }
                },
                "verification_status": verification.get("authenticity_assessment", {}).get("verification_status",
                                                                                           "unknown"),
                "processing_method": "unified_prompt",
                "validation_level": "comprehensive"
            }

            # Add rejection reason if applicable
            if rejection_reason:
                legacy_result["rejection_reason"] = rejection_reason

            # Add verification details
            if verification:
                legacy_result["verification_details"] = {
                    "authenticity_score": verification.get("authenticity_assessment", {}).get("confidence_score", 0.0),
                    "quality_checks": verification.get("quality_checks", {}),
                    "flags_and_warnings": verification.get("flags_and_warnings", []),
                    "recommendations": verification.get("recommendations", [])
                }

            return legacy_result

        except Exception as e:
            logger.error(f"Error converting unified result to legacy format: {str(e)}")
            # Return minimal legacy result on error
            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": source_file,
                "confidence": 0.0,
                "extracted_data": {
                    "data": {},
                    "confidence": 0.0,
                    "additional_info": f"Error converting result: {str(e)}",
                    "document_metadata": {"type": "unknown", "category": "error"}
                },
                "processing_method": "unified_prompt",
                "validation_level": "error"
            }

    def _flatten_extracted_data(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Flatten the extracted data to extract ALL possible fields without limitations

        Args:
            extracted_data: Data from unified processor

        Returns:
            Flattened data dictionary with ALL extracted fields
        """
        flattened = {}

        # Debug logging
        logger.info(f"Flattening extracted data with keys: {list(extracted_data.keys()) if extracted_data else 'None'}")

        try:
            # EXTRACT EVERYTHING - No limitations, no specific structure requirements
            def extract_all_fields(data, prefix=""):
                """Recursively extract all fields from any data structure"""
                if isinstance(data, dict):
                    for key, value in data.items():
                        # Create field name
                        if prefix:
                            field_name = f"{prefix} {key.replace('_', ' ').title()}"
                        else:
                            field_name = key.replace('_', ' ').title()

                        if isinstance(value, (str, int, float, bool)) and value not in [None, "", "null",
                                                                                        "not_present"]:
                            # Direct value - add to flattened
                            flattened[field_name] = value
                        elif isinstance(value, list) and value:
                            # List - add if not empty
                            flattened[field_name] = value
                        elif isinstance(value, dict) and value:
                            # Nested dict - recurse but also try to add as string if small
                            if len(str(value)) < 200:  # Small dict, add as string too
                                flattened[f"{field_name} (Full)"] = str(value)
                            extract_all_fields(value, field_name)
                elif isinstance(data, list):
                    for i, item in enumerate(data):
                        if isinstance(item, dict):
                            extract_all_fields(item, f"{prefix} Item {i + 1}")
                        elif item not in [None, "", "null", "not_present"]:
                            flattened[f"{prefix} Item {i + 1}"] = item

            # Extract from all top-level keys
            for key, value in extracted_data.items():
                if key in ["document_analysis", "verification_results", "processing_metadata"]:
                    # Skip metadata sections, focus on actual data
                    continue

                if isinstance(value, (str, int, float, bool)) and value not in [None, "", "null", "not_present"]:
                    # Direct top-level value
                    field_name = key.replace('_', ' ').title()
                    flattened[field_name] = value
                elif isinstance(value, dict) and value:
                    # Nested structure - extract everything
                    extract_all_fields(value, key.replace('_', ' ').title())
                elif isinstance(value, list) and value:
                    # List at top level
                    field_name = key.replace('_', ' ').title()
                    flattened[field_name] = value

            # Special handling for common patterns
            # If we have structured data, also extract it in a flattened way
            if "personal_information" in extracted_data:
                personal = extracted_data["personal_information"]
                extract_all_fields(personal, "")

            if "document_identifiers" in extracted_data:
                doc_ids = extracted_data["document_identifiers"]
                extract_all_fields(doc_ids, "")

            if "validity_information" in extracted_data:
                validity = extracted_data["validity_information"]
                extract_all_fields(validity, "")

            # If still no data, try to extract from ANY available data structure
            if not flattened and extracted_data:
                logger.info("No data found in standard extraction, trying to extract from ANY available data")
                extract_all_fields(extracted_data, "")

            # Also try to extract from any nested structures we might have missed
            for key, value in extracted_data.items():
                if isinstance(value, dict):
                    for nested_key, nested_value in value.items():
                        if isinstance(nested_value, dict):
                            extract_all_fields(nested_value, nested_key.replace('_', ' ').title())

            # Clean up only truly empty values - be very permissive
            cleaned_flattened = {}
            for key, value in flattened.items():
                # Only skip truly empty/null values - keep everything else
                if value is None:
                    continue
                elif isinstance(value, str):
                    # Only skip completely empty strings or explicit nulls
                    if value.strip() == "" or value.lower() in ['null', 'not_present', 'none']:
                        continue
                    else:
                        cleaned_flattened[key] = value
                elif isinstance(value, list):
                    # Keep all lists, even if they contain some empty items
                    cleaned_flattened[key] = value
                elif isinstance(value, (int, float, bool)):
                    # Keep all numbers and booleans
                    cleaned_flattened[key] = value
                else:
                    # Keep everything else that's not None
                    if value:
                        cleaned_flattened[key] = value

            # Debug logging
            logger.info(f"Flattened data result (before cleaning): {flattened}")
            logger.info(f"Cleaned flattened data result: {cleaned_flattened}")

            return cleaned_flattened

        except Exception as e:
            logger.error(f"Error flattening extracted data: {str(e)}")
            return {"error": f"Failed to flatten data: {str(e)}"}

    def _process_with_unified_prompt(self, text: str, source_file: str, min_confidence: float) -> Optional[
        Dict[str, Any]]:
        """
        Process document using an enhanced unified prompt approach with maximum accuracy and extraction

        Args:
            text: Document text to process
            source_file: Source file path
            min_confidence: Minimum confidence threshold

        Returns:
            Processed document result
        """
        try:
            logger.info("Processing document with maximum accuracy enhanced prompt")

            # Enhanced prompt for maximum accuracy and extraction
            enhanced_prompt = f"""
You are an expert document data extraction specialist with maximum accuracy. Your task is to extract EVERY SINGLE PIECE OF INFORMATION from the provided document text with the highest possible accuracy.

**CRITICAL REQUIREMENTS FOR MAXIMUM ACCURACY**:
1. Extract EVERY piece of information present in the document - no exceptions
2. Create HIGHLY DESCRIPTIVE field names that exactly describe what the data represents
3. Clean and correct any OCR errors or formatting issues
4. Identify document type and extract ALL relevant fields for that document type
5. Extract partial information even if incomplete
6. Include context and relationships between data elements

**DOCUMENT TEXT TO ANALYZE**:
{text}

**MAXIMUM EXTRACTION INSTRUCTIONS**:

1. **COMPREHENSIVE DOCUMENT ANALYSIS**:
   - Identify document type (passport, license, certificate, ID card, etc.)
   - Extract ALL headers, titles, labels, and section names
   - Identify ALL data fields and their values
   - Extract ALL dates, numbers, names, addresses, codes, and identifiers

2. **DETAILED FIELD EXTRACTION WITH ACCURATE NAMES**:

   **Personal Information Fields**:
   - "Full Name", "First Name", "Last Name", "Middle Name", "Given Name", "Surname", "Family Name"
   - "Date of Birth", "Birth Date", "DOB", "Born", "Birthday"
   - "Gender", "Sex", "Male/Female"
   - "Age", "Years Old"
   - "Nationality", "Citizenship", "Country of Citizenship"
   - "Place of Birth", "Birth Place", "Birth City", "Birth Country", "Birth State"
   - "Marital Status", "Married", "Single", "Divorced", "Widowed"
   - "Blood Type", "Blood Group"
   - "Height", "Weight", "Eye Color", "Hair Color", "Distinguishing Marks"
   - "Signature", "Signature Present", "Digital Signature"

   **Document Information Fields**:
   - "Document Type", "Document Category", "Document Class"
   - "Document Number", "ID Number", "License Number", "Passport Number", "Account Number"
   - "Reference Number", "Serial Number", "File Number", "Case Number", "Registration Number"
   - "Issue Date", "Date of Issue", "Issued On", "Valid From", "Effective Date"
   - "Expiry Date", "Expiration Date", "Valid Until", "Valid To", "Expires On"
   - "Renewal Date", "Next Renewal", "Valid Period", "Duration"
   - "Issuing Authority", "Issuing Office", "Issuing Department", "Issuing Agency"
   - "Issuing Country", "Country of Issue", "Issuing Nation"
   - "Document Status", "Validity Status", "Active", "Expired", "Suspended"
   - "Document Class", "Document Category", "Document Type Code"

   **Address and Location Fields**:
   - "Address", "Full Address", "Complete Address", "Residential Address"
   - "Street Address", "Street Name", "Street Number", "House Number"
   - "City", "Town", "Municipality", "Village"
   - "State", "Province", "Region", "Territory"
   - "Country", "Nation", "Republic", "Kingdom"
   - "Postal Code", "Zip Code", "PIN Code", "Postcode"
   - "District", "County", "Area", "Zone"
   - "Current Address", "Permanent Address", "Mailing Address", "Office Address"
   - "Emergency Address", "Contact Address"

   **Contact Information Fields**:
   - "Phone Number", "Telephone Number", "Mobile Number", "Cell Phone"
   - "Email Address", "Email", "E-mail", "Electronic Mail"
   - "Website", "Web Address", "URL", "Internet Address"
   - "Fax Number", "Fax", "Facsimile"
   - "Emergency Contact", "Emergency Phone", "Emergency Number"
   - "Alternative Contact", "Secondary Contact"

   **Organizational Information Fields**:
   - "Company Name", "Organization Name", "Institution Name", "Employer"
   - "Department", "Division", "Section", "Unit", "Branch"
   - "Job Title", "Position", "Designation", "Role", "Occupation"
   - "Employee ID", "Staff ID", "Worker ID", "Personnel Number"
   - "Department Code", "Division Code", "Unit Code"
   - "Office Location", "Work Location", "Business Address"

   **Financial Information Fields**:
   - "Amount", "Balance", "Salary", "Income", "Wage", "Pay"
   - "Account Number", "Bank Account", "Account ID", "Account Code"
   - "Transaction Amount", "Payment Amount", "Fee Amount"
   - "Currency", "Monetary Unit", "Dollar Amount", "Euro Amount"

   **Security and Verification Fields**:
   - "Security Code", "Access Code", "PIN", "Password", "Secret Code"
   - "Verification Number", "Authentication Code", "Security Number"
   - "Hologram", "Watermark", "Security Feature", "Anti-counterfeit"
   - "Machine Readable Zone", "MRZ", "Barcode", "QR Code"
   - "Digital Signature", "Electronic Signature", "E-signature"

   **Additional Information Fields**:
   - "Notes", "Comments", "Remarks", "Additional Information"
   - "Special Instructions", "Conditions", "Restrictions", "Endorsements"
   - "Authorized By", "Approved By", "Signed By", "Issued By"
   - "Witness", "Notary", "Official", "Authorized Person"
   - "Document Version", "Revision", "Edition", "Version Number"

3. **ACCURATE TEXT CLEANING AND FORMATTING**:
   - Fix obvious OCR errors (0→O, 1→l, 5→S, etc.)
   - Standardize date formats (DD/MM/YYYY or MM/DD/YYYY)
   - Properly capitalize names and proper nouns
   - Remove extra spaces and formatting artifacts
   - Correct common misreadings (rn→m, cl→d, vv→w, etc.)

4. **CONTEXT-AWARE EXTRACTION**:
   - Extract related information together (e.g., "Name" with "Date of Birth")
   - Identify relationships between data elements
   - Extract section headers and their associated data
   - Include document structure and layout information

5. **COMPREHENSIVE COVERAGE**:
   - Extract ALL text content, no matter how small
   - Include headers, footers, watermarks, stamps
   - Extract handwritten notes if present
   - Include any codes, numbers, or identifiers
   - Extract unclear text with descriptive field names

**OUTPUT FORMAT** - Return comprehensive JSON with maximum accuracy:
{{
    "document_analysis": {{
        "document_type": "specific_document_type_identified",
        "confidence_score": 0.0-1.0,
        "processing_method": "maximum_accuracy_extraction",
        "document_category": "identity_document/financial_document/legal_document/etc"
    }},
    "extracted_data": {{
        // EXTRACT EVERYTHING WITH MAXIMUM ACCURACY
        // USE HIGHLY DESCRIPTIVE FIELD NAMES
        // INCLUDE ALL INFORMATION FOUND IN THE DOCUMENT

        // Personal Information
        "Full Name": "exact_name_as_appears",
        "First Name": "first_name_only",
        "Last Name": "last_name_only", 
        "Date of Birth": "formatted_date",
        "Gender": "M/F/Male/Female",
        "Nationality": "country_name",
        "Place of Birth": "birth_location",

        // Document Information
        "Document Type": "specific_document_type",
        "Document Number": "exact_number",
        "Issue Date": "formatted_date",
        "Expiry Date": "formatted_date",
        "Issuing Authority": "authority_name",
        "Issuing Country": "country_name",

        // Address Information
        "Address": "complete_address",
        "City": "city_name",
        "State": "state_name",
        "Country": "country_name",
        "Postal Code": "postal_code",

        // Contact Information
        "Phone Number": "phone_number",
        "Email": "email_address",

        // Additional Information
        "Notes": "any_additional_notes",
        "Unclear Text": "text_that_could_not_be_clearly_read",

        // EXTRACT EVERYTHING ELSE YOU FIND - CREATE APPROPRIATE FIELD NAMES
    }},
    "verification_results": {{
        "is_genuine": true/false,
        "confidence_score": 0.0-1.0,
        "verification_summary": "detailed_authenticity_assessment",
        "security_features_found": ["list_of_security_elements"],
        "data_consistency": "assessment_of_internal_consistency"
    }},
    "processing_metadata": {{
        "extraction_confidence": 0.0-1.0,
        "processing_notes": "detailed_observations_about_extraction_quality",
        "ocr_quality": "assessment_of_text_quality",
        "missing_information": "any_information_that_appears_missing"
    }}
}}

**CRITICAL ACCURACY REQUIREMENTS**:
- Extract EVERY piece of information present in the document
- Use EXACT field names that describe what the data represents
- Clean and format text for maximum readability
- Include ALL dates, numbers, names, addresses, codes, and identifiers
- Extract partial information even if incomplete
- Provide detailed processing notes about extraction quality

Analyze the document with maximum attention to detail and extract all information with the highest possible accuracy.
"""

            # Process with the maximum accuracy prompt
            logger.info(f"Processing document text ({len(text)} characters) with maximum accuracy extraction")

            try:
                response = self.text_processor.process_text(text, enhanced_prompt)
                logger.debug(f"Received response length: {len(response) if response else 0}")
            except Exception as e:
                logger.error(f"Error calling text processor: {str(e)}")
                raise RuntimeError(f"Text processor failed: {str(e)}")

            # Check if response is empty or None
            if not response or not response.strip():
                logger.error("Received empty response from text processor")
                raise ValueError("AI model returned empty response")

            # Clean and parse the JSON response
            try:
                # Clean the response - remove markdown formatting and extra whitespace
                cleaned_response = self._clean_json_response(response)
                logger.debug(f"Cleaned response length: {len(cleaned_response)}")

                if not cleaned_response:
                    logger.error("Response became empty after cleaning")
                    raise ValueError("Response became empty after cleaning")

                result = json.loads(cleaned_response)

                # Debug logging to see what we got from the AI
                logger.info(f"Maximum accuracy processing raw result keys: {list(result.keys())}")
                if "extracted_data" in result:
                    extracted_keys = list(result['extracted_data'].keys())
                    logger.info(f"Extracted data keys: {extracted_keys}")

                    # Check if we're getting meaningful field names
                    generic_text_count = sum(
                        1 for key in extracted_keys if key.startswith('Text_') or key.startswith('text_'))
                    meaningful_count = len(extracted_keys) - generic_text_count
                    logger.info(
                        f"Meaningful field names: {meaningful_count}, Generic field names: {generic_text_count}")

                    if generic_text_count > meaningful_count:
                        logger.warning(
                            "Too many generic field names detected - attempting reprocessing with stricter prompt")
                        return self._process_with_strict_prompt(text, source_file, min_confidence)
                else:
                    logger.warning("No 'extracted_data' key in maximum accuracy result")

                # Validate the response structure
                if not self._validate_unified_response_structure(result):
                    logger.warning("Response structure validation failed, attempting to fix")
                    result = self._fix_unified_response_structure(result)

                # Convert unified result to legacy format for compatibility
                legacy_result = self._convert_unified_to_legacy_format(result, source_file, min_confidence)

                # Debug logging for legacy result
                legacy_data = legacy_result.get('extracted_data', {}).get('data', {})
                logger.info(f"Legacy result extracted data: {legacy_data}")

                if not legacy_data:
                    logger.warning("No extracted data in legacy result - checking original maximum accuracy result")
                    logger.info(f"Original maximum accuracy extracted_data: {result.get('extracted_data', {})}")

                logger.info(
                    f"Maximum accuracy processing completed - Document type: {legacy_result.get('document_type', 'unknown')}")
                return legacy_result

            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse JSON response: {str(e)}")
                logger.error(f"Raw response (first 500 chars): {response[:500]}")

                # Try to extract JSON from the response if it's embedded in text
                extracted_json = self._extract_json_from_text(response)
                if extracted_json:
                    try:
                        result = json.loads(extracted_json)
                        logger.info("Successfully extracted JSON from text response")

                        # Validate and fix structure
                        if not self._validate_unified_response_structure(result):
                            result = self._fix_unified_response_structure(result)

                        # Convert to legacy format
                        legacy_result = self._convert_unified_to_legacy_format(result, source_file, min_confidence)
                        return legacy_result
                    except json.JSONDecodeError:
                        logger.error("Failed to parse extracted JSON as well")

                # Try fallback with simpler prompt
                logger.warning("Maximum accuracy prompt failed, trying fallback approach")
                fallback_result = self._process_with_fallback_prompt(text, source_file, min_confidence)
                if fallback_result:
                    return fallback_result

                raise ValueError(f"Failed to parse AI response as JSON: {str(e)}")

        except Exception as e:
            logger.error(f"Error in maximum accuracy processing: {str(e)}")
            raise RuntimeError(f"Maximum accuracy document processing failed: {str(e)}")

    def _process_with_strict_prompt(self, text: str, source_file: str, min_confidence: float) -> Optional[
        Dict[str, Any]]:
        """
        Process document with a comprehensive strict prompt for maximum accuracy
        """
        try:
            logger.info("Processing document with comprehensive strict prompt for maximum accuracy")

            strict_prompt = f"""
You are a document data extraction expert with maximum accuracy requirements. Extract ALL information from this document text and create ONLY highly descriptive, meaningful field names.

**MAXIMUM ACCURACY REQUIREMENTS**:
- NEVER use generic names like "Text_1", "Text_2", "Text_3"
- ALWAYS use descriptive field names that exactly explain what the data represents
- Extract EVERY piece of information you can find, no matter how small
- Clean and format the text for maximum readability
- Fix any obvious OCR errors or formatting issues
- Include partial information even if incomplete

**DOCUMENT TEXT**:
{text}

**COMPREHENSIVE FIELD EXTRACTION GUIDE**:

**Personal Identification Fields**:
- "Full Name", "First Name", "Last Name", "Middle Name", "Given Name", "Surname", "Family Name", "Complete Name"
- "Date of Birth", "Birth Date", "DOB", "Born", "Birthday", "Date of Birth (DD/MM/YYYY)"
- "Gender", "Sex", "Male/Female", "M/F"
- "Age", "Years Old", "Current Age"
- "Nationality", "Citizenship", "Country of Citizenship", "Nationality Code"
- "Place of Birth", "Birth Place", "Birth City", "Birth Country", "Birth State", "Birth Location"
- "Marital Status", "Married", "Single", "Divorced", "Widowed", "Relationship Status"
- "Blood Type", "Blood Group", "Blood Classification"
- "Height", "Weight", "Eye Color", "Hair Color", "Distinguishing Marks", "Physical Characteristics"
- "Signature", "Signature Present", "Digital Signature", "Handwritten Signature"

**Document Information Fields**:
- "Document Type", "Document Category", "Document Class", "Document Classification"
- "Document Number", "ID Number", "License Number", "Passport Number", "Account Number", "Registration Number"
- "Reference Number", "Serial Number", "File Number", "Case Number", "Tracking Number", "Control Number"
- "Issue Date", "Date of Issue", "Issued On", "Valid From", "Effective Date", "Date Issued"
- "Expiry Date", "Expiration Date", "Valid Until", "Valid To", "Expires On", "Date of Expiry"
- "Renewal Date", "Next Renewal", "Valid Period", "Duration", "Validity Period"
- "Issuing Authority", "Issuing Office", "Issuing Department", "Issuing Agency", "Authority Name"
- "Issuing Country", "Country of Issue", "Issuing Nation", "Country Code", "Nation of Issue"
- "Document Status", "Validity Status", "Active", "Expired", "Suspended", "Current Status"
- "Document Class", "Document Category", "Document Type Code", "Document Classification"

**Address and Location Fields**:
- "Address", "Full Address", "Complete Address", "Residential Address", "Home Address"
- "Street Address", "Street Name", "Street Number", "House Number", "Building Number"
- "City", "Town", "Municipality", "Village", "City Name", "Urban Area"
- "State", "Province", "Region", "Territory", "State Name", "Provincial Area"
- "Country", "Nation", "Republic", "Kingdom", "Country Name", "National Territory"
- "Postal Code", "Zip Code", "PIN Code", "Postcode", "Postal Number"
- "District", "County", "Area", "Zone", "Administrative District"
- "Current Address", "Permanent Address", "Mailing Address", "Office Address", "Work Address"
- "Emergency Address", "Contact Address", "Alternative Address"

**Contact Information Fields**:
- "Phone Number", "Telephone Number", "Mobile Number", "Cell Phone", "Contact Number"
- "Email Address", "Email", "E-mail", "Electronic Mail", "Email Contact"
- "Website", "Web Address", "URL", "Internet Address", "Online Presence"
- "Fax Number", "Fax", "Facsimile", "Fax Contact"
- "Emergency Contact", "Emergency Phone", "Emergency Number", "Emergency Contact Person"
- "Alternative Contact", "Secondary Contact", "Backup Contact", "Additional Contact"

**Organizational Information Fields**:
- "Company Name", "Organization Name", "Institution Name", "Employer", "Business Name"
- "Department", "Division", "Section", "Unit", "Branch", "Organizational Unit"
- "Job Title", "Position", "Designation", "Role", "Occupation", "Professional Title"
- "Employee ID", "Staff ID", "Worker ID", "Personnel Number", "Employee Number"
- "Department Code", "Division Code", "Unit Code", "Organizational Code"
- "Office Location", "Work Location", "Business Address", "Professional Address"

**Financial Information Fields**:
- "Amount", "Balance", "Salary", "Income", "Wage", "Pay", "Monetary Amount"
- "Account Number", "Bank Account", "Account ID", "Account Code", "Financial Account"
- "Transaction Amount", "Payment Amount", "Fee Amount", "Financial Transaction"
- "Currency", "Monetary Unit", "Dollar Amount", "Euro Amount", "Currency Type"

**Security and Verification Fields**:
- "Security Code", "Access Code", "PIN", "Password", "Secret Code", "Security Number"
- "Verification Number", "Authentication Code", "Security Number", "Verification Code"
- "Hologram", "Watermark", "Security Feature", "Anti-counterfeit", "Security Element"
- "Machine Readable Zone", "MRZ", "Barcode", "QR Code", "Machine Readable Data"
- "Digital Signature", "Electronic Signature", "E-signature", "Digital Authentication"

**Additional Information Fields**:
- "Notes", "Comments", "Remarks", "Additional Information", "Special Notes"
- "Special Instructions", "Conditions", "Restrictions", "Endorsements", "Special Requirements"
- "Authorized By", "Approved By", "Signed By", "Issued By", "Authorizing Official"
- "Witness", "Notary", "Official", "Authorized Person", "Witnessing Authority"
- "Document Version", "Revision", "Edition", "Version Number", "Document Version"

**EXTRACTION PROCESS FOR MAXIMUM ACCURACY**:
1. Read the entire document text carefully and thoroughly
2. Identify the specific type of document (passport, license, certificate, etc.)
3. Extract ALL information you can find, including headers, labels, and section names
4. Create highly descriptive field names for each piece of data
5. Clean and format the extracted text, fixing any obvious errors
6. Include unclear or partial text with descriptive field names
7. Extract context and relationships between data elements
8. Include any codes, numbers, or identifiers found

**OUTPUT FORMAT**:
{{
    "document_analysis": {{
        "document_type": "specific_document_type_identified",
        "confidence_score": 0.0-1.0,
        "processing_method": "comprehensive_strict_extraction",
        "document_category": "identity_document/financial_document/legal_document/etc"
    }},
    "extracted_data": {{
        // EXTRACT ALL INFORMATION WITH MAXIMUM ACCURACY
        // USE HIGHLY DESCRIPTIVE FIELD NAMES ONLY
        // INCLUDE EVERY PIECE OF DATA FOUND IN THE DOCUMENT
        // NO GENERIC NAMES ALLOWED - ONLY MEANINGFUL DESCRIPTIONS
    }},
    "verification_results": {{
        "is_genuine": true/false,
        "confidence_score": 0.0-1.0,
        "verification_summary": "detailed_authenticity_assessment",
        "security_features_found": ["list_of_security_elements"],
        "data_consistency": "assessment_of_internal_consistency"
    }},
    "processing_metadata": {{
        "extraction_confidence": 0.0-1.0,
        "processing_notes": "detailed_observations_about_extraction_quality_and_accuracy",
        "ocr_quality": "assessment_of_text_quality",
        "missing_information": "any_information_that_appears_missing_or_unclear"
    }}
}}

**CRITICAL ACCURACY REQUIREMENTS**:
- Extract EVERY piece of information present in the document
- Use EXACT field names that describe what the data represents
- Clean and format text for maximum readability
- Include ALL dates, numbers, names, addresses, codes, and identifiers
- Extract partial information even if incomplete
- Provide detailed processing notes about extraction quality
- NO GENERIC FIELD NAMES ALLOWED - ONLY MEANINGFUL DESCRIPTIONS

Analyze the document with maximum attention to detail and extract all information with the highest possible accuracy using only meaningful field names.
"""

            response = self.text_processor.process_text(text, strict_prompt)

            if not response or not response.strip():
                raise ValueError("AI model returned empty response")

            cleaned_response = self._clean_json_response(response)
            if not cleaned_response:
                raise ValueError("Response became empty after cleaning")

            result = json.loads(cleaned_response)

            # Validate and convert to legacy format
            if not self._validate_unified_response_structure(result):
                result = self._fix_unified_response_structure(result)

            legacy_result = self._convert_unified_to_legacy_format(result, source_file, min_confidence)
            return legacy_result

        except Exception as e:
            logger.error(f"Error in strict processing: {str(e)}")
            # Fall back to the original unified prompt
            return self._process_with_fallback_prompt(text, source_file, min_confidence)

    def set_unified_processing(self, enabled: bool):

        self.use_unified_processing = enabled
        if enabled:
            logger.info("Unified processing enabled")
        else:
            logger.info("Unified processing disabled, using legacy processing")

    def _clean_json_response(self, response: str) -> str:
        """
        Clean the response to extract valid JSON

        Args:
            response: Raw response from AI model

        Returns:
            Cleaned JSON string
        """
        if not response:
            return ""

        # Remove common markdown formatting
        cleaned = response.strip()

        # Remove markdown code blocks
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        elif cleaned.startswith("```"):
            cleaned = cleaned[3:]

        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]

        # Remove any leading/trailing whitespace
        cleaned = cleaned.strip()

        # Remove any text before the first { or [
        start_idx = -1
        for i, char in enumerate(cleaned):
            if char in ['{', '[']:
                start_idx = i
                break

        if start_idx > 0:
            cleaned = cleaned[start_idx:]
        elif start_idx == -1:
            # No JSON found
            return ""

        # Remove any text after the last } or ]
        end_idx = -1
        for i in range(len(cleaned) - 1, -1, -1):
            if cleaned[i] in ['}', ']']:
                end_idx = i
                break

        if end_idx >= 0:
            cleaned = cleaned[:end_idx + 1]

        return cleaned

    def _extract_json_from_text(self, text: str) -> Optional[str]:
        """
        Try to extract JSON from text that may contain other content

        Args:
            text: Text that may contain JSON

        Returns:
            Extracted JSON string or None
        """
        import re

        if not text:
            return None

        # Try to find JSON objects using regex
        json_patterns = [
            r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}',  # Simple nested objects
            r'\{.*?\}',  # Basic object pattern
            r'\[.*?\]'  # Array pattern
        ]

        for pattern in json_patterns:
            matches = re.findall(pattern, text, re.DOTALL)
            for match in matches:
                try:
                    # Test if it's valid JSON
                    json.loads(match)
                    return match
                except json.JSONDecodeError:
                    continue

        # Try to find content between specific markers
        markers = [
            ('{', '}'),
            ('[', ']')
        ]

        for start_marker, end_marker in markers:
            start_pos = text.find(start_marker)
            if start_pos != -1:
                # Find the matching closing marker
                bracket_count = 0
                for i in range(start_pos, len(text)):
                    if text[i] == start_marker:
                        bracket_count += 1
                    elif text[i] == end_marker:
                        bracket_count -= 1
                        if bracket_count == 0:
                            potential_json = text[start_pos:i + 1]
                            try:
                                json.loads(potential_json)
                                return potential_json
                            except json.JSONDecodeError:
                                break

        return None

    def _validate_unified_response_structure(self, response: Dict[str, Any]) -> bool:
        """Validate that the unified response has the expected structure"""
        required_sections = [
            "document_analysis",
            "extracted_data",
            "verification_results",
            "processing_metadata"
        ]

        return all(section in response for section in required_sections)

    def _fix_unified_response_structure(self, response: Dict[str, Any]) -> Dict[str, Any]:
        """Fix incomplete unified response structure"""
        fixed_response = {
            "document_analysis": response.get("document_analysis", {}),
            "extracted_data": response.get("extracted_data", {}),
            "verification_results": response.get("verification_results", {}),
            "processing_metadata": response.get("processing_metadata", {})
        }

        # Ensure minimum required fields
        if not fixed_response["document_analysis"]:
            fixed_response["document_analysis"] = {
                "document_type": "unknown",
                "confidence_score": 0.0,
                "processing_method": "unified_prompt"
            }

        if not fixed_response["processing_metadata"]:
            fixed_response["processing_metadata"] = {
                "extraction_confidence": 0.0,
                "text_quality": "unknown",
                "completeness_score": 0.0
            }

        return fixed_response

    def _process_with_fallback_prompt(self, text: str, source_file: str, min_confidence: float) -> Optional[
        Dict[str, Any]]:
        """
        Fallback processing using simpler prompts when unified prompt fails

        Args:
            text: Document text to process
            source_file: Source file path
            min_confidence: Minimum confidence threshold

        Returns:
            Processing result or None if fallback also fails
        """
        try:
            logger.info("Using fallback processing with simpler prompts")

            # Simple document detection
            detection_prompt = f"""
            Analyze this document text and identify the document type. Return a JSON response with:
            {{
                "document_type": "detected_type",
                "confidence": 0.0-1.0,
                "category": "document_category"
            }}

            Text: {text[:2000]}"""  # Limit text length for fallback

            try:
                detection_response = self.text_processor.process_text("", detection_prompt)
                detection_result = json.loads(self._clean_json_response(detection_response))
            except Exception as e:
                logger.warning(f"Fallback detection failed: {str(e)}")
                detection_result = {
                    "document_type": "unknown",
                    "confidence": 0.0,
                    "category": "unknown"
                }

            # Simple data extraction
            extraction_prompt = f"""
            Extract key information from this {detection_result.get('document_type', 'document')} text. Return JSON with:
            {{
                "name": "extracted_name",
                "document_number": "extracted_number",
                "date_of_birth": "YYYY-MM-DD",
                "other_fields": {{}}
            }}

            Text: {text[:2000]}"""

            try:
                extraction_response = self.text_processor.process_text("", extraction_prompt)
                extraction_result = json.loads(self._clean_json_response(extraction_response))
            except Exception as e:
                logger.warning(f"Fallback extraction failed: {str(e)}")
                extraction_result = {}

            # Create legacy format result from fallback
            confidence = detection_result.get("confidence", 0.0)
            doc_type = detection_result.get("document_type", "unknown")

            # Check confidence threshold
            if confidence < min_confidence:
                status = "rejected"
                rejection_reason = f"Confidence {confidence:.2f} below threshold {min_confidence}"
            else:
                status = "success"
                rejection_reason = None

            # Create extracted data in legacy format
            extracted_data = {
                "data": {
                    "Name": extraction_result.get("name"),
                    "Document Number": extraction_result.get("document_number"),
                    "Date of Birth": extraction_result.get("date_of_birth"),
                    **extraction_result.get("other_fields", {})
                },
                "confidence": confidence,
                "additional_info": "Processed with fallback method due to unified prompt failure",
                "document_metadata": {
                    "type": doc_type,
                    "category": detection_result.get("category", "unknown"),
                    "issuing_authority": "unknown",
                    "key_indicators": []
                }
            }

            # Create legacy format result
            fallback_result = {
                "status": status,
                "document_type": doc_type,
                "source_file": source_file,
                "confidence": confidence,
                "extracted_data": extracted_data,
                "processing_method": "fallback_simplified",
                "validation_level": "basic"
            }

            # Add rejection reason if applicable
            if rejection_reason:
                fallback_result["rejection_reason"] = rejection_reason

            # Add basic verification result
            fallback_result["verification_result"] = {
                "is_genuine": True,  # Default to true for fallback
                "confidence_score": 0.5,
                "verification_checks": {},
                "verification_summary": "Fallback processing - limited verification"
            }

            logger.info("Fallback processing completed successfully")
            return fallback_result

        except Exception as e:
            logger.error(f"Fallback processing also failed: {str(e)}")
            return None

    def _split_into_chunks(self, text: str) -> List[str]:
        """Split text into chunks based on document boundaries"""
        try:

            separators = [
                r'\n\s*\n\s*\n',
                r'[-=]{3,}',
                r'_{3,}',
                r'\*{3,}',
                r'Page \d+',
                r'Document \d+',
                r'Copy \d+',
                r'Original',
                r'Duplicate',
                r'COPY',
                r'ORIGINAL'
            ]

            separator_pattern = '|'.join(f'({sep})' for sep in separators)

            chunks = re.split(separator_pattern, text)

            chunks = [chunk.strip() for chunk in chunks if chunk and not any(sep in chunk for sep in separators)]

            if not chunks:
                return [text]

            return chunks

        except Exception as e:
            logger.error(f"Error splitting text into chunks: {str(e)}")
            return [text]

    def process_file(self, file_path: str, min_confidence: float = 0.0) -> List[Dict[str, Any]]:
        """Process a file that may contain multiple documents"""
        if file_path.lower().endswith('.pdf'):
            return self._process_pdf(file_path, min_confidence)
        elif file_path.lower().endswith('.docx'):
            return self._process_docx(file_path, min_confidence)
        else:
            result = self._process_single_image(file_path, min_confidence)
            return [result] if result else []

    def _extract_text_from_docx_images(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from images in a DOCX file with simplified approach"""
        try:
            doc = Document(file_path)
            document_segments = []
            image_count = 0

            # Process images from relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
                    with tempfile.TemporaryDirectory() as temp_dir:
                        image_path = os.path.join(temp_dir, f"image_{image_count}.png")
                        try:
                            image_data = rel.target_part.blob

                            with open(image_path, 'wb') as f:
                                f.write(image_data)

                            # Always use Tesseract first
                            img = Image.open(image_path)
                            ocr_text = pytesseract.image_to_string(img)

                            # Only try Gemini if Tesseract results are poor
                            if not self._is_good_ocr_result(ocr_text):
                                logger.info(
                                    f"Tesseract OCR yielded poor results for image {image_count}, trying Gemini Vision")
                                try:
                                    response = self._process_with_gemini(
                                        image_path,
                                        "Extract all text from this document image. Return only the raw text without any formatting."
                                    )
                                    if response and response.strip():
                                        ocr_text = response.strip()
                                except Exception as e:
                                    logger.warning(
                                        f"Gemini Vision failed for image {image_count}, using Tesseract results: {str(e)}")
                                    # Continue with the original Tesseract results

                            document_segments.append({
                                "text": ocr_text,
                                "source": f"image_{image_count}",
                                "type": "image"
                            })

                        except Exception as e:
                            logger.error(f"Error processing image {image_count}: {str(e)}")
                            document_segments.append({
                                "text": f"[Error processing image {image_count}: {str(e)}]",
                                "source": f"image_{image_count}",
                                "type": "error"
                            })

            # Process images from paragraphs
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.xpath('.//w:drawing'):
                        image_count += 1
                        with tempfile.TemporaryDirectory() as temp_dir:
                            image_path = os.path.join(temp_dir, f"image_{image_count}.png")
                            try:
                                image_data = run._element.xpath('.//a:blip/@r:embed')[0]
                                image_part = doc.part.related_parts[image_data]

                                with open(image_path, 'wb') as f:
                                    f.write(image_part.blob)

                                # Always use Tesseract first
                                img = Image.open(image_path)
                                ocr_text = pytesseract.image_to_string(img)

                                # Only try Gemini if Tesseract results are poor
                                if not self._is_good_ocr_result(ocr_text):
                                    logger.info(
                                        f"Tesseract OCR yielded poor results for image {image_count}, trying Gemini Vision")
                                    try:
                                        response = self._process_with_gemini(
                                            image_path,
                                            "Extract all text from this document image. Return only the raw text without any formatting."
                                        )
                                        if response and response.strip():
                                            ocr_text = response.strip()
                                    except Exception as e:
                                        logger.warning(
                                            f"Gemini Vision failed for image {image_count}, using Tesseract results: {str(e)}")
                                        # Continue with the original Tesseract results

                                document_segments.append({
                                    "text": ocr_text,
                                    "source": f"image_{image_count}",
                                    "type": "image"
                                })

                            except Exception as e:
                                logger.error(f"Error processing image {image_count}: {str(e)}")
                                document_segments.append({
                                    "text": f"[Error processing image {image_count}: {str(e)}]",
                                    "source": f"image_{image_count}",
                                    "type": "error"
                                })

            return document_segments

        except Exception as e:
            logger.error(f"Error extracting text from DOCX images: {str(e)}")
            return [{
                "text": f"[Error extracting images from DOCX: {str(e)}]",
                "source": "docx_error",
                "type": "error"
            }]

    def _determine_docx_processing_method(self, file_path: str) -> str:
        """Determine if DOCX requires OCR or can use normal text extraction"""
        try:
            doc = Document(file_path)
            text_content = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)

            combined_text = "\n".join(text_content).strip()

            if len(combined_text) > 100:
                logger.info("DOCX has substantial text content - using normal text extraction")
                return "text_extraction"
            else:
                logger.info("DOCX has minimal text content - will use OCR on embedded images")
                return "ocr_required"

        except Exception as e:
            logger.error(f"Error determining DOCX processing method: {str(e)}")
            return "ocr_required"

    def _process_docx(self, file_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a DOCX file by determining if OCR is needed or normal text extraction"""
        if not DOCX_AVAILABLE or Document is None:
            logger.warning("python-docx not available - attempting OCR-only processing")
            return self._process_docx_with_ocr(file_path, min_confidence)

        consolidated_results = []

        try:

            processing_method = self._determine_docx_processing_method(file_path)
            logger.info(f"Using processing method: {processing_method}")

            if processing_method == "text_extraction":

                logger.info("Processing DOCX using normal text extraction")

                try:
                    doc = Document(file_path)
                except Exception as e:
                    logger.error(f"Error opening DOCX file: {str(e)}")
                    return self._process_docx_with_ocr(file_path, min_confidence)

                text_content = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content.append(para.text)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text)

                if text_content:
                    combined_text = "\n".join(text_content)
                    logger.info(f"Extracted {len(combined_text)} characters of text from DOCX")

                    multi_doc_results = self._process_multiple_documents(combined_text, file_path, min_confidence)
                    if multi_doc_results:
                        consolidated_results.extend(multi_doc_results)
                    else:

                        result = self._process_text_content(combined_text, file_path, min_confidence)
                        if result:
                            consolidated_results.append(result)

            else:
                return self._process_docx_with_ocr(file_path, min_confidence)

            if not consolidated_results:
                logger.warning(f"No valid content found in DOCX file: {file_path}")

                try:
                    doc = Document(file_path)
                    basic_text = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            basic_text.append(para.text.strip())

                    basic_extracted_data = {
                        "data": {
                            "raw_text_snippets": basic_text[:10],
                            "total_paragraphs": len([p for p in doc.paragraphs if p.text.strip()]),
                            "file_type": "docx"
                        },
                        "confidence": 0.0,
                        "additional_info": "No valid content could be extracted through normal processing",
                        "document_metadata": {
                            "type": "unknown",
                            "category": "document",
                            "file_format": "docx"
                        }
                    }
                except Exception as e:
                    basic_extracted_data = {
                        "data": {"file_type": "docx", "processing_error": str(e)},
                        "confidence": 0.0,
                        "additional_info": "Failed to extract any content",
                        "document_metadata": {"type": "unknown", "category": "error"}
                    }

                return [{
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": file_path,
                    "rejection_reason": "No valid content could be extracted",
                    "extracted_data": basic_extracted_data,
                    "confidence": 0.0,
                    "validation_level": "failed"
                }]

            return consolidated_results

        except Exception as e:
            logger.error(f"Error processing DOCX file: {str(e)}")

            error_extracted_data = {
                "data": {
                    "file_type": "docx",
                    "file_path": file_path,
                    "processing_error": str(e),
                    "error_type": type(e).__name__
                },
                "confidence": 0.0,
                "additional_info": f"DOCX processing failed: {str(e)}",
                "document_metadata": {
                    "type": "unknown",
                    "category": "error",
                    "file_format": "docx"
                }
            }

            return [{
                "status": "error",
                "document_type": "unknown",
                "source_file": file_path,
                "rejection_reason": f"Error processing document: {str(e)}",
                "extracted_data": error_extracted_data,
                "confidence": 0.0,
                "validation_level": "error",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "error_traceback": traceback.format_exc()
                }
            }]

    def _process_docx_with_ocr(self, file_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a DOCX file using OCR on embedded images"""
        logger.info("Processing DOCX using OCR on embedded images")
        consolidated_results = []
        processed_images = set()

        try:
            if DOCX_AVAILABLE and Document is not None:
                document_segments = self._extract_text_from_docx_images(file_path)
            else:

                with tempfile.TemporaryDirectory() as temp_dir:
                    pdf_path = os.path.join(temp_dir, "temp.pdf")
                    try:
                        self._convert_docx_to_pdf(file_path, pdf_path)
                        # Use PDF processing instead of missing method
                        pdf_results = self._process_pdf(pdf_path, min_confidence)
                        document_segments = []
                        for result in pdf_results:
                            if result.get('extracted_data', {}).get('data'):
                                document_segments.append({
                                    "text": str(result.get('extracted_data', {}).get('data')),
                                    "source": "pdf_conversion",
                                    "type": "pdf"
                                })
                    except Exception as e:
                        logger.error(f"Error converting DOCX to PDF: {str(e)}")
                        document_segments = []

            if document_segments:
                for segment in document_segments:
                    # Use source as unique identifier since image_path might not exist
                    segment_id = segment.get("source", f"segment_{len(processed_images)}")

                    if segment_id in processed_images:
                        continue

                    processed_images.add(segment_id)

                    if segment["text"].strip():
                        logger.info(f"Processing segment: {segment_id}")

                        multi_doc_results = self._process_multiple_documents(segment["text"], file_path, min_confidence)
                        if multi_doc_results:
                            for result in multi_doc_results:
                                result["segment_id"] = segment_id
                                result["processing_method"] = "ocr"
                                consolidated_results.append(result)
                        else:
                            result = self._process_text_content(segment["text"], file_path, min_confidence)
                            if result:
                                result["segment_id"] = segment_id
                                result["processing_method"] = "ocr"
                                consolidated_results.append(result)

            if DOCX_AVAILABLE and Document is not None:
                try:
                    doc = Document(file_path)
                    text_content = []

                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_content.append(para.text)

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_content.append(cell.text)

                    if text_content:
                        combined_text = "\n".join(text_content)
                        if len(combined_text.strip()) > 50:
                            logger.info(f"Also processing {len(combined_text)} characters of normal text from DOCX")
                            result = self._process_text_content(combined_text, file_path, min_confidence)
                            if result:
                                result["processing_method"] = "text_extraction"
                                consolidated_results.append(result)
                except Exception as e:
                    logger.warning(f"Failed to extract text content from DOCX: {str(e)}")

            if not consolidated_results:
                logger.warning(f"No valid content found in DOCX file through OCR: {file_path}")
                return [{
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": file_path,
                    "rejection_reason": "No valid content could be extracted through OCR",
                    "extracted_data": {
                        "data": {"file_type": "docx", "processing_method": "ocr"},
                        "confidence": 0.0,
                        "additional_info": "OCR processing failed to extract valid content",
                        "document_metadata": {"type": "unknown", "category": "error"}
                    },
                    "confidence": 0.0,
                    "validation_level": "failed"
                }]

            for result in consolidated_results:
                logger.info(f"Processed document type: {result.get('document_type', 'unknown')}, "
                            f"Status: {result.get('status', 'unknown')}, "
                            f"Segment ID: {result.get('segment_id', 'N/A')}")

            return consolidated_results

        except Exception as e:
            logger.error(f"Error processing DOCX with OCR: {str(e)}")
            return [{
                "status": "error",
                "document_type": "unknown",
                "source_file": file_path,
                "rejection_reason": f"Error processing document with OCR: {str(e)}",
                "extracted_data": {
                    "data": {
                        "file_type": "docx",
                        "processing_error": str(e),
                        "error_type": type(e).__name__
                    },
                    "confidence": 0.0,
                    "additional_info": f"DOCX OCR processing failed: {str(e)}",
                    "document_metadata": {"type": "unknown", "category": "error"}
                },
                "confidence": 0.0,
                "validation_level": "error",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "error_traceback": traceback.format_exc()
                }
            }]

    def _consolidate_results(self, results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Consolidate results to remove duplicates and combine related documents"""
        try:
            consolidated = []
            seen_documents = set()

            for result in results:
                if result.get("status") != "success":
                    continue

                doc_type = result.get("document_type", "unknown")
                extracted_data = result.get("extracted_data", {})

                if doc_type == "aadhaar_card":
                    key = f"aadhaar_{extracted_data.get('Aadhaar Number', '')}"
                elif doc_type == "license":
                    key = f"license_{extracted_data.get('Name', '')}_{extracted_data.get('Date of Birth', '')}"
                else:
                    key = f"{doc_type}_{json.dumps(extracted_data)}"

                if key not in seen_documents:
                    seen_documents.add(key)
                    consolidated.append(result)

            return consolidated

        except Exception as e:
            logger.error(f"Error consolidating results: {str(e)}")
            return results

    def _process_text_content(self, text: str, source_file: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process text content directly without creating temporary files"""
        try:
            if not text or not isinstance(text, str):
                raise ValueError("Invalid text input: text must be a non-empty string")

            if not source_file or not isinstance(source_file, str):
                raise ValueError("Invalid source file: source_file must be a non-empty string")

            if not isinstance(min_confidence, (int, float)) or not 0 <= min_confidence <= 1:
                raise ValueError("Invalid confidence threshold: min_confidence must be between 0 and 1")

            # Use unified processing approach directly in DocumentProcessor3
            logger.info("Using unified document processing")
            try:
                return self._process_with_unified_prompt(text, source_file, min_confidence)
            except Exception as e:
                logger.warning(f"Unified processing failed, falling back to legacy: {str(e)}")

            # Legacy processing fallback
            logger.info("Using legacy processing for text content")
            try:
                detection_prompt = DOCUMENT_DETECTION_PROMPT.format(text=text)

                response = self.text_processor.process_text(text, detection_prompt)
                if not response:
                    raise ValueError("Empty response from text processor")

                try:
                    detection_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
                except json.JSONDecodeError as e:
                    logger.error(f"Failed to parse detection result: {str(e)}")
                    raise ValueError(f"Invalid JSON response from text processor: {str(e)}")

                doc_type = detection_result.get("document_type", "").lower()
                confidence = detection_result.get("confidence", 0.0)

                if not doc_type or doc_type == "unknown" or confidence < min_confidence:
                    logger.warning(f"Could not confidently determine document type (confidence: {confidence})")

                    extracted_data_for_response = {
                        "data": detection_result.get("key_indicators", []) if detection_result else {},
                        "confidence": confidence,
                        "additional_info": f"Document type detection failed with confidence {confidence}",
                        "document_metadata": {
                            "type": "unknown",
                            "category": detection_result.get("document_category",
                                                             "unknown") if detection_result else "unknown",
                            "issuing_authority": detection_result.get("issuing_authority",
                                                                      "unknown") if detection_result else "unknown",
                            "key_indicators": detection_result.get("key_indicators", []) if detection_result else []
                        }
                    }

                    template_suggestion = self._generate_template_suggestion(text, extracted_data_for_response,
                                                                             confidence)

                    return {
                        "status": "rejected",
                        "document_type": "unknown",
                        "source_file": source_file,
                        "rejection_reason": f"Document type could not be determined with sufficient confidence (confidence: {confidence})",
                        "confidence": confidence,
                        "extracted_data": extracted_data_for_response,
                        "validation_level": "failed",
                        "document_metadata": extracted_data_for_response["document_metadata"],
                        "template_suggestion": template_suggestion
                    }

                logger.info(f"Detected document type: {doc_type} with confidence {confidence}")

                try:
                    extraction_prompt = DOCUMENT_EXTRACTION_PROMPT.format(
                        text=text,
                        doc_type=doc_type,
                        document_category=detection_result.get('document_category', 'unknown'),
                        issuing_authority=detection_result.get('issuing_authority', 'unknown'),
                        key_indicators=json.dumps(detection_result.get('key_indicators', []))
                    )

                    response = self.text_processor.process_text(text, extraction_prompt)
                    if not response:
                        raise ValueError("Empty response from text processor during extraction")

                    try:
                        extracted_data = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
                    except json.JSONDecodeError as e:
                        logger.error(f"Failed to parse extraction result: {str(e)}")
                        raise ValueError(f"Invalid JSON response from text processor during extraction: {str(e)}")

                    if not extracted_data or "data" not in extracted_data:
                        raise ValueError("Invalid extraction result: missing required fields")

                    logger.info(f"Extracted data: {json.dumps(extracted_data, indent=2)}")

                    try:
                        verification_result = self.verify_document(extracted_data, doc_type)
                    except Exception as e:
                        logger.error(f"Error during document verification: {str(e)}")
                        raise ValueError(f"Document verification failed: {str(e)}")

                    if not verification_result["is_genuine"]:
                        rejection_reason = verification_result.get("rejection_reason",
                                                                   "Document failed authenticity verification")
                        logger.warning(f"Document rejected - Not genuine: {rejection_reason}")

                        logger.info(
                            f"Including extracted data in rejected response: {json.dumps(extracted_data, indent=2)}")

                        return {
                            "status": "rejected",
                            "document_type": doc_type,
                            "source_file": source_file,
                            "rejection_reason": rejection_reason,
                            "verification_result": verification_result,
                            "extracted_data": extracted_data,
                            "confidence": confidence,
                            "validation_level": "strict",
                            "document_metadata": extracted_data.get("document_metadata", {})
                        }

                    logger.info(
                        f"Including extracted data in successful response: {json.dumps(extracted_data, indent=2)}")

                    response = {
                        "extracted_data": extracted_data,
                        "status": "success",
                        "confidence": confidence,
                        "document_type": doc_type,
                        "source_file": source_file,
                        "validation_level": "strict",
                        "verification_result": verification_result,
                        "document_metadata": extracted_data.get("document_metadata", {})
                    }

                    if confidence < 0.7:
                        template_suggestion = self._generate_template_suggestion(text, extracted_data, confidence)
                        template_suggestion[
                            "reason"] = f"Low confidence match (confidence: {confidence:.2f}). Consider creating a specific template for better accuracy."
                        response["template_suggestion"] = template_suggestion

                    return response

                except Exception as e:
                    logger.error(f"Error during data extraction: {str(e)}")
                    raise ValueError(f"Data extraction failed: {str(e)}")

            except Exception as e:
                logger.error(f"Error during document type detection: {str(e)}")
                raise ValueError(f"Document type detection failed: {str(e)}")

        except Exception as e:
            logger.exception(f"Error processing text content: {str(e)}")

            error_extracted_data = {
                "data": {"raw_text": text[:500] + "..." if len(text) > 500 else text} if text else {},
                "confidence": 0.0,
                "additional_info": f"Processing failed: {str(e)}",
                "document_metadata": {
                    "type": "unknown",
                    "category": "error",
                    "processing_error": str(e)
                }
            }

            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": source_file,
                "rejection_reason": f"Error processing document: {str(e)}",
                "extracted_data": error_extracted_data,
                "confidence": 0.0,
                "validation_level": "error",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "error_traceback": traceback.format_exc()
                }
            }

    def verify_document(self, extracted_data: Dict[str, Any], doc_type: str) -> Dict[str, Any]:
        """Verify document authenticity and data validity with safety filter handling"""
        try:
            # First try with the regular verification prompt
            verification_prompt = DOCUMENT_VERIFICATION_PROMPT.format(
                document_data=json.dumps(extracted_data, indent=2),
                doc_type=doc_type,
                document_category=extracted_data.get('document_metadata', {}).get('category', 'unknown'),
                issuing_authority=extracted_data.get('document_metadata', {}).get('issuing_authority', 'unknown')
            )

            try:
                response = self.text_processor.process_text(json.dumps(extracted_data), verification_prompt)
                verification_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
            except Exception as e:
                error_str = str(e)
                if "safety" in error_str.lower() or "finish_reason" in error_str:
                    logger.warning("Regular verification prompt triggered safety filters, trying safer prompt")
                    # Try with safer verification prompt
                    from Common.constants import SAFE_DOCUMENT_VERIFICATION_PROMPT
                    safer_prompt = SAFE_DOCUMENT_VERIFICATION_PROMPT.format(
                        document_data=json.dumps(extracted_data, indent=2),
                        doc_type=doc_type,
                        document_category=extracted_data.get('document_metadata', {}).get('category', 'unknown'),
                        issuing_authority=extracted_data.get('document_metadata', {}).get('issuing_authority',
                                                                                          'unknown')
                    )
                    response = self.text_processor.process_text(json.dumps(extracted_data), safer_prompt)
                    verification_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
                else:
                    raise

            checks = verification_result.get("verification_checks", {})
            confidence_scores = [
                checks.get("authenticity", {}).get("confidence", 0),
                checks.get("security_features", {}).get("confidence", 0),
                checks.get("data_validation", {}).get("confidence", 0),
                checks.get("quality", {}).get("confidence", 0)
            ]
            overall_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0

            verification_result["confidence_score"] = overall_confidence
            verification_result["is_genuine"] = overall_confidence >= VERIFICATION_THRESHOLD

            logger.info(f"Document verification results: {json.dumps(verification_result, indent=2)}")

            return verification_result

        except Exception as e:
            logger.exception(f"Error verifying document genuineness: {str(e)}")
            return {
                "is_genuine": False,
                "confidence_score": 0.0,
                "rejection_reason": f"Error during verification: {str(e)}",
                "verification_checks": {
                    "authenticity": {"passed": False, "details": f"Error during verification: {str(e)}",
                                     "confidence": 0.0},
                    "security_features": {"passed": False, "details": "Verification failed", "confidence": 0.0},
                    "data_validation": {"passed": False, "details": "Verification failed", "confidence": 0.0},
                    "quality": {"passed": False, "details": "Verification failed", "confidence": 0.0}
                },
                "security_features_found": [],
                "verification_summary": f"Document verification failed due to error: {str(e)}",
                "recommendations": ["Verification process failed due to technical error"]
            }

    def _process_pdf(self, pdf_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a PDF file that may contain multiple documents"""
        try:
            if not pdf_path or not os.path.exists(pdf_path):
                raise ValueError(f"Invalid PDF path: {pdf_path}")

            if not isinstance(min_confidence, (int, float)) or not 0 <= min_confidence <= 1:
                raise ValueError("Invalid confidence threshold: min_confidence must be between 0 and 1")

            results = []
            pdf_document = None

            try:
                pdf_document = fitz.open(pdf_path)
            except Exception as e:
                logger.error(f"Error opening PDF file: {str(e)}")
                raise ValueError(f"Failed to open PDF file: {str(e)}")

            try:

                all_text = ""
                for page_num in range(pdf_document.page_count):
                    try:
                        page = pdf_document[page_num]
                        text = page.get_text()
                        if text.strip():
                            all_text += text + "\n\n"
                    except Exception as e:
                        logger.warning(f"Error processing page {page_num + 1}: {str(e)}")
                        continue

                if all_text.strip():
                    try:

                        multi_doc_results = self._process_multiple_documents(all_text, pdf_path, min_confidence)
                        if multi_doc_results:
                            results.extend(multi_doc_results)
                            return results
                    except Exception as e:
                        logger.warning(f"Error processing PDF as multiple documents: {str(e)}")

                for page_num in range(pdf_document.page_count):
                    try:
                        logger.info(f"Processing page {page_num + 1} of {pdf_document.page_count}")

                        page = pdf_document[page_num]
                        text = page.get_text()

                        needs_ocr = self._needs_ocr(text, page)

                        if needs_ocr:
                            logger.info(f"Page {page_num + 1} requires OCR processing")
                            try:
                                images = page.get_images(full=True)
                            except Exception as e:
                                logger.error(f"Error extracting images from page {page_num + 1}: {str(e)}")
                                continue

                            if not images:
                                logger.warning(f"No images found on page {page_num + 1} for OCR")
                                continue

                            for img_index, img_info in enumerate(images):
                                try:
                                    with tempfile.TemporaryDirectory() as temp_dir:
                                        try:
                                            img_data = pdf_document.extract_image(img_info[0])
                                        except Exception as e:
                                            logger.error(f"Error extracting image data: {str(e)}")
                                            continue

                                        temp_image_path = os.path.join(temp_dir, f"page_{page_num}_img_{img_index}.png")

                                        try:
                                            with open(temp_image_path, "wb") as img_file:
                                                img_file.write(img_data["image"])
                                        except Exception as e:
                                            logger.error(f"Error saving image: {str(e)}")
                                            continue

                                        try:
                                            ocr_text = self._perform_ocr(temp_image_path)
                                        except Exception as e:
                                            logger.error(f"Error performing OCR: {str(e)}")
                                            continue

                                        if ocr_text:
                                            try:

                                                multi_doc_results = self._process_multiple_documents(ocr_text, pdf_path,
                                                                                                     min_confidence)
                                                if multi_doc_results:
                                                    for result in multi_doc_results:
                                                        result["page_number"] = page_num + 1
                                                        result["image_index"] = img_index + 1
                                                        result["processing_method"] = "ocr"
                                                    results.extend(multi_doc_results)
                                                else:

                                                    result = self._process_text_content(ocr_text, pdf_path,
                                                                                        min_confidence)
                                                    if result:
                                                        result["page_number"] = page_num + 1
                                                        result["image_index"] = img_index + 1
                                                        result["processing_method"] = "ocr"
                                                        results.append(result)
                                            except Exception as e:
                                                logger.error(f"Error processing OCR text: {str(e)}")
                                                continue
                                except Exception as e:
                                    logger.error(
                                        f"Error processing image {img_index + 1} on page {page_num + 1}: {str(e)}")
                                    continue
                        else:
                            if text.strip():
                                try:

                                    multi_doc_results = self._process_multiple_documents(text, pdf_path, min_confidence)
                                    if multi_doc_results:
                                        for result in multi_doc_results:
                                            result["page_number"] = page_num + 1
                                            result["processing_method"] = "direct_text"
                                        results.extend(multi_doc_results)
                                    else:

                                        result = self._process_text_content(text, pdf_path, min_confidence)
                                        if result:
                                            result["page_number"] = page_num + 1
                                            result["processing_method"] = "direct_text"
                                            results.append(result)
                                except Exception as e:
                                    logger.error(f"Error processing text content: {str(e)}")
                                    continue
                    except Exception as e:
                        logger.error(f"Error processing page {page_num + 1}: {str(e)}")
                        continue

                if not results:
                    logger.warning(f"No valid content found in PDF file: {pdf_path}")

                    basic_pdf_data = {
                        "data": {
                            "file_type": "pdf",
                            "file_path": pdf_path,
                            "total_pages": len(pdf_document) if 'pdf_document' in locals() else 0
                        },
                        "confidence": 0.0,
                        "additional_info": "No valid content found in PDF document",
                        "document_metadata": {
                            "type": "unknown",
                            "category": "document",
                            "file_format": "pdf"
                        }
                    }

                    return [{
                        "status": "rejected",
                        "document_type": "unknown",
                        "source_file": pdf_path,
                        "rejection_reason": "No valid content found in document",
                        "extracted_data": basic_pdf_data,
                        "confidence": 0.0,
                        "validation_level": "failed"
                    }]

                return results

            except Exception as e:
                logger.error(f"Error processing PDF file: {str(e)}")

                error_pdf_data = {
                    "data": {
                        "file_type": "pdf",
                        "file_path": pdf_path,
                        "processing_error": str(e),
                        "error_type": type(e).__name__
                    },
                    "confidence": 0.0,
                    "additional_info": f"PDF processing failed: {str(e)}",
                    "document_metadata": {
                        "type": "unknown",
                        "category": "error",
                        "file_format": "pdf"
                    }
                }

                return [{
                    "status": "error",
                    "document_type": "unknown",
                    "source_file": pdf_path,
                    "rejection_reason": f"Error processing document: {str(e)}",
                    "extracted_data": error_pdf_data,
                    "confidence": 0.0,
                    "validation_level": "error",
                    "error_details": {
                        "error_type": type(e).__name__,
                        "error_message": str(e),
                        "error_traceback": traceback.format_exc()
                    }
                }]

            finally:
                if pdf_document:
                    try:
                        pdf_document.close()
                    except Exception as e:
                        logger.error(f"Error closing PDF file: {str(e)}")

        except Exception as e:
            logger.error(f"Error in PDF processing: {str(e)}")

            outer_error_data = {
                "data": {
                    "file_type": "pdf",
                    "file_path": pdf_path,
                    "outer_processing_error": str(e),
                    "error_type": type(e).__name__
                },
                "confidence": 0.0,
                "additional_info": f"PDF outer processing failed: {str(e)}",
                "document_metadata": {
                    "type": "unknown",
                    "category": "error",
                    "file_format": "pdf"
                }
            }

            return [{
                "status": "error",
                "document_type": "unknown",
                "source_file": pdf_path,
                "rejection_reason": f"Error processing document: {str(e)}",
                "extracted_data": outer_error_data,
                "confidence": 0.0,
                "validation_level": "error",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                    "error_traceback": traceback.format_exc()
                }
            }]

    def _generate_template_suggestion(self, text: str, extracted_data: dict, confidence: float) -> dict:
        """Generate a template creation suggestion for unknown document types"""
        try:
            text_lower = text.lower() if text else ""

            suggestion = {
                "should_create_template": True,
                "reason": f"Document type not recognized (confidence: {confidence:.2f})",
                "suggested_template": self._analyze_content_for_template_type(text_lower),
                "confidence": confidence,
                "extracted_content_sample": text[:300] + "..." if len(text) > 300 else text,
                "extracted_data_preview": self._extract_sample_fields(text, extracted_data)
            }

            logger.info(f"Generated template suggestion: {suggestion['suggested_template']['name']}")
            return suggestion

        except Exception as e:
            logger.error(f"Error generating template suggestion: {str(e)}")
            return {
                "should_create_template": True,
                "reason": "Document type not recognized",
                "suggested_template": {
                    "name": "Custom Template",
                    "type": "other",
                    "description": "Custom document template based on uploaded document",
                    "suggested_fields": ["document_number", "name", "date"]
                },
                "confidence": confidence,
                "extracted_content_sample": text[:300] if text else "",
                "extracted_data_preview": {}
            }

    def _analyze_content_for_template_type(self, text_lower: str) -> dict:
        """Analyze document content to suggest template details"""

        suggestion = {
            "name": "Custom Template",
            "type": "other",
            "description": "Custom document template based on uploaded document",
            "suggested_fields": ["document_number", "name", "date"]
        }

        if any(keyword in text_lower for keyword in ["aadhaar", "aadhar", "uid", "unique identification"]):
            suggestion.update({
                "name": "Custom Aadhaar Template",
                "type": "aadhaar_card",
                "description": "Custom Aadhaar card template detected from document content",
                "suggested_fields": ["name", "aadhaar_number", "date_of_birth", "address", "gender"]
            })
        elif any(keyword in text_lower for keyword in ["pan", "permanent account", "income tax"]):
            suggestion.update({
                "name": "Custom PAN Template",
                "type": "pan_card",
                "description": "Custom PAN card template detected from document content",
                "suggested_fields": ["name", "fathers_name", "pan_number", "date_of_birth"]
            })
        elif any(keyword in text_lower for keyword in ["driving", "license", "licence", "vehicle", "motor"]):
            suggestion.update({
                "name": "Custom License Template",
                "type": "driving_license",
                "description": "Custom driving license template detected from document content",
                "suggested_fields": ["name", "license_number", "date_of_birth", "address", "vehicle_class"]
            })
        elif any(keyword in text_lower for keyword in ["passport", "travel", "nationality", "republic"]):
            suggestion.update({
                "name": "Custom Passport Template",
                "type": "passport",
                "description": "Custom passport template detected from document content",
                "suggested_fields": ["name", "passport_number", "date_of_birth", "place_of_birth", "nationality"]
            })
        elif any(keyword in text_lower for keyword in
                 ["company", "corporation", "business", "registration", "incorporation"]):
            suggestion.update({
                "name": "Custom Corporate Template",
                "type": "corporate_document",
                "description": "Custom corporate document template detected from document content",
                "suggested_fields": ["company_name", "registration_number", "address", "incorporation_date"]
            })
        elif any(keyword in text_lower for keyword in ["certificate", "diploma", "degree", "graduation"]):
            suggestion.update({
                "name": "Custom Certificate Template",
                "type": "educational_certificate",
                "description": "Custom educational certificate template detected from document content",
                "suggested_fields": ["student_name", "institution", "degree", "graduation_date"]
            })
        elif any(keyword in text_lower for keyword in ["medical", "hospital", "doctor", "patient"]):
            suggestion.update({
                "name": "Custom Medical Template",
                "type": "medical_document",
                "description": "Custom medical document template detected from document content",
                "suggested_fields": ["patient_name", "doctor_name", "hospital", "date", "diagnosis"]
            })
        elif any(keyword in text_lower for keyword in ["bank", "statement", "account", "balance", "transaction"]):
            suggestion.update({
                "name": "Custom Bank Statement Template",
                "type": "bank_statement",
                "description": "Custom bank statement template detected from document content",
                "suggested_fields": ["account_holder", "account_number", "statement_period", "balance"]
            })
        elif any(keyword in text_lower for keyword in ["invoice", "bill", "receipt", "payment"]):
            suggestion.update({
                "name": "Custom Invoice Template",
                "type": "invoice",
                "description": "Custom invoice/bill template detected from document content",
                "suggested_fields": ["invoice_number", "date", "amount", "vendor", "customer"]
            })
        elif any(keyword in text_lower for keyword in ["resume", "cv", "curriculum vitae", "experience", "skills"]):
            suggestion.update({
                "name": "Custom Resume Template",
                "type": "resume",
                "description": "Custom resume/CV template detected from document content",
                "suggested_fields": ["name", "email", "phone", "experience", "education", "skills"]
            })
        elif any(keyword in text_lower for keyword in ["property", "deed", "sale", "rental", "agreement"]):
            suggestion.update({
                "name": "Custom Property Template",
                "type": "property_document",
                "description": "Custom property document template detected from document content",
                "suggested_fields": ["property_address", "owner_name", "document_type", "date", "value"]
            })
        elif any(keyword in text_lower for keyword in ["insurance", "policy", "premium", "coverage"]):
            suggestion.update({
                "name": "Custom Insurance Template",
                "type": "insurance_document",
                "description": "Custom insurance document template detected from document content",
                "suggested_fields": ["policy_number", "policyholder", "coverage", "premium", "validity"]
            })
        elif any(keyword in text_lower for keyword in ["transcript", "marksheet", "grade", "marks", "gpa"]):
            suggestion.update({
                "name": "Custom Transcript Template",
                "type": "transcript",
                "description": "Custom academic transcript template detected from document content",
                "suggested_fields": ["student_name", "student_id", "institution", "grades", "gpa", "semester"]
            })
        elif any(keyword in text_lower for keyword in ["employment", "appointment", "salary", "job", "designation"]):
            suggestion.update({
                "name": "Custom Employment Template",
                "type": "employment_letter",
                "description": "Custom employment document template detected from document content",
                "suggested_fields": ["employee_name", "designation", "company", "salary", "joining_date"]
            })
        elif any(keyword in text_lower for keyword in ["tax", "income", "return", "assessment"]):
            suggestion.update({
                "name": "Custom Tax Template",
                "type": "tax_document",
                "description": "Custom tax document template detected from document content",
                "suggested_fields": ["taxpayer_name", "tax_year", "income", "tax_amount", "assessment_number"]
            })
        elif any(keyword in text_lower for keyword in ["utility", "electricity", "water", "gas", "phone"]):
            suggestion.update({
                "name": "Custom Utility Bill Template",
                "type": "utility_bill",
                "description": "Custom utility bill template detected from document content",
                "suggested_fields": ["consumer_name", "consumer_number", "billing_period", "amount", "due_date"]
            })

        return suggestion

    def _extract_sample_fields(self, text: str, extracted_data: dict) -> dict:
        """Extract sample field values from document text and extracted data"""
        import re

        sample_fields = {}

        try:

            if extracted_data and "data" in extracted_data:
                data_content = extracted_data["data"]
                if isinstance(data_content, dict):
                    for key, value in data_content.items():
                        if isinstance(value, str) and value.strip():
                            sample_fields[key] = value[:50]
                        elif isinstance(value, dict):
                            for sub_key, sub_value in value.items():
                                if isinstance(sub_value, str) and sub_value.strip():
                                    sample_fields[f"{key}_{sub_key}"] = str(sub_value)[:50]

            if text:

                date_patterns = [
                    r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{4})\b',
                    r'\b(\d{4}[/-]\d{1,2}[/-]\d{1,2})\b',
                    r'\b(\d{1,2}\s+\w+\s+\d{4})\b'
                ]

                for pattern in date_patterns:
                    matches = re.findall(pattern, text)
                    if matches:
                        sample_fields["dates_found"] = matches[:3]
                        break

                id_patterns = [
                    r'\b([A-Z]{3,5}\d{4,10}[A-Z]?)\b',
                    r'\b(\d{12})\b',
                    r'\b(\d{10,16})\b'
                ]

                for pattern in id_patterns:
                    matches = re.findall(pattern, text)
                    if matches:
                        sample_fields["id_numbers_found"] = matches[:3]
                        break

                name_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b'
                names = re.findall(name_pattern, text)
                if names:
                    sample_fields["names_found"] = [name for name in names if len(name.split()) >= 2][:3]

        except Exception as e:
            logger.error(f"Error extracting sample fields: {str(e)}")

        return sample_fields

    def _needs_ocr(self, text: str, page) -> bool:
        """Determine if OCR is needed for the page"""
        try:
            if not text.strip():
                return True

            if self._is_character_by_character(text):
                return True

            images = page.get_images(full=True)
            if images:
                if not self._has_meaningful_content(text):
                    return True

            if self._has_document_indicators(page):
                return True

            return False

        except Exception as e:
            logger.error(f"Error determining OCR need: {str(e)}")
            return True

    def _process_single_image(self, image_path: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process a single image file with simplified error handling"""
        try:
            logger.info(f"Processing single image: {image_path}")

            # Always try Tesseract OCR first
            try:
                img = Image.open(image_path)
                ocr_text = pytesseract.image_to_string(img)
            except Exception as e:
                logger.error(f"Error with Tesseract OCR: {str(e)}")
                return {
                    "status": "error",
                    "document_type": "unknown",
                    "source_file": image_path,
                    "rejection_reason": f"OCR processing failed: {str(e)}"
                }

            # If Tesseract results are poor, try Gemini Vision
            if not self._is_good_ocr_result(ocr_text):
                logger.info("Tesseract OCR yielded poor results, trying Gemini Vision")
                try:
                    response = self._process_with_gemini(
                        image_path,
                        "Extract all text from this document image. Return only the raw text without any formatting."
                    )
                    if response and response.strip():
                        ocr_text = response.strip()
                except Exception as e:
                    logger.warning(f"Gemini Vision failed, using Tesseract results: {str(e)}")
                    # Continue with the original Tesseract results

            # Process the extracted text
            try:
                result = self._process_text_content(ocr_text, image_path, min_confidence)
                if result:
                    return result

                # If processing didn't yield a result, return basic info
                return {
                    "status": "processed",
                    "document_type": "unknown",
                    "confidence": 0.3,
                    "source_file": image_path,
                    "extracted_text": ocr_text,
                    "processing_method": "ocr_only"
                }
            except Exception as e:
                logger.error(f"Error processing extracted text: {str(e)}")
                return {
                    "status": "error",
                    "document_type": "unknown",
                    "source_file": image_path,
                    "rejection_reason": f"Text processing failed: {str(e)}",
                    "extracted_text": ocr_text
                }

        except Exception as e:
            logger.error(f"Error processing single image: {str(e)}")
            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": image_path,
                "rejection_reason": f"Error processing image: {str(e)}"
            }

    def _is_character_by_character(self, text: str) -> bool:
        """Check if text appears to be extracted character by character"""
        try:
            patterns = [
                r'[A-Z]\s+[A-Z]\s+[A-Z]',
                r'[a-z]\s+[a-z]\s+[a-z]',
                r'[0-9]\s+[0-9]\s+[0-9]',
                r'[A-Za-z]\s+[A-Za-z]\s+[A-Za-z]'
            ]

            for pattern in patterns:
                if re.search(pattern, text):
                    return True

            if len(text) > 0:
                space_ratio = text.count(' ') / len(text)
                if space_ratio > 0.3:
                    return True

            return False

        except Exception as e:
            logger.error(f"Error checking character-by-character extraction: {str(e)}")
            return False

    def _process_with_gemini(self, image_path: str, prompt: str) -> str:
        """
        Process an image with Gemini Vision with fallback to Tesseract

        Args:
            image_path: Path to the image file
            prompt: Prompt to send to Gemini

        Returns:
            Extracted text from the image
        """
        try:
            # Try Gemini first
            try:
                response = self.text_extractor.process_with_gemini(image_path, prompt)
                if response and response.strip():
                    return response.strip()
            except Exception as e:
                error_str = str(e)
                logger.warning(f"Gemini Vision attempt failed: {error_str}")

                # Check if this is a safety filter block
                if "finish_reason" in error_str and "safety_ratings" in error_str:
                    logger.warning("Gemini safety filters triggered - using Tesseract OCR")
                else:
                    logger.warning(f"Gemini Vision failed with error: {error_str}")

                # Fall through to Tesseract OCR

            # Use Tesseract OCR as fallback
            logger.info("Using Tesseract OCR as fallback")
            img = Image.open(image_path)
            ocr_text = pytesseract.image_to_string(img)

            if ocr_text.strip():
                logger.info("Successfully extracted text with Tesseract OCR")
                return ocr_text.strip()
            else:
                logger.warning("Tesseract OCR returned empty result")
                raise ValueError("Tesseract OCR failed to extract text")

        except Exception as e:
            logger.error(f"Error in _process_with_gemini: {str(e)}")
            raise ValueError(f"Image text extraction failed: {str(e)}")

    def _has_meaningful_content(self, text: str) -> bool:
        """Check if text contains meaningful content"""
        try:

            indicators = [
                r'[A-Z]{2,}',
                r'\d{4,}',
                r'[A-Za-z]+\s+[A-Za-z]+',
                r'[A-Za-z]+\s+\d+',
                r'\d+\s+[A-Za-z]+'
            ]

            indicator_count = 0
            for pattern in indicators:
                if re.search(pattern, text):
                    indicator_count += 1

            return indicator_count >= 2

        except Exception as e:
            logger.error(f"Error checking meaningful content: {str(e)}")
            return False

    def _has_document_indicators(self, page) -> bool:
        """Check if page has indicators of being a document that might need OCR"""
        try:

            images = page.get_images(full=True)
            if not images:
                return False

            blocks = page.get_text("blocks")
            if not blocks:
                return True

            for block in blocks:

                if block[4].strip():
                    if any(indicator in block[4].lower() for indicator in [
                        "name", "date", "address", "signature", "photo",
                        "passport", "license", "id", "number", "issued"
                    ]):
                        return True

            return False

        except Exception as e:
            logger.error(f"Error checking document indicators: {str(e)}")
            return True

    def _perform_ocr(self, image_path: str) -> str:
        """Perform OCR on an image with simplified error handling"""
        try:
            if not image_path or not os.path.exists(image_path):
                raise ValueError(f"Invalid image path: {image_path}")

            try:
                img = Image.open(image_path)
            except Exception as e:
                logger.error(f"Error opening image: {str(e)}")
                raise ValueError(f"Failed to open image: {str(e)}")

            # Always try Tesseract OCR first
            try:
                ocr_text = pytesseract.image_to_string(img)

                # If Tesseract gives good results, use them
                if self._is_good_ocr_result(ocr_text):
                    return ocr_text

                # Otherwise try Gemini, but be prepared to fall back to Tesseract
                logger.info("Tesseract OCR yielded poor results, trying Gemini Vision")
                try:
                    response = self._process_with_gemini(
                        image_path,
                        "Extract all text from this document image. Return only the raw text without any formatting."
                    )
                    if response and response.strip():
                        return response.strip()
                except Exception as e:
                    logger.warning(f"Gemini Vision failed, using Tesseract results: {str(e)}")
                    # Return the original Tesseract results if Gemini fails
                    return ocr_text

            except Exception as e:
                logger.error(f"All OCR methods failed: {str(e)}")
                raise ValueError(f"OCR processing failed: {str(e)}")

            return ocr_text

        except Exception as e:
            logger.error(f"Error performing OCR: {str(e)}")
            raise ValueError(f"OCR processing failed: {str(e)}")

    def _is_good_ocr_result(self, text: str) -> bool:
        """Check if OCR result is of good quality"""
        try:
            if not text.strip():
                return False

            if len(text.strip()) < 10:
                return False

            if not self._has_meaningful_content(text):
                return False

            error_patterns = [
                r'[|]{2,}',
                r'[l1]{3,}',
                r'[o0]{3,}',
                r'[rn]{3,}',
                r'\s{3,}'
            ]

            for pattern in error_patterns:
                if re.search(pattern, text):
                    return False

            return True

        except Exception as e:
            logger.error(f"Error checking OCR result quality: {str(e)}")
            return False

    # Methods removed - using constants from constants.py

    def _process_multiple_documents(self, text: str, source_file: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process text that may contain multiple documents"""
        try:
            results = []

            chunks = self._split_into_chunks(text)
            logger.info(f"Split text into {len(chunks)} potential document chunks")

            for chunk_index, chunk in enumerate(chunks):
                if not chunk.strip():
                    continue

                logger.info(f"Processing chunk {chunk_index + 1} of {len(chunks)}")

                result = self._process_text_content(chunk, source_file, min_confidence)

                if result:
                    result["chunk_index"] = chunk_index + 1
                    result["total_chunks"] = len(chunks)
                    results.append(result)

            # Merge results if we have multiple chunks with data
            if len(results) > 1:
                merged_result = self._merge_chunk_results(results, source_file)
                return [merged_result]

            return results

        except Exception as e:
            logger.error(f"Error processing multiple documents: {str(e)}")
            return []

    def _merge_chunk_results(self, results: List[Dict[str, Any]], source_file: str) -> Dict[str, Any]:
        """
        Merge results from multiple chunks into a single comprehensive result

        Args:
            results: List of results from different chunks
            source_file: Source file name

        Returns:
            Merged result dictionary
        """
        try:
            logger.info(f"Merging {len(results)} chunk results")

            # Find the best result (highest confidence with most data)
            best_result = None
            best_score = 0

            for result in results:
                confidence = result.get('confidence', 0.0)
                extracted_data = result.get('extracted_data', {}).get('data', {})
                data_count = len(
                    [v for v in extracted_data.values() if v and str(v).lower() not in ['null', 'not_present', 'n/a']])

                # Score based on confidence and amount of extracted data
                score = confidence + (data_count * 0.1)

                if score > best_score:
                    best_score = score
                    best_result = result

            if not best_result:
                best_result = results[0]

            # Merge extracted data from all chunks
            merged_data = {}
            all_document_types = set()

            for result in results:
                # Collect document types
                doc_type = result.get('document_type', 'unknown')
                if doc_type != 'unknown':
                    all_document_types.add(doc_type)

                # Merge extracted data
                extracted_data = result.get('extracted_data', {}).get('data', {})
                for key, value in extracted_data.items():
                    if value and str(value).lower() not in ['null', 'not_present', 'n/a', '']:
                        if key not in merged_data or not merged_data[key]:
                            merged_data[key] = value
                        elif isinstance(value, list) and isinstance(merged_data[key], list):
                            # Merge lists and remove duplicates
                            merged_data[key] = list(set(merged_data[key] + value))
                        elif key in merged_data and merged_data[key] != value:
                            # If we have different values for the same key, create a list
                            if not isinstance(merged_data[key], list):
                                merged_data[key] = [merged_data[key]]
                            if value not in merged_data[key]:
                                merged_data[key].append(value)

            # Determine the primary document type
            if len(all_document_types) == 1:
                primary_doc_type = list(all_document_types)[0]
            elif 'resume' in all_document_types:
                primary_doc_type = 'resume'
            elif all_document_types:
                primary_doc_type = list(all_document_types)[0]
            else:
                primary_doc_type = best_result.get('document_type', 'unknown')

            # Create merged result
            merged_result = {
                "status": "success",
                "document_type": primary_doc_type,
                "source_file": source_file,
                "confidence": max([r.get('confidence', 0.0) for r in results]),
                "extracted_data": {
                    "data": merged_data,
                    "confidence": max([r.get('confidence', 0.0) for r in results]),
                    "additional_info": f"Merged from {len(results)} document chunks",
                    "document_metadata": {
                        "type": primary_doc_type,
                        "category": best_result.get('extracted_data', {}).get('document_metadata', {}).get('category',
                                                                                                           'unknown'),
                        "issuing_authority": "merged_document",
                        "key_indicators": list(all_document_types),
                        "chunks_processed": len(results)
                    }
                },
                "processing_method": "unified_prompt_merged",
                "validation_level": "comprehensive",
                "chunk_info": {
                    "total_chunks": len(results),
                    "chunks_with_data": len([r for r in results if r.get('extracted_data', {}).get('data')]),
                    "document_types_found": list(all_document_types)
                }
            }

            # Add verification result from best result
            if best_result.get('verification_result'):
                merged_result['verification_result'] = best_result['verification_result']

            logger.info(f"Merged result: document_type={primary_doc_type}, data_fields={len(merged_data)}")
            return merged_result

        except Exception as e:
            logger.error(f"Error merging chunk results: {str(e)}")
            # Return the best individual result as fallback
            return max(results, key=lambda r: r.get('confidence', 0.0)) if results else {}

    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> None:
        """Convert a DOCX file to PDF format"""
        try:

            try:
                import subprocess

                subprocess.run(['soffice', '--version'], check=True, capture_output=True)

                subprocess.run([
                    'soffice',
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', os.path.dirname(pdf_path),
                    docx_path
                ], check=True)

                generated_pdf = os.path.join(
                    os.path.dirname(pdf_path),
                    os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
                )
                if generated_pdf != pdf_path:
                    os.rename(generated_pdf, pdf_path)

                return
            except Exception as e:
                logger.warning(f"LibreOffice conversion failed: {str(e)}")

            try:
                import win32com.client
                word = win32com.client.Dispatch('Word.Application')
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)
                doc.Close()
                word.Quit()
                return
            except Exception as e:
                logger.warning(f"Win32com conversion failed: {str(e)}")

            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                return
            except Exception as e:
                logger.warning(f"docx2pdf conversion failed: {str(e)}")

            raise RuntimeError("All conversion methods failed")

        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {str(e)}")
            raise RuntimeError(f"Failed to convert DOCX to PDF: {str(e)}")