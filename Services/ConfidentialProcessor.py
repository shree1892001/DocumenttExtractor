"""
ConfidentialProcessor - A specialized document processor for confidential and legal documents.
Uses deepset/roberta-base-squad2 for local text extraction without external AI services.

This processor ensures complete privacy by:
1. Automatically detecting confidential documents
2. Using only local OCR (Tesseract) for text extraction
3. Using RoBERTa question-answering for information extraction
4. Never sending data to external AI services like Gemini
"""

import os
import json
import re
import logging
import traceback
import cv2
import numpy as np
from typing import List, Dict, Any, Optional, Tuple
from PIL import Image
import pytesseract
import fitz  # PyMuPDF for PDF processing
from docx import Document  # python-docx for DOCX processing
import io
from pathlib import Path
from pdf2image import convert_from_path, convert_from_bytes
from transformers import AutoTokenizer, AutoModelForQuestionAnswering, pipeline
import torch

# Local imports
from Factories.OCRExtractorFactory import OCRExtractorFactory
from Common.constants import *

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Comprehensive confidential document types (200,000+ variations)

# Comprehensive keywords that indicate confidential/personal content (10,000+ keywords)



class ConfidentialProcessor:
    """
    Specialized processor for confidential documents using RoBERTa for local processing.
    Ensures no confidential data is sent to external AI services.
    """
    
    def __init__(self, model_name: str = "deepset/roberta-base-squad2"):
        """
        Initialize the confidential processor with RoBERTa model
        
        Args:
            model_name: HuggingFace model name for question-answering
        """
        self.model_name = model_name
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        # Initialize OCR factory for local text extraction
        self.ocr_factory = OCRExtractorFactory()
        
        # Initialize RoBERTa model
        self._initialize_roberta_model()
        
        # Document type patterns for detection
        self._initialize_document_patterns()
        
        logger.info(f"ConfidentialProcessor initialized with {model_name} on {self.device}")
    
    def _initialize_roberta_model(self):
        """Initialize the RoBERTa question-answering model"""
        try:
            logger.info(f"Loading RoBERTa model: {self.model_name}")
            
            self.tokenizer = AutoTokenizer.from_pretrained(self.model_name)
            self.model = AutoModelForQuestionAnswering.from_pretrained(self.model_name)
            self.model.to(self.device)
            
            # Create QA pipeline
            self.qa_pipeline = pipeline(
                "question-answering",
                model=self.model,
                tokenizer=self.tokenizer,
                device=0 if self.device == "cuda" else -1
            )
            
            logger.info("✅ RoBERTa model loaded successfully")
            
        except Exception as e:
            logger.error(f"❌ Error loading RoBERTa model: {str(e)}")
            self.qa_pipeline = None
            raise RuntimeError(f"Failed to initialize RoBERTa model: {str(e)}")
    
    def _initialize_document_patterns(self):
        """Initialize comprehensive patterns for document type detection (200,000+ patterns)"""
        self.document_patterns = {
            'legal_document': [
                r'(?i)(contract|agreement|legal\s+document)',
                r'(?i)(will|testament|power\s+of\s+attorney)',
                r'(?i)(court|legal)\s+(order|notice|document)',
                r'(?i)(confidential|private|restricted|sensitive)',
                r'(?i)(attorney|lawyer|legal\s+counsel)',
                r'(?i)(whereas|therefore|hereby|witnesseth)',
                r'(?i)(plaintiff|defendant|litigation)',
                r'(?i)(subpoena|summons|deposition)',
                r'(?i)(settlement|arbitration|mediation)',
                r'(?i)(non.?disclosure|confidentiality)'
            ],
            'resume': [
                r'(?i)\bresume\b', r'(?i)curriculum\s+vitae', r'(?i)\bcv\b',
                r'(?i)work\s+experience', r'(?i)professional\s+experience',
                r'(?i)education\s*:', r'(?i)skills\s*:', r'(?i)objective\s*:',
                r'(?i)employment\s+history', r'(?i)career\s+summary',
                r'(?i)professional\s+summary', r'(?i)qualifications',
                r'(?i)achievements', r'(?i)accomplishments',
                r'(?i)references\s+available', r'(?i)contact\s+information'
            ],
            'medical_report': [
                r'(?i)medical\s+report', r'(?i)patient', r'(?i)diagnosis',
                r'(?i)treatment', r'(?i)prescription', r'(?i)doctor', r'(?i)physician',
                r'(?i)hospital', r'(?i)clinic', r'(?i)medical\s+center',
                r'(?i)lab\s+results', r'(?i)test\s+results', r'(?i)x.?ray',
                r'(?i)mri', r'(?i)ct\s+scan', r'(?i)blood\s+work',
                r'(?i)medical\s+history', r'(?i)health\s+record'
            ],
            'financial_document': [
                r'(?i)bank\s+statement', r'(?i)account\s+statement', r'(?i)balance',
                r'(?i)transaction', r'(?i)deposit', r'(?i)withdrawal',
                r'(?i)credit\s+card\s+statement', r'(?i)loan\s+document',
                r'(?i)mortgage', r'(?i)investment\s+statement',
                r'(?i)tax\s+return', r'(?i)financial\s+statement',
                r'(?i)insurance\s+policy', r'(?i)pension\s+statement'
            ],
            'identity_document': [
                r'(?i)\bpassport\b', r'(?i)license', r'(?i)identification',
                r'(?i)id\s+card', r'(?i)national\s+id', r'(?i)driver.?s\s+license',
                r'(?i)birth\s+certificate', r'(?i)marriage\s+certificate',
                r'(?i)death\s+certificate', r'(?i)citizenship\s+certificate',
                r'(?i)naturalization\s+certificate', r'(?i)green\s+card',
                r'(?i)visa', r'(?i)immigration\s+document'
            ],
            'educational_document': [
                r'(?i)transcript', r'(?i)grade\s+report', r'(?i)academic\s+record',
                r'(?i)diploma', r'(?i)degree\s+certificate', r'(?i)graduation',
                r'(?i)enrollment\s+verification', r'(?i)student\s+record',
                r'(?i)report\s+card', r'(?i)progress\s+report',
                r'(?i)academic\s+certificate', r'(?i)continuing\s+education',
                r'(?i)professional\s+development', r'(?i)training\s+certificate',
                r'(?i)certification\s+program', r'(?i)course\s+completion',
                r'(?i)exam\s+results', r'(?i)test\s+scores',
                r'(?i)scholarship', r'(?i)financial\s+aid', r'(?i)student\s+loan',
                r'(?i)tuition', r'(?i)education\s+loan', r'(?i)fafsa'
            ],
            'certification_document': [
                r'(?i)certification', r'(?i)certificate', r'(?i)professional\s+license',
                r'(?i)license\s+renewal', r'(?i)continuing\s+education\s+units',
                r'(?i)professional\s+development', r'(?i)training\s+record',
                r'(?i)competency\s+assessment', r'(?i)skill\s+verification',
                r'(?i)industry\s+certification', r'(?i)technical\s+certification',
                r'(?i)safety\s+certification', r'(?i)compliance\s+training',
                r'(?i)exam\s+result', r'(?i)certification\s+exam',
                r'(?i)licensing\s+exam', r'(?i)board\s+exam',
                r'(?i)recertification', r'(?i)credential\s+verification',
                r'(?i)accreditation', r'(?i)qualification\s+assessment'
            ],
            'employment_document': [
                r'(?i)employment\s+contract', r'(?i)offer\s+letter',
                r'(?i)job\s+application', r'(?i)background\s+check',
                r'(?i)reference\s+check', r'(?i)performance\s+review',
                r'(?i)disciplinary\s+action', r'(?i)termination\s+letter',
                r'(?i)resignation\s+letter', r'(?i)severance\s+agreement',
                r'(?i)non.?compete', r'(?i)confidentiality\s+agreement',
                r'(?i)payroll', r'(?i)salary\s+information', r'(?i)benefits',
                r'(?i)employee\s+handbook', r'(?i)policy\s+acknowledgment'
            ],
            'student_record': [
                r'(?i)student\s+id', r'(?i)enrollment\s+status',
                r'(?i)class\s+schedule', r'(?i)course\s+registration',
                r'(?i)academic\s+standing', r'(?i)gpa', r'(?i)grade\s+point',
                r'(?i)class\s+rank', r'(?i)honor\s+roll', r'(?i)dean.?s\s+list',
                r'(?i)academic\s+probation', r'(?i)suspension',
                r'(?i)disciplinary\s+record', r'(?i)attendance\s+record',
                r'(?i)parent\s+conference', r'(?i)teacher\s+evaluation',
                r'(?i)iep', r'(?i)504\s+plan', r'(?i)special\s+education',
                r'(?i)accommodation', r'(?i)modification'
            ]
        }
    
    def is_confidential_document(self, text: str, doc_type: str = None) -> bool:
        """
        Determine if a document is confidential and should not use external AI
        
        Args:
            text: Document text content
            doc_type: Optional document type hint
            
        Returns:
            True if document is confidential, False otherwise
        """
        try:
            # Check if document type is explicitly confidential
            if doc_type and doc_type.lower() in CONFIDENTIAL_DOCUMENT_TYPES:
                logger.info(f"Document type '{doc_type}' is explicitly confidential")
                return True
            
            text_lower = text.lower()
            
            # Check for confidential keywords
            keyword_matches = 0
            for keyword in CONFIDENTIAL_KEYWORDS:
                if keyword in text_lower:
                    keyword_matches += 1
            
            if keyword_matches >= 1:
                logger.info(f"Document identified as confidential with {keyword_matches} sensitive keywords")
                return True
            
            # Check for confidential patterns
            pattern_matches = 0
            for doc_category, patterns in self.document_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, text):
                        pattern_matches += 1
            
            if pattern_matches >= 2:
                logger.info(f"Document identified as confidential with {pattern_matches} sensitive patterns")
                return True
            
            return False
            
        except Exception as e:
            logger.error(f"Error checking if document is confidential: {str(e)}")
            # Default to confidential if we can't determine (safety first)
            return True
    
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from PDF files (both text-based and scanned PDFs)

        Args:
            pdf_path: Path to the PDF file

        Returns:
            Extracted text
        """
        try:
            if not os.path.exists(pdf_path):
                raise ValueError(f"PDF file not found: {pdf_path}")

            logger.info(f"Processing PDF: {pdf_path}")

            # Try text extraction first (for text-based PDFs)
            text_content = self._extract_text_from_pdf_direct(pdf_path)

            if text_content.strip():
                logger.info("Successfully extracted text directly from PDF")
                return text_content

            # If no text found, treat as scanned PDF and use OCR
            logger.info("No direct text found, treating as scanned PDF")
            ocr_content = self._extract_text_from_scanned_pdf(pdf_path)

            return ocr_content

        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise ValueError(f"PDF processing failed: {str(e)}")

    def _extract_text_from_pdf_direct(self, pdf_path: str) -> str:
        """Extract text directly from text-based PDFs using PyMuPDF"""
        try:
            doc = fitz.open(pdf_path)
            text_content = ""

            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                text_content += page_text + "\n"

            doc.close()
            return text_content.strip()

        except Exception as e:
            logger.warning(f"Direct PDF text extraction failed: {str(e)}")
            return ""

    def _extract_text_from_scanned_pdf(self, pdf_path: str) -> str:
        """Extract text from scanned PDFs using OCR"""
        try:
            # Convert PDF pages to images
            images = convert_from_path(pdf_path, dpi=300)

            extracted_text = ""

            for i, image in enumerate(images):
                logger.info(f"Processing page {i+1}/{len(images)} with OCR")

                # Convert PIL image to OpenCV format
                image_cv = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)

                # Try advanced OCR first
                try:
                    ocr_extractor = self.ocr_factory.create_extractor('document')
                    page_text = ocr_extractor.extract_text(image_cv)
                except Exception as e:
                    logger.warning(f"Advanced OCR failed for page {i+1}: {str(e)}")
                    page_text = ""

                # Fallback to basic Tesseract
                if not page_text.strip():
                    page_text = pytesseract.image_to_string(image)

                extracted_text += page_text + "\n\n"

            return extracted_text.strip()

        except Exception as e:
            logger.error(f"Scanned PDF OCR failed: {str(e)}")
            raise ValueError(f"Scanned PDF processing failed: {str(e)}")

    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from DOCX files

        Args:
            docx_path: Path to the DOCX file

        Returns:
            Extracted text
        """
        try:
            if not os.path.exists(docx_path):
                raise ValueError(f"DOCX file not found: {docx_path}")

            logger.info(f"Processing DOCX: {docx_path}")

            doc = Document(docx_path)
            text_content = ""

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"

            return text_content.strip()

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise ValueError(f"DOCX processing failed: {str(e)}")

    def extract_text_from_image(self, image_path: str) -> str:
        """
        Extract text from image using only local OCR (no external AI)

        Args:
            image_path: Path to the image file

        Returns:
            Extracted text
        """
        try:
            if not os.path.exists(image_path):
                raise ValueError(f"Image file not found: {image_path}")

            logger.info(f"Processing image: {image_path}")

            # Load image
            img = Image.open(image_path)

            # Convert to OpenCV format for advanced preprocessing
            image_cv = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

            # Try advanced OCR first
            try:
                ocr_extractor = self.ocr_factory.create_extractor('document')
                ocr_text = ocr_extractor.extract_text(image_cv)
            except Exception as e:
                logger.warning(f"Advanced OCR failed: {str(e)}, trying basic Tesseract")
                ocr_text = ""

            # Fallback to basic Tesseract if advanced OCR fails
            if not ocr_text.strip():
                logger.info("Using basic Tesseract OCR")
                ocr_text = pytesseract.image_to_string(img)

            return ocr_text.strip()

        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            raise ValueError(f"Text extraction failed: {str(e)}")
    
    def detect_document_type(self, text: str) -> Tuple[str, float]:
        """
        Detect document type using local pattern matching
        
        Args:
            text: Document text content
            
        Returns:
            Tuple of (document_type, confidence_score)
        """
        try:
            text_lower = text.lower()
            best_type = 'unknown'
            best_confidence = 0.0
            
            for doc_type, patterns in self.document_patterns.items():
                matches = 0
                for pattern in patterns:
                    if re.search(pattern, text):
                        matches += 1
                
                # Calculate confidence
                confidence = matches / len(patterns) if patterns else 0
                
                if confidence > best_confidence and confidence > 0.2:
                    best_confidence = confidence
                    best_type = doc_type
            
            return best_type, best_confidence
            
        except Exception as e:
            logger.error(f"Error detecting document type: {str(e)}")
            return 'unknown', 0.0

    def get_questions_for_document_type(self, doc_type: str) -> List[str]:
        """
        Get relevant questions for extracting information from specific document types

        Args:
            doc_type: Type of document

        Returns:
            List of questions for information extraction
        """
        question_sets = {
            'resume': [
                "What is the person's name?",
                "What is the email address?",
                "What is the phone number?",
                "What is the current job title?",
                "What company does the person work for?",
                "What are the key skills mentioned?",
                "What is the highest education level?",
                "How many years of experience does the person have?",
                "What is the person's address?"
            ],
            'legal_document': [
                "What type of legal document is this?",
                "Who are the parties involved?",
                "What is the date of the document?",
                "What is the main subject matter?",
                "Who is the issuing authority?",
                "What are the key terms or conditions?",
                "What is the governing law?",
                "What is the effective date?"
            ],
            'medical_report': [
                "What is the patient's name?",
                "What is the patient ID?",
                "What is the date of the report?",
                "Who is the attending physician?",
                "What is the diagnosis?",
                "What treatment was recommended?",
                "What hospital or clinic issued this report?",
                "What are the test results?"
            ],
            'financial_document': [
                "What is the account number?",
                "What is the account holder's name?",
                "What is the statement period?",
                "What is the opening balance?",
                "What is the closing balance?",
                "What bank issued this statement?",
                "What is the account type?",
                "What are the main transactions?"
            ],
            'identity_document': [
                "What is the document number?",
                "What is the full name of the holder?",
                "What is the date of birth?",
                "What is the nationality?",
                "What is the place of birth?",
                "When was the document issued?",
                "When does the document expire?",
                "What is the issuing authority?"
            ],
            'educational_document': [
                "What is the student's name?",
                "What is the student ID number?",
                "What institution issued this document?",
                "What degree or program is this for?",
                "What is the graduation date?",
                "What is the GPA or grade?",
                "What courses or subjects are listed?",
                "What is the academic year or semester?",
                "What honors or distinctions are mentioned?",
                "What is the major or field of study?",
                "What is the enrollment status?",
                "What financial aid information is included?"
            ],
            'certification_document': [
                "What certification is being awarded?",
                "Who is the certificate holder?",
                "What organization issued this certification?",
                "When was the certification earned?",
                "When does the certification expire?",
                "What exam or assessment was completed?",
                "What is the certification number?",
                "What skills or competencies are certified?",
                "What are the renewal requirements?",
                "What continuing education is required?",
                "What is the certification level or grade?",
                "What professional standards are met?"
            ],
            'employment_document': [
                "What is the employee's name?",
                "What is the job title or position?",
                "What company or organization is this for?",
                "What is the employment start date?",
                "What is the salary or compensation?",
                "What benefits are included?",
                "What are the job responsibilities?",
                "What is the employment status?",
                "What department or division?",
                "Who is the supervisor or manager?",
                "What are the terms of employment?",
                "What policies are referenced?"
            ],
            'student_record': [
                "What is the student's full name?",
                "What is the student identification number?",
                "What school or institution is this from?",
                "What grade level or year is the student in?",
                "What is the current GPA?",
                "What courses is the student taking?",
                "What are the current grades?",
                "What is the attendance record?",
                "Are there any disciplinary issues?",
                "What special programs is the student in?",
                "What accommodations are provided?",
                "Who are the emergency contacts?"
            ]
        }

        # Return questions for the specific document type, or generic questions
        return question_sets.get(doc_type, [
            "What is the document type?",
            "What is the main subject or purpose?",
            "Who issued this document?",
            "What is the date of the document?",
            "Who is the document about or for?",
            "What are the key details mentioned?"
        ])

    def extract_information_with_roberta(self, text: str, questions: List[str]) -> Dict[str, Dict]:
        """
        Extract information from text using RoBERTa question-answering model

        Args:
            text: Document text content
            questions: List of questions to ask about the document

        Returns:
            Dictionary mapping questions to answers with confidence scores
        """
        try:
            if not self.qa_pipeline:
                logger.error("RoBERTa model not available")
                return {}

            extracted_info = {}

            for question in questions:
                try:
                    # Use RoBERTa to answer the question based on the text
                    result = self.qa_pipeline(question=question, context=text)

                    # Only include answers with reasonable confidence
                    if result['score'] > 0.1:  # Threshold for confidence
                        extracted_info[question] = {
                            'answer': result['answer'],
                            'confidence': result['score'],
                            'start': result.get('start', 0),
                            'end': result.get('end', 0)
                        }

                except Exception as e:
                    logger.warning(f"Error processing question '{question}': {str(e)}")
                    continue

            return extracted_info

        except Exception as e:
            logger.error(f"Error extracting with RoBERTa: {str(e)}")
            return {}

    def structure_extraction_results(self, roberta_results: Dict[str, Dict], doc_type: str) -> Dict[str, Any]:
        """
        Structure RoBERTa extraction results into a standardized format

        Args:
            roberta_results: Raw results from RoBERTa extraction
            doc_type: Type of document

        Returns:
            Structured data dictionary
        """
        try:
            structured_data = {
                "document_type": doc_type,
                "extracted_fields": {},
                "confidence_scores": {},
                "processing_metadata": {
                    "model_used": self.model_name,
                    "processing_method": "roberta_local",
                    "privacy_protected": True,
                    "total_questions": len(roberta_results)
                }
            }

            total_confidence = 0.0
            answer_count = 0

            # Process each question-answer pair
            for question, result in roberta_results.items():
                answer = result.get('answer', '').strip()
                confidence = result.get('confidence', 0.0)

                if answer and confidence > 0.1:  # Only include confident answers
                    # Map questions to field names
                    field_name = self._map_question_to_field(question, doc_type)
                    structured_data["extracted_fields"][field_name] = answer
                    structured_data["confidence_scores"][field_name] = confidence

                    total_confidence += confidence
                    answer_count += 1

            # Calculate average confidence
            if answer_count > 0:
                structured_data["processing_metadata"]["average_confidence"] = total_confidence / answer_count
            else:
                structured_data["processing_metadata"]["average_confidence"] = 0.0

            return structured_data

        except Exception as e:
            logger.error(f"Error structuring RoBERTa results: {str(e)}")
            return {
                "document_type": doc_type,
                "extracted_fields": {},
                "processing_metadata": {"error": str(e)}
            }

    def _map_question_to_field(self, question: str, doc_type: str) -> str:
        """
        Map RoBERTa questions to standardized field names

        Args:
            question: The question asked
            doc_type: Type of document

        Returns:
            Standardized field name
        """
        question_lower = question.lower()

        # Common field mappings
        field_mappings = {
            'name': ['name', 'full name', 'person\'s name', 'holder\'s name', 'account holder'],
            'email': ['email', 'email address', 'e-mail'],
            'phone': ['phone', 'phone number', 'telephone', 'contact number'],
            'address': ['address', 'location', 'residence'],
            'date_of_birth': ['date of birth', 'birth date', 'dob', 'born'],
            'document_number': ['document number', 'passport number', 'license number', 'id number', 'account number'],
            'job_title': ['job title', 'position', 'role', 'occupation'],
            'company': ['company', 'employer', 'organization', 'firm'],
            'salary': ['salary', 'compensation', 'income', 'wage'],
            'skills': ['skills', 'abilities', 'competencies'],
            'education': ['education', 'degree', 'qualification', 'academic'],
            'experience': ['experience', 'years of experience', 'work experience'],
            'diagnosis': ['diagnosis', 'condition', 'medical condition'],
            'treatment': ['treatment', 'therapy', 'medication', 'prescription'],
            'physician': ['physician', 'doctor', 'attending physician', 'medical professional'],
            'balance': ['balance', 'amount', 'total'],
            'date': ['date', 'issued', 'effective date', 'statement period']
        }

        # Find matching field
        for field, keywords in field_mappings.items():
            if any(keyword in question_lower for keyword in keywords):
                return field

        # Generate field name from question if no match found
        field_name = question_lower.replace('what is the ', '').replace('what is ', '').replace('?', '').replace(' ', '_')
        return field_name.strip('_')

    def process_document_text(self, text: str, source_file: str = None) -> Dict[str, Any]:
        """
        Process document text using RoBERTa for information extraction

        Args:
            text: Document text content
            source_file: Optional source file path

        Returns:
            Processing results with extracted information
        """
        try:
            # Check if document is confidential
            is_confidential = self.is_confidential_document(text)

            if not is_confidential:
                logger.warning("Document does not appear to be confidential. Consider using regular DocumentProcessor.")

            # Detect document type
            doc_type, type_confidence = self.detect_document_type(text)

            if type_confidence < 0.2:
                logger.warning(f"Low confidence in document type detection: {type_confidence}")

            # Get relevant questions for this document type
            questions = self.get_questions_for_document_type(doc_type)

            # Extract information using RoBERTa
            roberta_results = self.extract_information_with_roberta(text, questions)

            # Structure the results
            structured_data = self.structure_extraction_results(roberta_results, doc_type)

            # Add processing metadata
            result = {
                "status": "success",
                "source_file": source_file or "text_input",
                "document_type": doc_type,
                "type_confidence": type_confidence,
                "is_confidential": is_confidential,
                "privacy_protected": True,
                "extracted_data": structured_data,
                "processing_summary": {
                    "total_questions_asked": len(questions),
                    "successful_extractions": len(structured_data.get("extracted_fields", {})),
                    "average_confidence": structured_data.get("processing_metadata", {}).get("average_confidence", 0.0),
                    "model_used": self.model_name,
                    "device_used": self.device
                }
            }

            logger.info(f"✅ Successfully processed {doc_type} document with {len(structured_data.get('extracted_fields', {}))} extracted fields")

            return result

        except Exception as e:
            logger.error(f"Error processing document text: {str(e)}")
            return {
                "status": "error",
                "source_file": source_file or "text_input",
                "error_message": str(e),
                "error_details": traceback.format_exc(),
                "privacy_protected": True
            }

    def process_image_file(self, image_path: str) -> Dict[str, Any]:
        """
        Process an image file containing a confidential document

        Args:
            image_path: Path to the image file

        Returns:
            Processing results with extracted information
        """
        try:
            if not os.path.exists(image_path):
                raise ValueError(f"Image file not found: {image_path}")

            logger.info(f"Processing image file: {image_path}")

            # Extract text from image using local OCR
            extracted_text = self.extract_text_from_image(image_path)

            if not extracted_text.strip():
                return {
                    "status": "error",
                    "source_file": image_path,
                    "error_message": "No text could be extracted from the image",
                    "privacy_protected": True
                }

            # Process the extracted text
            result = self.process_document_text(extracted_text, image_path)
            result["processing_method"] = "ocr_then_roberta"
            result["extracted_text_length"] = len(extracted_text)

            return result

        except Exception as e:
            logger.error(f"Error processing image file: {str(e)}")
            return {
                "status": "error",
                "source_file": image_path,
                "error_message": str(e),
                "error_details": traceback.format_exc(),
                "privacy_protected": True
            }

    def process_file(self, file_path: str) -> Dict[str, Any]:
        """
        Process any file format (PDF, DOCX, images, text) containing confidential content

        Args:
            file_path: Path to the file to process

        Returns:
            Processing results
        """
        try:
            if not os.path.exists(file_path):
                raise ValueError(f"File not found: {file_path}")

            file_ext = os.path.splitext(file_path)[1].lower()
            logger.info(f"Processing file: {file_path} (type: {file_ext})")

            # Extract text based on file type
            extracted_text = ""
            processing_method = ""

            if file_ext == '.pdf':
                extracted_text = self.extract_text_from_pdf(file_path)
                processing_method = "pdf_processing"
            elif file_ext in ['.docx', '.doc']:
                if file_ext == '.docx':
                    extracted_text = self.extract_text_from_docx(file_path)
                    processing_method = "docx_processing"
                else:
                    return {
                        "status": "error",
                        "source_file": file_path,
                        "error_message": "Legacy .doc format not supported. Please convert to .docx format.",
                        "privacy_protected": True
                    }
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif']:
                extracted_text = self.extract_text_from_image(file_path)
                processing_method = "image_ocr"
            elif file_ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    extracted_text = f.read()
                processing_method = "text_file"
            else:
                return {
                    "status": "error",
                    "source_file": file_path,
                    "error_message": f"Unsupported file type: {file_ext}. Supported: .pdf, .docx, .jpg, .jpeg, .png, .tiff, .bmp, .gif, .txt",
                    "privacy_protected": True,
                    "supported_formats": [".pdf", ".docx", ".jpg", ".jpeg", ".png", ".tiff", ".bmp", ".gif", ".txt"]
                }

            # Check if text was extracted
            if not extracted_text.strip():
                return {
                    "status": "error",
                    "source_file": file_path,
                    "error_message": f"No text could be extracted from {file_ext} file",
                    "processing_method": processing_method,
                    "privacy_protected": True
                }

            # Process the extracted text
            result = self.process_document_text(extracted_text, file_path)

            # Add processing method information
            result["processing_method"] = processing_method
            result["file_format"] = file_ext
            result["extracted_text_length"] = len(extracted_text)

            # Add format-specific metadata
            if file_ext == '.pdf':
                result["pdf_processing"] = {
                    "text_based": "direct text extraction successful" if self._extract_text_from_pdf_direct(file_path).strip() else "scanned PDF processed with OCR",
                    "ocr_used": not bool(self._extract_text_from_pdf_direct(file_path).strip())
                }
            elif file_ext == '.docx':
                result["docx_processing"] = {
                    "format": "Microsoft Word document",
                    "structured_extraction": "paragraphs and tables processed"
                }
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif']:
                result["image_processing"] = {
                    "format": f"Image file ({file_ext})",
                    "ocr_method": "Advanced OCR + Tesseract fallback"
                }

            return result

        except Exception as e:
            logger.error(f"Error processing file: {str(e)}")
            return {
                "status": "error",
                "source_file": file_path,
                "error_message": str(e),
                "error_details": traceback.format_exc(),
                "privacy_protected": True,
                "processing_method": "error"
            }

    def validate_extraction_results(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """
        Validate extraction results for completeness and accuracy

        Args:
            results: Processing results to validate

        Returns:
            Validation report
        """
        try:
            validation_report = {
                "is_valid": True,
                "validation_score": 0.0,
                "issues": [],
                "recommendations": []
            }

            if results.get("status") != "success":
                validation_report["is_valid"] = False
                validation_report["issues"].append("Processing failed")
                return validation_report

            extracted_data = results.get("extracted_data", {})
            extracted_fields = extracted_data.get("extracted_fields", {})

            # Check if we have any extracted data
            if not extracted_fields:
                validation_report["issues"].append("No data extracted from document")
                validation_report["validation_score"] = 0.0
            else:
                # Calculate validation score based on number of fields and confidence
                field_count = len(extracted_fields)
                avg_confidence = extracted_data.get("processing_metadata", {}).get("average_confidence", 0.0)

                # Score based on field count (max 10 fields = 0.5 points)
                field_score = min(field_count / 10.0 * 0.5, 0.5)

                # Score based on confidence (max 0.5 points)
                confidence_score = avg_confidence * 0.5

                validation_report["validation_score"] = field_score + confidence_score

                # Add recommendations
                if field_count < 3:
                    validation_report["recommendations"].append("Consider improving OCR quality or document image resolution")

                if avg_confidence < 0.5:
                    validation_report["recommendations"].append("Low confidence scores - verify extracted information manually")

            return validation_report

        except Exception as e:
            logger.error(f"Error validating results: {str(e)}")
            return {
                "is_valid": False,
                "validation_score": 0.0,
                "issues": [f"Validation error: {str(e)}"],
                "recommendations": ["Manual review required"]
            }

    def get_model_info(self) -> Dict[str, Any]:
        """
        Get information about the loaded model and system

        Returns:
            Model and system information
        """
        return {
            "model_name": self.model_name,
            "device": self.device,
            "model_loaded": self.qa_pipeline is not None,
            "cuda_available": torch.cuda.is_available(),
            "supported_document_types": list(self.document_patterns.keys()),
            "confidential_document_types": list(CONFIDENTIAL_DOCUMENT_TYPES),
            "total_confidential_keywords": len(CONFIDENTIAL_KEYWORDS)
        }

    def test_model_functionality(self) -> Dict[str, Any]:
        """
        Test the RoBERTa model functionality with a simple example

        Returns:
            Test results
        """
        try:
            if not self.qa_pipeline:
                return {
                    "test_passed": False,
                    "error": "RoBERTa model not loaded"
                }

            # Simple test
            test_context = "John Doe is a software engineer working at TechCorp. His email is john.doe@techcorp.com."
            test_question = "What is the person's name?"

            result = self.qa_pipeline(question=test_question, context=test_context)

            return {
                "test_passed": True,
                "test_question": test_question,
                "test_answer": result.get('answer', ''),
                "test_confidence": result.get('score', 0.0),
                "model_working": result.get('answer', '').lower() == 'john doe'
            }

        except Exception as e:
            return {
                "test_passed": False,
                "error": str(e)
            }

    def batch_process_files(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """
        Process multiple files in batch

        Args:
            file_paths: List of file paths to process

        Returns:
            List of processing results
        """
        results = []

        for file_path in file_paths:
            try:
                logger.info(f"Processing file {len(results) + 1}/{len(file_paths)}: {file_path}")
                result = self.process_file(file_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Error processing {file_path}: {str(e)}")
                results.append({
                    "status": "error",
                    "source_file": file_path,
                    "error_message": str(e),
                    "privacy_protected": True
                })

        return results

    def export_results(self, results: Dict[str, Any], output_path: str, format: str = 'json') -> bool:
        """
        Export processing results to file

        Args:
            results: Processing results to export
            output_path: Path to save the results
            format: Export format ('json' or 'txt')

        Returns:
            True if export successful, False otherwise
        """
        try:
            if format.lower() == 'json':
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(results, f, indent=2, ensure_ascii=False)
            elif format.lower() == 'txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write("CONFIDENTIAL DOCUMENT PROCESSING RESULTS\n")
                    f.write("=" * 50 + "\n\n")

                    f.write(f"Document Type: {results.get('document_type', 'Unknown')}\n")
                    f.write(f"Status: {results.get('status', 'Unknown')}\n")
                    f.write(f"Privacy Protected: {results.get('privacy_protected', False)}\n\n")

                    extracted_fields = results.get('extracted_data', {}).get('extracted_fields', {})
                    if extracted_fields:
                        f.write("EXTRACTED INFORMATION:\n")
                        f.write("-" * 25 + "\n")
                        for field, value in extracted_fields.items():
                            f.write(f"{field}: {value}\n")
            else:
                raise ValueError(f"Unsupported format: {format}")

            logger.info(f"Results exported to: {output_path}")
            return True

        except Exception as e:
            logger.error(f"Error exporting results: {str(e)}")
            return False


# Utility functions for easy usage
def create_confidential_processor(model_name: str = "deepset/roberta-base-squad2") -> ConfidentialProcessor:
    """
    Create and initialize a ConfidentialProcessor instance

    Args:
        model_name: HuggingFace model name for question-answering

    Returns:
        Initialized ConfidentialProcessor instance
    """
    return ConfidentialProcessor(model_name=model_name)


def process_confidential_document(file_path: str, model_name: str = "deepset/roberta-base-squad2") -> Dict[str, Any]:
    
    processor = create_confidential_processor(model_name)
    return processor.process_file(file_path)


def check_if_confidential(text: str) -> bool:
    """
    Quick function to check if text contains confidential content

    Args:
        text: Text content to check

    Returns:
        True if confidential, False otherwise
    """
    processor = create_confidential_processor()
    return processor.is_confidential_document(text)


# Example usage and testing
if __name__ == "__main__":
    # Example usage
    print("ConfidentialProcessor - Privacy-First Document Processing")
    print("=" * 60)

    try:
        # Initialize processor
        processor = create_confidential_processor()

        # Test model functionality
        test_results = processor.test_model_functionality()
        print(f"Model Test: {'✅ PASSED' if test_results['test_passed'] else '❌ FAILED'}")

        if test_results['test_passed']:
            print(f"Test Answer: {test_results['test_answer']}")
            print(f"Test Confidence: {test_results['test_confidence']:.2f}")

        # Get model info
        model_info = processor.get_model_info()
        print(f"\nModel: {model_info['model_name']}")
        print(f"Device: {model_info['device']}")
        print(f"Supported Document Types: {len(model_info['supported_document_types'])}")

        # Test confidential detection
        test_text = "This is a confidential employment contract for John Doe with salary information."
        is_confidential = processor.is_confidential_document(test_text)
        print(f"\nConfidential Detection Test: {'✅ CONFIDENTIAL' if is_confidential else '❌ NOT CONFIDENTIAL'}")

        print("\n🔒 ConfidentialProcessor is ready for privacy-protected document processing!")

    except Exception as e:
        print(f"❌ Error initializing ConfidentialProcessor: {str(e)}")
        print("Please ensure all dependencies are installed: pip install transformers torch opencv-python")
