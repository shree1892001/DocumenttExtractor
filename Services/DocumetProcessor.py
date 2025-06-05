import fitz
import tempfile
from typing import List, Dict, Any, Optional, Tuple
import os
import json
import re
from docx import Document
from Common.constants import *
from dataclasses import dataclass
from Extractor.ImageExtractor import ImageTextExtractor
from Extractor.Paddle import flatten_json
from Factories.DocumentFactory import (
    DocumentExtractorFactory,
    DocumentExtractor,
    TextExtractorFactory,
    BaseTextExtractor,
    DocxTextExtractor,
    PdfTextExtractor,
    ImageTextExtractor,
    TextExtractor
)
import logging
import google.generativeai as genai
from abc import ABC, abstractmethod
import pytesseract
from PIL import Image

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class DocumentInfo:
    document_type: str
    confidence_score: float
    extracted_data: Dict[str, Any]
    matched_fields: Dict[str, Any]


class BaseTextExtractor(ABC):
    def __init__(self, api_key: str):
        self.api_key = api_key
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        pass


class TextProcessor:
    def __init__(self, api_key: str):
        self.api_key = api_key
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    def process_text(self, text: str, prompt: str) -> str:
        """Process text using Gemini without image handling"""
        try:
            response = self.model.generate_content([prompt, text])
            return response.text
        except Exception as e:
            logger.error(f"Error processing text with Gemini: {str(e)}")
            raise


class TemplateMatcher:
    document_type_mapping = {
        "aadhaar": "aadhaar_card",
        "aadhaarcard": "aadhaar_card",
        "aadhar": "aadhaar_card",
        "aadharcard": "aadhaar_card",
        "pan": "pan_card",
        "pancard": "pan_card",
        "license": "license",
        "driving license": "license",
        "dl": "license"
    }

    SUPPORTED_EXTENSIONS = {'.docx', '.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.bmp'}

    MIN_CONFIDENCE_THRESHOLD = 0.4
    HIGH_CONFIDENCE_THRESHOLD = 0.8

    def __init__(self, api_key: str, templates_dir: str = "D:\\imageextractor\\identites\\Templates"):
        self.templates_dir = templates_dir
        self.api_key = api_key
        self.text_processor = TextProcessor(api_key)

        if not os.path.exists(self.templates_dir):
            logger.error(f"Templates directory does not exist: {self.templates_dir}")
            raise ValueError(f"Templates directory not found: {self.templates_dir}")

        self.templates = self._load_templates()
        if not self.templates:
            logger.error(f"No valid templates found in directory: {self.templates_dir}")
            raise ValueError(f"No valid templates found in directory: {self.templates_dir}")

    def _load_templates(self) -> Dict[str, Dict[str, Any]]:
        """Load all templates from the templates directory"""
        templates = {}
        try:
            files = os.listdir(self.templates_dir)
            if not files:
                logger.error(f"Templates directory is empty: {self.templates_dir}")
                return templates

            logger.info(f"Found {len(files)} files in templates directory")

            for filename in files:
                file_path = os.path.join(self.templates_dir, filename)
                file_ext = os.path.splitext(filename)[1].lower()

                if file_ext not in self.SUPPORTED_EXTENSIONS:
                    logger.debug(f"Skipping unsupported file format: {filename}")
                    continue

                if os.path.isfile(file_path):
                    try:

                        doc_type = filename.replace('sample_', '').replace('.docx', '').replace('.pdf', '').replace(
                            '.jpg', '').replace('.png', '')
                        doc_type = self._standardize_document_type(doc_type)

                        logger.info(f"Processing template: {filename} as {doc_type}")

                        text_extractor = TextExtractorFactory.create_extractor(file_path, self.api_key)
                        template_text = text_extractor.extract_text(file_path)

                        if template_text:
                            templates[doc_type] = {
                                'content': template_text,
                                'fields': self._extract_fields_from_text(template_text),
                                'structure': template_text
                            }
                            logger.info(f"Successfully loaded template: {filename}")
                        else:
                            logger.warning(f"No text extracted from template: {filename}")
                    except Exception as e:
                        logger.warning(f"Error processing template file {filename}: {str(e)}")
                        continue

            if not templates:
                logger.error(f"No valid templates found in directory: {self.templates_dir}")
            else:
                logger.info(f"Successfully loaded {len(templates)} templates")

            return templates

        except Exception as e:
            logger.error(f"Error loading templates from directory {self.templates_dir}: {str(e)}")
            raise

    def _extract_fields_from_text(self, text: str) -> set:
        """Extract potential field names from text"""
        fields = set()

        field_patterns = [
            r'\{([^}]+)\}',
            r'([^:]+):',
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'
        ]

        for pattern in field_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    field = match[0] if match[0] else match[1]
                else:
                    field = match
                fields.add(field.strip())

        return fields

    def _standardize_document_type(self, doc_type: str) -> str:
        """Standardize document type to match factory mappings"""
        doc_type = doc_type.lower().strip()
        mapping = {
            "aadhaar": "aadhaar_card",
            "aadhaarcard": "aadhaar_card",
            "aadhar": "aadhaar_card",
            "aadharcard": "aadhaar_card",
            "pan": "pan_card",
            "pancard": "pan_card",
            "license": "license",
            "driving license": "license",
            "dl": "license",
            "passport": "passport"
        }
        return mapping.get(doc_type, doc_type)

    def _create_template_matching_prompt(self, input_text: str, template_content: str, doc_type: str) -> str:
        """Create a prompt for Gemini to match document against template"""

        doc_info = self.document_system.get_document_type(input_text)

        passport_instructions = ""
        if doc_type.lower() == 'passport':
            passport_instructions = """
            For passport documents, pay special attention to:
            1. Passport number format and location
            2. Type/Category of passport
            3. Country code and issuing country
            4. Personal information (name, nationality, DOB, etc.)
            5. Dates (issue, expiry)
            6. Authority information
            7. Machine readable zone (MRZ) if present
            8. Security features and official elements

            Common passport layouts:
            - Standard passport layout with personal details
            - Machine readable zone at the bottom
            - Official stamps and seals
            - Security features like holograms
            - Country-specific formatting

            Be lenient with:
            - Different passport formats from various countries
            - Variations in field labels and positions
            - Different date formats
            - Various security feature implementations
            - Regional variations in layout
            """

        return f"""
            You are a document analysis expert specializing in document verification. Your task is to analyze if the given document matches the {doc_type} template structure.

            Document Information:
            Type: {doc_info['document_type']}
            Category: {doc_info.get('category', 'unknown')}
            Confidence: {doc_info.get('confidence', 0.0)}

            Document Text:
            {input_text}

            Template Structure:
            {template_content}

            {passport_instructions}

            Please analyze the document and return a JSON response with:
            1. Whether this document matches the template (true/false)
            2. Confidence score (0-1) based on:
               - Key field presence (0.3 weight)
               - Content similarity (0.3 weight)
               - Structure match (0.2 weight)
               - Format consistency (0.2 weight)

            3. Extracted fields based on the template structure
            4. Any additional relevant information found

            Format the response as:
            {{
                "matches_template": true/false,
                "confidence_score": 0.0-1.0,
                "extracted_fields": {{
                    "field1": "value1",
                    "field2": "value2"
                }},
                "additional_info": "any additional relevant information",
                "matching_details": {{
                    "field_match": 0.0-1.0,
                    "structure_match": 0.0-1.0,
                    "format_match": 0.0-1.0,
                    "content_match": 0.0-1.0
                }},
                "document_info": {{
                    "type": "{doc_info['document_type']}",
                    "category": "{doc_info.get('category', 'unknown')}",
                    "confidence": {doc_info.get('confidence', 0.0)}
                }}
            }}

            Important:
            - Consider partial matches if key fields are present
            - Account for variations in formatting and layout
            - Be lenient with confidence scoring if key identifiers are present
            - Consider a match if confidence score is above 0.4 and key fields are present
            - Look for document-specific security features and official elements
            - Consider regional variations in document formats
            - Account for different versions of the same document type
            """

    def _get_document_specific_indicators(self, doc_type: str) -> str:
        """Get document-specific indicators for matching"""
        indicators = {
            "aadhaar_card": """
                        - 12-digit Aadhaar number
                        - Name of the holder
                        - Date of birth
                        - Gender
                        - Address
                        - UIDAI logo or text
                        - QR code or barcode
                        """,
            "license": """
                        - License number
                        - Name of the holder
                        - Date of birth
                        - Address
                        - License type/class
                        - Issue and expiry dates
                        - Issuing authority
                        """,
            "pan_card": """
                        - 10-character PAN number
                        - Name of the holder
                        - Father's name
                        - Date of birth
                        - PAN card number format
                        - Income Tax Department text
                        """
        }
        return indicators.get(doc_type, "Look for standard identity document fields and structure.")

    def match_document(self, file_path: str, section_text: str = None) -> Dict[str, Any]:
        """Match document against templates using Gemini"""
        if not self.templates:
            raise ValueError("No templates available for matching")

        best_match = {
            'document_type': None,
            'confidence': 0.0,
            'matched_fields': {},
            'additional_info': None,
            'template_id': None,
            'matching_details': {}
        }

        try:

            input_text = section_text
            if not input_text and file_path:
                text_extractor = TextExtractorFactory.create_extractor(file_path, self.api_key)
                max_retries = 3
                for attempt in range(max_retries):
                    try:
                        input_text = text_extractor.extract_text(file_path)
                        if input_text:
                            break
                    except Exception as e:
                        if attempt == max_retries - 1:
                            raise
                        logger.warning(f"Attempt {attempt + 1} failed: {str(e)}")
                        continue

            if not input_text:
                raise ValueError("No text content available for matching")

            logger.info(f"Extracted text from document: {len(input_text)} characters")

            for template_id, template in self.templates.items():
                try:
                    template_doc_type = template.get('document_type', 'unknown')
                    prompt = self._create_template_matching_prompt(
                        input_text,
                        template['content'],
                        template_doc_type
                    )
                    response = self.text_processor.process_text(input_text, prompt)
                    match_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())

                    confidence = match_result.get('confidence_score', 0)
                    logger.info(f"Template match result for {template_id}: {confidence}")

                    if confidence > best_match['confidence']:
                        best_match = {
                            'document_type': match_result.get('document_type'),
                            'confidence': confidence,
                            'matched_fields': match_result.get('extracted_fields', {}),
                            'additional_info': match_result.get('additional_info'),
                            'template_id': template_id,
                            'matching_details': match_result.get('matching_details', {})
                        }
                        logger.info(f"Found better match: {template_id} with confidence {confidence}")
                        logger.info(
                            f"Matching details: {json.dumps(match_result.get('matching_details', {}), indent=2)}")
                except Exception as e:
                    logger.warning(f"Error matching template for {template_id}: {str(e)}")
                    continue

            if not best_match['document_type']:
                logger.error("No matching template found for the document")
                raise ValueError("No matching template found for the document")

            logger.info(f"Best match found: {best_match['document_type']} with confidence {best_match['confidence']}")
            return best_match

        except Exception as e:
            logger.error(f"Error matching document: {str(e)}")
            raise


class DocumentClassifier:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.template_matcher = TemplateMatcher(api_key)

    def identify_document_type(self, file_path: str) -> DocumentInfo:
        try:

            template_match = self.template_matcher.match_document(file_path)

            if template_match['document_type']:
                return DocumentInfo(
                    document_type=template_match['document_type'],
                    confidence_score=template_match['confidence'],
                    extracted_data=template_match['matched_fields'],
                    matched_fields=template_match['matched_fields']
                )
            else:
                raise ValueError("Document type not identified")

        except Exception as e:
            logger.exception("Failed to classify document")
            raise ValueError(f"Failed to classify document: {str(e)}")


class DocumentProcessor:
    def __init__(self, api_key: str, templates_dir: str = "D:\\imageextractor\\identites\\Templates"):
        self.api_key = api_key
        self.templates_dir = templates_dir
        self.classifier = DocumentClassifier(api_key)
        self.text_extractor = TextExtractor(api_key)
        self.text_processor = TextProcessor(api_key)
        self.template_matcher = TemplateMatcher(api_key, templates_dir)

        self.document_patterns = {
            'license': [
                r'(?i)license|dl|driving|permit|rto|dmv',
                r'(?i)vehicle|motor|transport',
                r'(?i)driver|driving'
            ],
            'aadhaar_card': [
                r'(?i)aadhaar|aadhar|uidai|unique\s*id',
                r'(?i)आधार|यूआईडीएआई'
            ],
            'pan_card': [
                r'(?i)pan|permanent\s*account|income\s*tax',
                r'(?i)tax\s*id|tax\s*number'
            ],
            'passport': [
                r'(?i)passport|travel\s*doc|nationality',
                r'(?i)immigration|border|customs'
            ]
        }

        self.compiled_patterns = {
            doc_type: [re.compile(pattern) for pattern in patterns]
            for doc_type, patterns in self.document_patterns.items()
        }

        self.MIN_CONFIDENCE_THRESHOLD = 0.4
        self.MIN_GENUINENESS_SCORE = 0.6
        self.VERIFICATION_THRESHOLD = 0.5

    def _standardize_document_type(self, doc_type: str) -> str:
        """Standardize document type using pattern matching"""
        try:
            doc_type = doc_type.lower().strip()

            common_mappings = {
                "aadhaar": "aadhaar_card",
                "aadhaarcard": "aadhaar_card",
                "aadhar": "aadhaar_card",
                "aadharcard": "aadhaar_card",
                "pan": "pan_card",
                "pancard": "pan_card",
                "license": "license",
                "driving license": "license",
                "dl": "license",
                "passport": "passport",
                "travel document": "passport",
                "national passport": "passport",
                "diplomatic passport": "passport",
                "official passport": "passport",
                "service passport": "passport"
            }

            if doc_type in common_mappings:
                return common_mappings[doc_type]

            for standardized_type, patterns in self.compiled_patterns.items():
                for pattern in patterns:
                    if pattern.search(doc_type):
                        logger.debug(f"Matched {doc_type} to {standardized_type} using pattern {pattern.pattern}")
                        return standardized_type

            base_type = self._extract_base_type(doc_type)
            if base_type:
                return base_type

            logger.warning(f"No pattern match found for document type: {doc_type}")
            return doc_type

        except Exception as e:
            logger.error(f"Error standardizing document type: {str(e)}")
            return doc_type

    def _extract_base_type(self, doc_type: str) -> Optional[str]:
        """Extract base document type from filename or text"""
        try:

            doc_type = re.sub(r'(?i)sample_|template_|example_', '', doc_type)
            doc_type = re.sub(r'\.(docx|pdf|jpg|png|jpeg)$', '', doc_type)

            doc_type = re.sub(r'^(indian|florida|california|texas|new_york|uk|us|canada)_', '', doc_type)

            doc_type = re.sub(r'_v\d+', '', doc_type)

            doc_type = re.sub(r'_\d{4}(_\d{2}){0,2}', '', doc_type)

            doc_type = re.sub(r'[^a-z0-9]', '_', doc_type.lower())
            doc_type = re.sub(r'_+', '_', doc_type)
            doc_type = doc_type.strip('_')

            for standardized_type, patterns in self.compiled_patterns.items():
                for pattern in patterns:
                    if pattern.search(doc_type):
                        return standardized_type

            return None

        except Exception as e:
            logger.error(f"Error extracting base type: {str(e)}")
            return None

    def add_document_pattern(self, doc_type: str, pattern: str) -> None:
        """Add a new pattern for document type detection"""
        try:
            if doc_type not in self.document_patterns:
                self.document_patterns[doc_type] = []
                self.compiled_patterns[doc_type] = []

            self.document_patterns[doc_type].append(pattern)
            self.compiled_patterns[doc_type].append(re.compile(pattern))
            logger.info(f"Added new pattern for {doc_type}: {pattern}")

        except Exception as e:
            logger.error(f"Error adding document pattern: {str(e)}")

    def get_document_patterns(self) -> Dict[str, List[str]]:
        """Get current document patterns"""
        return {
            doc_type: [pattern.pattern for pattern in patterns]
            for doc_type, patterns in self.compiled_patterns.items()
        }

    def _extract_text_from_docx_images(self, file_path: str) -> List[Dict[str, Any]]:
        """Extract text from images in a DOCX file and return list of document segments"""
        try:
            doc = Document(file_path)
            document_segments = []
            image_count = 0

            # First, extract all images from the document
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
                    with tempfile.TemporaryDirectory() as temp_dir:
                        image_path = os.path.join(temp_dir, f"image_{image_count}.png")
                        try:
                            # Get the image data
                            image_data = rel.target_part.blob

                            # Save the image
                            with open(image_path, 'wb') as f:
                                f.write(image_data)

                            # Process the image
                            img = Image.open(image_path)

                            # Try Tesseract OCR first
                            ocr_text = pytesseract.image_to_string(img)

                            # If Tesseract results are poor, use Gemini Vision
                            if not self._is_good_ocr_result(ocr_text):
                                logger.info(
                                    f"Tesseract OCR yielded poor results for image {image_count}, trying Gemini Vision")
                                response = self.text_extractor.process_with_gemini(
                                    image_path,
                                    "Extract all text from this document image. Return only the raw text without any formatting."
                                )
                                ocr_text = response.strip()

                            if ocr_text.strip():
                                document_segments.append({
                                    "text": ocr_text.strip(),
                                    "image_path": image_path,
                                    "segment_index": len(document_segments),
                                    "image_number": image_count
                                })
                                logger.info(f"Extracted text from image {image_count}: {len(ocr_text)} characters")
                            else:
                                logger.warning(f"No text could be extracted from image {image_count}")
                        except Exception as e:
                            logger.warning(f"Error processing image {image_count}: {str(e)}")
                            continue

            # Also check for images in paragraphs and tables
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.xpath('.//w:drawing'):
                        image_count += 1
                        with tempfile.TemporaryDirectory() as temp_dir:
                            image_path = os.path.join(temp_dir, f"image_{image_count}.png")
                            try:
                                # Get the image data
                                image_data = run._element.xpath('.//a:blip/@r:embed')[0]
                                image_part = doc.part.related_parts[image_data]

                                # Save the image
                                with open(image_path, 'wb') as f:
                                    f.write(image_part.blob)

                                # Process the image
                                img = Image.open(image_path)

                                # Try Tesseract OCR first
                                ocr_text = pytesseract.image_to_string(img)

                                # If Tesseract results are poor, use Gemini Vision
                                if not self._is_good_ocr_result(ocr_text):
                                    logger.info(
                                        f"Tesseract OCR yielded poor results for image {image_count}, trying Gemini Vision")
                                    response = self.text_extractor.process_with_gemini(
                                        image_path,
                                        "Extract all text from this document image. Return only the raw text without any formatting."
                                    )
                                    ocr_text = response.strip()

                                if ocr_text.strip():
                                    document_segments.append({
                                        "text": ocr_text.strip(),
                                        "image_path": image_path,
                                        "segment_index": len(document_segments),
                                        "image_number": image_count
                                    })
                                    logger.info(f"Extracted text from image {image_count}: {len(ocr_text)} characters")
                                else:
                                    logger.warning(f"No text could be extracted from image {image_count}")
                            except Exception as e:
                                logger.warning(f"Error processing image {image_count}: {str(e)}")
                                continue

            if not document_segments:
                logger.warning("No text could be extracted from images in DOCX")
                return []

            logger.info(f"Successfully extracted text from {len(document_segments)} images")
            return document_segments

        except Exception as e:
            logger.error(f"Error extracting text from DOCX images: {str(e)}")
            return []

    def _split_into_chunks(self, text: str) -> List[str]:
        """Split text into chunks based on document boundaries"""
        try:
            # Common document separators
            separators = [
                r'\n\s*\n\s*\n',  # Multiple blank lines
                r'[-=]{3,}',  # Lines of dashes or equals
                r'_{3,}',  # Lines of underscores
                r'\*{3,}',  # Lines of asterisks
                r'Page \d+',  # Page numbers
                r'Document \d+',  # Document numbers
                r'Copy \d+',  # Copy numbers
                r'Original',  # Original document marker
                r'Duplicate',  # Duplicate document marker
                r'COPY',  # COPY marker
                r'ORIGINAL'  # ORIGINAL marker
            ]

            # Combine separators into a single pattern
            separator_pattern = '|'.join(f'({sep})' for sep in separators)

            # Split text into chunks
            chunks = re.split(separator_pattern, text)

            # Filter out empty chunks and separators
            chunks = [chunk.strip() for chunk in chunks if chunk and not any(sep in chunk for sep in separators)]

            # If no chunks were found, return the whole text as one chunk
            if not chunks:
                return [text]

            return chunks

        except Exception as e:
            logger.error(f"Error splitting text into chunks: {str(e)}")
            return [text]

    def _process_multiple_documents(self, text: str, source_file: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process text that may contain multiple documents"""
        try:
            results = []

            # Split text into potential document chunks
            chunks = self._split_into_chunks(text)
            logger.info(f"Split text into {len(chunks)} potential document chunks")

            for chunk_index, chunk in enumerate(chunks):
                if not chunk.strip():
                    continue

                logger.info(f"Processing chunk {chunk_index + 1} of {len(chunks)}")

                # Process each chunk as a separate document
                result = self._process_text_content(chunk, source_file, min_confidence)

                if result:
                    # Add chunk information to the result
                    result["chunk_index"] = chunk_index + 1
                    result["total_chunks"] = len(chunks)
                    results.append(result)

            return results

        except Exception as e:
            logger.error(f"Error processing multiple documents: {str(e)}")
            return []

    def process_file(self, file_path: str, min_confidence: float = 0.0) -> List[Dict[str, Any]]:
        """Process a file that may contain multiple documents"""
        if file_path.lower().endswith('.pdf'):
            return self._process_pdf(file_path, min_confidence)
        elif file_path.lower().endswith('.docx'):
            return self._process_docx(file_path, min_confidence)
        else:
            result = self._process_single_image(file_path, min_confidence)
            return [result] if result else []

    def _process_docx(self, file_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a DOCX file that may contain multiple documents"""
        results = []
        processed_images = set()  # Track processed images to avoid duplicates

        try:
            # First, try to extract and process images
            logger.info("Attempting to extract and process images from DOCX")
            document_segments = self._extract_text_from_docx_images(file_path)

            if document_segments:
                logger.info(f"Found {len(document_segments)} images in the document")
                for segment in document_segments:
                    # Skip if we've already processed this image
                    if segment["image_path"] in processed_images:
                        continue

                    processed_images.add(segment["image_path"])

                    if segment["text"].strip():
                        logger.info(f"Processing image {segment['image_number']}")
                        # Try to process as multiple documents first
                        multi_doc_results = self._process_multiple_documents(segment["text"], file_path, min_confidence)
                        if multi_doc_results:
                            for result in multi_doc_results:
                                result["image_number"] = segment["image_number"]
                                result["processing_method"] = "ocr"
                            results.extend(multi_doc_results)
                        else:
                            # Process as single document if multiple document detection failed
                            result = self._process_text_content(segment["text"], file_path, min_confidence)
                            if result:
                                result["image_number"] = segment["image_number"]
                                result["processing_method"] = "ocr"
                                results.append(result)

            # Then process the document text
            doc = Document(file_path)
            text_content = []

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())

            # Combine all text
            text = '\n'.join(text_content)

            if text.strip():
                logger.info(f"Processing text content from DOCX: {len(text)} characters")

                # Try to process as multiple documents first
                multi_doc_results = self._process_multiple_documents(text, file_path, min_confidence)
                if multi_doc_results:
                    for result in multi_doc_results:
                        result["processing_method"] = "direct_text"
                    results.extend(multi_doc_results)
                else:
                    # If multiple document processing didn't yield results, process as single document
                    result = self._process_text_content(text, file_path, min_confidence)
                    if result:
                        result["processing_method"] = "direct_text"
                        results.append(result)

            if not results:
                logger.warning(f"No valid content found in DOCX file: {file_path}")
                return [{
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": file_path,
                    "rejection_reason": "No valid content found in document"
                }]

            # Consolidate results to remove duplicates
            consolidated_results = self._consolidate_results(results)

            # Log the results
            logger.info(f"Successfully processed {len(consolidated_results)} unique documents from DOCX file")
            for result in consolidated_results:
                logger.info(f"Document type: {result.get('document_type')}, "
                            f"Processing method: {result.get('processing_method')}, "
                            f"Image number: {result.get('image_number', 'N/A')}")

            return consolidated_results

        except Exception as e:
            logger.error(f"Error processing DOCX file: {str(e)}")
            return [{
                "status": "error",
                "document_type": "unknown",
                "source_file": file_path,
                "rejection_reason": f"Error processing document: {str(e)}"
            }]

    def _consolidate_results(self, results: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Consolidate results to remove duplicates and combine related documents"""
        try:
            consolidated = []
            seen_documents = set()

            for result in results:
                if result.get("status") != "success":
                    continue

                # Create a unique identifier for the document
                doc_type = result.get("document_type", "unknown")
                extracted_data = result.get("extracted_data", {})

                # Create a unique key based on document type and key fields
                if doc_type == "aadhaar_card":
                    key = f"aadhaar_{extracted_data.get('Aadhaar Number', '')}"
                elif doc_type == "license":
                    key = f"license_{extracted_data.get('Name', '')}_{extracted_data.get('Date of Birth', '')}"
                else:
                    key = f"{doc_type}_{json.dumps(extracted_data)}"

                if key not in seen_documents:
                    seen_documents.add(key)
                    consolidated.append(result)

            return consolidated

        except Exception as e:
            logger.error(f"Error consolidating results: {str(e)}")
            return results

    def _process_text_content(self, text: str, source_file: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process text content directly without creating temporary files"""
        try:

            detection_prompt = f"""
                        Analyze this text and determine what type of document it is (license, aadhaar_card, pan_card, or passport).
                        Look for these indicators:
                        - Aadhaar: "Aadhaar", "UIDAI", "आधार", "Unique Identification"
                        - PAN: "PAN", "Permanent Account Number", "Income Tax"
                        - License: "License", "Driving License", "DL", "RTO"
                        - Passport: "Passport", "Passport Number", "Nationality"

                        Text:
                        {text}

                        Return a JSON response with:
                        {{
                            "document_type": "detected_type",
                            "confidence": 0.0-1.0,
                            "reasoning": "explanation"
                        }}

                        Important:
                        - Only return document_type as one of: "license", "aadhaar_card", "pan_card", "passport"
                        - If you cannot confidently determine the type, return "unknown"
                        - Be strict in your confidence scoring
                        - Consider a document type only if you see clear indicators
                        """

            response = self.text_processor.process_text(text, detection_prompt)
            detection_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())

            doc_type = detection_result.get("document_type", "").lower()
            confidence = detection_result.get("confidence", 0.0)

            doc_type = self._standardize_document_type(doc_type)

            if doc_type == "unknown" or confidence < min_confidence:
                logger.warning(f"Could not confidently determine document type (confidence: {confidence})")
                return {
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": source_file,
                    "rejection_reason": f"Document type could not be determined with sufficient confidence (confidence: {confidence})",
                    "confidence": confidence
                }

            logger.info(f"Detected document type: {doc_type} with confidence {confidence}")

            try:
                extractor = DocumentExtractorFactory.get_extractor(doc_type, self.api_key)
            except ValueError as e:
                logger.error(f"Error getting extractor: {str(e)}")
                return {
                    "status": "rejected",
                    "document_type": doc_type,
                    "source_file": source_file,
                    "rejection_reason": f"No extractor available for document type: {doc_type}",
                    "confidence": confidence
                }

            extraction_prompt = f"""
                        You are a document analysis expert. Extract all relevant information from this {doc_type} document.

                        Document Text:
                        {text}

                        For a {doc_type}, look for and extract these specific fields:
                        {self._get_document_specific_fields(doc_type)}

                        Important:
                        1. Extract ALL visible information from the document
                        2. If a field is not found, use "None" as the value
                        3. For dates, use YYYY-MM-DD format
                        4. For numbers, extract them exactly as shown
                        5. For names and addresses, preserve the exact spelling and formatting

                        Return the extracted information in JSON format with this structure:
                        {{
                            "data": {{
                                // All extracted fields
                            }},
                            "confidence": 0.0-1.0,
                            "additional_info": "Any additional relevant information"
                        }}
                        """

            response = self.text_processor.process_text(text, extraction_prompt)
            extracted_data = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
            logger.info(f"Extracted data: {json.dumps(extracted_data, indent=2)}")

            verification_result = self.verify_document(extracted_data, doc_type)

            if not verification_result["is_genuine"]:
                rejection_reason = verification_result.get("rejection_reason",
                                                           "Document failed authenticity verification")
                logger.warning(f"Document rejected - Not genuine: {rejection_reason}")
                return {
                    "status": "rejected",
                    "document_type": doc_type,
                    "source_file": source_file,
                    "rejection_reason": rejection_reason,
                    "verification_result": verification_result
                }

            if extractor.validate_fields(extracted_data):
                flattened_data = flatten_json(extracted_data)
                return {
                    "extracted_data": flattened_data,
                    "status": "success",
                    "confidence": confidence,
                    "document_type": doc_type,
                    "source_file": source_file,
                    "validation_level": "strict",
                    "verification_result": verification_result
                }
            else:
                logger.warning(f"Validation failed for {source_file}")
                return {
                    "status": "rejected",
                    "document_type": doc_type,
                    "source_file": source_file,
                    "rejection_reason": "Document failed validation",
                    "verification_result": verification_result
                }

        except Exception as e:
            logger.exception(f"Error processing text content: {str(e)}")
            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": source_file,
                "rejection_reason": f"Error processing document: {str(e)}"
            }

    def _process_pdf(self, pdf_path: str, min_confidence: float) -> List[Dict[str, Any]]:
        """Process a PDF file that may contain multiple documents"""
        results = []
        pdf_document = fitz.open(pdf_path)

        try:
            # First try to process the entire PDF as one document
            all_text = ""
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text = page.get_text()
                if text.strip():
                    all_text += text + "\n\n"

            if all_text.strip():
                # Try to process as multiple documents first
                multi_doc_results = self._process_multiple_documents(all_text, pdf_path, min_confidence)
                if multi_doc_results:
                    results.extend(multi_doc_results)
                    return results

            # If multiple document processing didn't yield results, process page by page
            for page_num in range(pdf_document.page_count):
                logger.info(f"Processing page {page_num + 1} of {pdf_document.page_count}")

                page = pdf_document[page_num]
                text = page.get_text()

                needs_ocr = self._needs_ocr(text, page)

                if needs_ocr:
                    logger.info(f"Page {page_num + 1} requires OCR processing")
                    images = page.get_images(full=True)

                    if not images:
                        logger.warning(f"No images found on page {page_num + 1} for OCR")
                        continue

                    for img_index, img_info in enumerate(images):
                        with tempfile.TemporaryDirectory() as temp_dir:
                            img_data = pdf_document.extract_image(img_info[0])
                            temp_image_path = os.path.join(temp_dir, f"page_{page_num}_img_{img_index}.png")

                            with open(temp_image_path, "wb") as img_file:
                                img_file.write(img_data["image"])

                            ocr_text = self._perform_ocr(temp_image_path)
                            if ocr_text:
                                # Try to process OCR text as multiple documents
                                multi_doc_results = self._process_multiple_documents(ocr_text, pdf_path, min_confidence)
                                if multi_doc_results:
                                    for result in multi_doc_results:
                                        result["page_number"] = page_num + 1
                                        result["image_index"] = img_index + 1
                                        result["processing_method"] = "ocr"
                                    results.extend(multi_doc_results)
                                else:
                                    # Process as single document if multiple document detection failed
                                    result = self._process_text_content(ocr_text, pdf_path, min_confidence)
                                    if result:
                                        result["page_number"] = page_num + 1
                                        result["image_index"] = img_index + 1
                                        result["processing_method"] = "ocr"
                                        results.append(result)
                else:
                    if text.strip():
                        # Try to process text as multiple documents
                        multi_doc_results = self._process_multiple_documents(text, pdf_path, min_confidence)
                        if multi_doc_results:
                            for result in multi_doc_results:
                                result["page_number"] = page_num + 1
                                result["processing_method"] = "direct_text"
                            results.extend(multi_doc_results)
                        else:
                            # Process as single document if multiple document detection failed
                            result = self._process_text_content(text, pdf_path, min_confidence)
                            if result:
                                result["page_number"] = page_num + 1
                                result["processing_method"] = "direct_text"
                                results.append(result)

            return results
        finally:
            pdf_document.close()

    def _needs_ocr(self, text: str, page) -> bool:
        """Determine if OCR is needed for the page"""
        try:

            if not text.strip():
                return True

            if self._is_character_by_character(text):
                return True

            images = page.get_images(full=True)
            if images:

                if not self._has_meaningful_content(text):
                    return True

            if self._has_document_indicators(page):
                return True

            return False

        except Exception as e:
            logger.error(f"Error determining OCR need: {str(e)}")
            return True

    def _is_character_by_character(self, text: str) -> bool:
        """Check if text appears to be extracted character by character"""
        try:

            patterns = [
                r'[A-Z]\s+[A-Z]\s+[A-Z]',
                r'[a-z]\s+[a-z]\s+[a-z]',
                r'[0-9]\s+[0-9]\s+[0-9]',
                r'[A-Za-z]\s+[A-Za-z]\s+[A-Za-z]'
            ]

            for pattern in patterns:
                if re.search(pattern, text):
                    return True

            if len(text) > 0:
                space_ratio = text.count(' ') / len(text)
                if space_ratio > 0.3:
                    return True

            return False

        except Exception as e:
            logger.error(f"Error checking character-by-character extraction: {str(e)}")
            return False

    def _has_meaningful_content(self, text: str) -> bool:
        """Check if text contains meaningful content"""
        try:

            indicators = [
                r'[A-Z]{2,}',
                r'\d{4,}',
                r'[A-Za-z]+\s+[A-Za-z]+',
                r'[A-Za-z]+\s+\d+',
                r'\d+\s+[A-Za-z]+'
            ]

            indicator_count = 0
            for pattern in indicators:
                if re.search(pattern, text):
                    indicator_count += 1

            return indicator_count >= 2

        except Exception as e:
            logger.error(f"Error checking meaningful content: {str(e)}")
            return False

    def _has_document_indicators(self, page) -> bool:
        """Check if page has indicators of being a document that might need OCR"""
        try:

            images = page.get_images(full=True)
            if not images:
                return False

            blocks = page.get_text("blocks")
            if not blocks:
                return True

            for block in blocks:

                if block[4].strip():
                    if any(indicator in block[4].lower() for indicator in [
                        "name", "date", "address", "signature", "photo",
                        "passport", "license", "id", "number", "issued"
                    ]):
                        return True

            return False

        except Exception as e:
            logger.error(f"Error checking document indicators: {str(e)}")
            return True

    def _perform_ocr(self, image_path: str) -> str:
        """Perform OCR on an image"""
        try:

            img = Image.open(image_path)
            ocr_text = pytesseract.image_to_string(img)

            if not self._is_good_ocr_result(ocr_text):
                logger.info("Tesseract OCR yielded poor results, trying Gemini Vision")
                response = self.text_extractor.process_with_gemini(
                    image_path,
                    "Extract all text from this document image. Return only the raw text without any formatting."
                )
                ocr_text = response.strip()

            return ocr_text

        except Exception as e:
            logger.error(f"Error performing OCR: {str(e)}")
            return ""

    def _is_good_ocr_result(self, text: str) -> bool:
        """Check if OCR result is of good quality"""
        try:
            if not text.strip():
                return False

            if len(text.strip()) < 10:
                return False

            if not self._has_meaningful_content(text):
                return False

            error_patterns = [
                r'[|]{2,}',
                r'[l1]{3,}',
                r'[o0]{3,}',
                r'[rn]{3,}',
                r'\s{3,}'
            ]

            for pattern in error_patterns:
                if re.search(pattern, text):
                    return False

            return True

        except Exception as e:
            logger.error(f"Error checking OCR result quality: {str(e)}")
            return False

    def verify_document(self, extracted_data: Dict[str, Any], doc_type: str) -> Dict[str, Any]:
        """Verify document authenticity and data validity"""
        try:

            verification_prompt = f"""
            You are a document verification expert specializing in {doc_type} verification. 
            Analyze this document data and determine if it is genuine.
            Provide a detailed verification report with specific checks and findings.

            Document Data:
            {json.dumps(extracted_data, indent=2)}

            Document Type: {doc_type}

            Perform the following checks:

            1. Document Authenticity:
               - Verify presence of required fields for {doc_type}
               - Check for official text and formatting
               - Validate document structure and layout
               - Check for security features specific to {doc_type}
               - Verify document quality and printing standards
               - Look for official headers and footers
               - Check for any official stamps or marks

            2. Security Feature Analysis:
               - Look for official seals and stamps
               - Check for watermarks or holograms
               - Verify presence of security patterns
               - Look for QR codes or barcodes
               - Check for any digital signatures
               - Verify any security numbers or codes
               - Look for any anti-counterfeit measures

            3. Data Validation:
               - Verify all identification numbers and their formats
               - Check date formats and their logical consistency
               - Validate name and address formatting
               - Look for any data inconsistencies
               - Verify field relationships and dependencies
               - Check for logical data patterns
               - Validate any reference numbers

            4. Document Quality:
               - Check printing quality
               - Verify text alignment
               - Check for any smudges or marks
               - Verify color consistency
               - Check for any signs of poor quality
               - Verify professional formatting
               - Check for any signs of manipulation

            Return a JSON response with:
            {{
                "is_genuine": true/false,
                "confidence_score": 0.0-1.0,
                "rejection_reason": "reason if not genuine",
                "verification_checks": {{
                    "authenticity": {{
                        "passed": true/false,
                        "details": "explanation",
                        "confidence": 0.0-1.0
                    }},
                    "security_features": {{
                        "passed": true/false,
                        "details": "explanation",
                        "confidence": 0.0-1.0
                    }},
                    "data_validation": {{
                        "passed": true/false,
                        "details": "explanation",
                        "confidence": 0.0-1.0
                    }},
                    "quality": {{
                        "passed": true/false,
                        "details": "explanation",
                        "confidence": 0.0-1.0
                    }}
                }},
                "security_features_found": ["list of security features"],
                "verification_summary": "Overall verification summary",
                "recommendations": ["list of recommendations for improvement"]
            }}

            Important:
            - Be thorough but fair in your analysis
            - Consider document-specific requirements for {doc_type}
            - Account for variations in official document formats
            - Consider both digital and physical security features
            - Look for logical consistency in the data
            - Consider the document's intended use
            - Be lenient with minor formatting variations
            - Focus on key security features and data validity
            - Consider official document standards
            - Account for regional variations in document formats
            """

            response = self.text_processor.process_text(json.dumps(extracted_data), verification_prompt)
            verification_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())

            checks = verification_result.get("verification_checks", {})
            confidence_scores = [
                checks.get("authenticity", {}).get("confidence", 0),
                checks.get("security_features", {}).get("confidence", 0),
                checks.get("data_validation", {}).get("confidence", 0),
                checks.get("quality", {}).get("confidence", 0)
            ]
            overall_confidence = sum(confidence_scores) / len(confidence_scores) if confidence_scores else 0

            verification_result["confidence_score"] = overall_confidence
            verification_result["is_genuine"] = overall_confidence >= self.VERIFICATION_THRESHOLD

            logger.info(f"Document verification results: {json.dumps(verification_result, indent=2)}")

            return verification_result

        except Exception as e:
            logger.exception(f"Error verifying document genuineness: {str(e)}")
            return {
                "is_genuine": False,
                "confidence_score": 0.0,
                "rejection_reason": f"Error during verification: {str(e)}",
                "verification_checks": {
                    "authenticity": {"passed": False, "details": f"Error during verification: {str(e)}",
                                     "confidence": 0.0},
                    "security_features": {"passed": False, "details": "Verification failed", "confidence": 0.0},
                    "data_validation": {"passed": False, "details": "Verification failed", "confidence": 0.0},
                    "quality": {"passed": False, "details": "Verification failed", "confidence": 0.0}
                },
                "security_features_found": [],
                "verification_summary": f"Document verification failed due to error: {str(e)}",
                "recommendations": ["Verification process failed due to technical error"]
            }

    def _validate_with_leniency(self, extracted_data: Dict[str, Any], doc_type: str) -> bool:
        """Validate extracted data with lenient rules for low confidence matches"""
        try:
            data = extracted_data.get("data", {})

            if not data:
                return False

            if doc_type == "license":

                required_fields = ["license_number", "name", "date_of_birth"]
                return any(field in data for field in required_fields)
            elif doc_type == "aadhaar_card":

                required_fields = ["aadhaar_number", "name", "date_of_birth"]
                return any(field in data for field in required_fields)
            elif doc_type == "pan_card":

                required_fields = ["pan_number", "name", "date_of_birth"]
                return any(field in data for field in required_fields)
            else:

                return len(data) > 0

        except Exception as e:
            logger.error(f"Error in lenient validation: {str(e)}")
            return False

    def _get_document_specific_fields(self, doc_type: str) -> str:
        """Get document-specific fields to extract"""
        fields = {
            "license": """
                        - License Number (look for DL number, license number, etc.)
                        - Name (full name of the license holder)
                        - Date of Birth (DOB, birth date)
                        - Address (current address)
                        - Valid From (issue date, date of issue)
                        - Valid Until (expiry date, valid till)
                        - Vehicle Categories (classes of vehicles allowed)
                        - Issuing Authority (RTO, DMV, etc.)
                        """,
            "aadhaar_card": """
                        - Aadhaar Number (12-digit number)
                        - Name (full name)
                        - Gender
                        - Date of Birth
                        - Address
                        - Father's Name
                        - Photo
                        """,
            "pan_card": """
                        - PAN Number (10-character alphanumeric)
                        - Name
                        - Father's Name
                        - Date of Birth
                        - Photo
                        """,
            "passport": """
                        - Passport Number (look for P, A, C, S prefixes)
                        - Type/Category (Regular, Diplomatic, Official, Service)
                        - Country Code (3-letter code)
                        - Surname/Last Name
                        - Given Names
                        - Nationality
                        - Date of Birth
                        - Place of Birth
                        - Gender
                        - Date of Issue
                        - Date of Expiry
                        - Authority/Issuing Country
                        - Personal Number (if present)
                        - Machine Readable Zone (MRZ)
                        - Photo
                        - Signature
                        - Security Features
                        """
        }
        return fields.get(doc_type, "Extract all visible text and information from the document.")

    def _verify_document_genuineness(self, text: str, doc_type: str) -> Dict[str, Any]:
        """Verify if the document is genuine before extraction"""
        verification_result = {
            "is_genuine": False,
            "confidence_score": 0.0,
            "rejection_reason": "",
            "verification_checks": [],
            "security_features_found": []
        }

        try:

            verification_prompt = f"""
                        You are an expert document verification AI. Your task is to determine if this document is genuine BEFORE any data extraction.

                        Document Text:
                        {text}

                        Perform a thorough genuineness check:

                        1. Document Authenticity:
                           - Verify if this is a genuine official document
                           - Check for official government/issuing authority text and formatting
                           - Look for security features and anti-counterfeit measures
                           - Validate document structure and layout
                           - Check for official logos and seals

                        2. Security Feature Analysis:
                           - Check for official seals, watermarks, or holograms
                           - Look for security patterns or microtext
                           - Verify presence of QR codes or barcodes
                           - Check for official government/issuing authority text
                           - Look for security threads or special paper

                        3. Red Flag Detection:
                           - Look for signs of forgery or tampering
                           - Check for missing security features
                           - Identify suspicious patterns
                           - Look for inconsistencies in formatting
                           - Check for sample/template indicators
                           - Look for "not for official use" or similar disclaimers

                        4. Document Quality:
                           - Check for professional printing quality
                           - Verify proper alignment and formatting
                           - Look for official document design elements
                           - Check for proper spacing and typography

                        Return a JSON response with:
                        {{
                            "is_genuine": true/false,
                            "confidence_score": 0.0-1.0,
                            "rejection_reason": "detailed explanation if not genuine",
                            "verification_checks": [
                                {{
                                    "check_type": "type of check",
                                    "status": "passed/failed",
                                    "details": "explanation"
                                }}
                            ],
                            "security_features_found": ["list of found features"],
                            "verification_summary": "detailed explanation"
                        }}

                        Important:
                        - Be extremely thorough in your analysis
                        - Reject any document that shows signs of being:
                          * A sample or template
                          * A photocopy or scan
                          * A digital copy
                          * A test document
                          * A demonstration document
                          * A training document
                        - Look for any signs of forgery or tampering
                        - Provide detailed explanations for your decisions
                        - If in doubt, reject the document
                        - Only accept documents that are clearly genuine official documents
                        """

            response = self.text_processor.process_text(text, verification_prompt)
            verification_data = json.loads(response.strip().replace("```json", "").replace("```", "").strip())

            verification_result["is_genuine"] = verification_data.get("is_genuine", False)
            verification_result["confidence_score"] = verification_data.get("confidence_score", 0.0)
            verification_result["rejection_reason"] = verification_data.get("rejection_reason", "")
            verification_result["verification_checks"] = verification_data.get("verification_checks", [])
            verification_result["security_features_found"] = verification_data.get("security_features_found", [])

            non_genuine_indicators = [
                "sample", "template", "example", "dummy", "test",
                "not for official use", "for demonstration",
                "training", "practice", "mock",
                "photocopy", "scan", "digital copy",
                "this is a", "this document is",
                "for testing", "for practice",
                "do not use", "invalid",
                "unofficial", "non-official"
            ]

            if any(indicator in text.lower() for indicator in non_genuine_indicators):
                verification_result["is_genuine"] = False
                verification_result["rejection_reason"] = "Document contains indicators of being non-genuine"

            if verification_result["confidence_score"] < self.MIN_GENUINENESS_SCORE:
                verification_result["is_genuine"] = False
                verification_result[
                    "rejection_reason"] = f"Low genuineness confidence: {verification_result['confidence_score']}"

            if not verification_result["security_features_found"]:
                verification_result["is_genuine"] = False
                verification_result["rejection_reason"] = "No security features found"

            return verification_result

        except Exception as e:
            logger.error(f"Error in document genuineness verification: {str(e)}")
            verification_result["rejection_reason"] = f"Verification error: {str(e)}"
            return verification_result

    def _identify_documents(self, text: str) -> List[Tuple[str, str]]:
        """Identify separate documents in the text using template matching"""
        try:

            chunks = self._split_into_chunks(text)
            logger.info(f"Split text into {len(chunks)} chunks")

            documents = []
            for chunk in chunks:
                if not chunk.strip():
                    continue

                try:

                    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as temp_file:
                        temp_file.write(chunk)
                        temp_file_path = temp_file.name

                    try:
                        template_match = self.template_matcher.match_document(temp_file_path, chunk)
                        if template_match and template_match['confidence'] >= self.MIN_CONFIDENCE_THRESHOLD:
                            documents.append((chunk, template_match['document_type']))
                            logger.info(f"Identified document of type: {template_match['document_type']}")
                    finally:

                        try:
                            os.unlink(temp_file_path)
                        except Exception as e:
                            logger.warning(f"Error cleaning up temporary file: {str(e)}")

                except Exception as e:
                    logger.warning(f"Error matching chunk: {str(e)}")
                    continue

            if not documents:
                documents = [(text, "unknown")]

            return documents

        except Exception as e:
            logger.error(f"Error identifying documents: {str(e)}")
            return [(text, "unknown")]

    def _process_single_image(self, image_path: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process a single image file"""
        try:
            logger.info(f"Processing single image: {image_path}")

            img = Image.open(image_path)
            ocr_text = pytesseract.image_to_string(img)

            if not self._is_good_ocr_result(ocr_text):
                logger.info("Tesseract OCR yielded poor results, trying Gemini Vision")
                response = self.text_extractor.process_with_gemini(
                    image_path,
                    "Extract all text from this document image. Return only the raw text without any formatting."
                )
                ocr_text = response.strip()

            if not ocr_text.strip():
                logger.warning("No text could be extracted from the image")
                return {
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": image_path,
                    "rejection_reason": "No text could be extracted from the image"
                }

            result = self._process_text_content(ocr_text, image_path, min_confidence)
            if result:
                result["processing_method"] = "ocr"
                result["source_file"] = image_path
                return result

            return None

        except Exception as e:
            logger.error(f"Error processing single image: {str(e)}")
            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": image_path,
                "rejection_reason": f"Error processing image: {str(e)}"
            }


def main():
    input_path = "D:\\imageextractor\\identites\\React_Resume.pdf"
    api_key = API_KEY_1

    processor = DocumentProcessor(api_key=api_key)

    results = processor.process_file(file_path=input_path)

    if results:

        grouped_results = {}
        for result in results:
            doc_type = result["document_type"]
            if doc_type not in grouped_results:
                grouped_results[doc_type] = []
            grouped_results[doc_type].append(result)

        with open("results.json", "w", encoding="utf-8") as f:
            json.dump({
                "total_documents": len(results),
                "document_types": list(grouped_results.keys()),
                "results": grouped_results
            }, f, indent=2, ensure_ascii=False)
        print(f"✅ Processed {len(results)} documents and saved to results.json")
    else:
        print("⚠️ No valid documents detected")


if __name__ == "__main__":
    main()