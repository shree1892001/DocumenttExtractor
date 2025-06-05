import os
import re
import json
import logging
import pytesseract
from PIL import Image
import cv2
import numpy as np
from pdf2image import convert_from_path
from docx import Document
import spacy
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Union, Any
import traceback
import tempfile
from Common.document_constants import (
    DOCUMENT_CATEGORIES,
    DOCUMENT_TYPE_MAPPING,
    DOCUMENT_INDICATORS,
    MIN_CONFIDENCE_THRESHOLD,
    HIGH_CONFIDENCE_THRESHOLD,
    MIN_GENUINENESS_SCORE,
    VERIFICATION_THRESHOLD,
    SUPPORTED_EXTENSIONS,
    FIELD_PATTERNS,
    DOCUMENT_SEPARATORS,
    OCR_ERROR_PATTERNS,
    CHARACTER_PATTERNS,
    CONTENT_INDICATORS,
    DOCUMENT_INDICATOR_KEYWORDS,
    NON_GENUINE_INDICATORS,
    DOCUMENT_PROMPTS,
    DOCUMENT_FIELD_TEMPLATES
)
from Factories.TextExtractorFactory import TextExtractorFactory
from Factories.FieldExtractorFactory import FieldExtractorFactory
from Factories.OCRExtractorFactory import OCRExtractorFactory

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor3:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.text_extractor_factory = TextExtractorFactory()
        self.field_extractor_factory = FieldExtractorFactory()
        self.ocr_factory = OCRExtractorFactory()
        self.nlp = spacy.load("en_core_web_sm")
        self.templates = {}
        self.load_templates()

    def load_templates(self):
        """Load all document templates from the templates directory"""
        try:
            templates_dir = os.path.join(os.path.dirname(__file__), 'templates')
            if not os.path.exists(templates_dir):
                logger.warning(f"Templates directory not found: {templates_dir}")
                return

            logger.info(f"Found {len(os.listdir(templates_dir))} files in templates directory")
            
            for filename in os.listdir(templates_dir):
                if filename.lower().endswith(('.jpg', '.jpeg', '.png', '.pdf', '.docx')):
                    template_path = os.path.join(templates_dir, filename)
                    template_name = os.path.splitext(filename)[0].lower()
                    
                    # Simply store the file path
                    self.templates[template_name] = template_path
                    logger.info(f"Loaded template: {filename}")
            
            logger.info(f"Successfully loaded {len(self.templates)} templates")
            
        except Exception as e:
            logger.error(f"Error loading templates: {str(e)}")
            logger.error(traceback.format_exc())

    def match_template(self, image: np.ndarray) -> Tuple[str, float]:
        """
        Match document against loaded templates and return best match with confidence score.
        This is the primary method for document type detection.
        """
        try:
            if not self.templates:
                logger.warning("No templates loaded for matching")
                return ("unknown", 0.0)

            # Convert to grayscale if needed
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image.copy()

            best_match = ("unknown", 0.0)
            
            # Try different template matching methods
            methods = [
                cv2.TM_CCOEFF_NORMED,
                cv2.TM_CCORR_NORMED
            ]
            
            for template_name, template_path in self.templates.items():
                try:
                    # Load template based on file type
                    if template_path.lower().endswith(('.jpg', '.jpeg', '.png')):
                        template = cv2.imread(template_path)
                    elif template_path.lower().endswith('.pdf'):
                        images = convert_from_path(template_path, first_page=1, last_page=1)
                        template = cv2.cvtColor(np.array(images[0]), cv2.COLOR_RGB2BGR)
                    elif template_path.lower().endswith('.docx'):
                        doc = Document(template_path)
                        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                            doc.save(temp_file.name)
                            images = convert_from_path(temp_file.name, first_page=1, last_page=1)
                            template = cv2.cvtColor(np.array(images[0]), cv2.COLOR_RGB2BGR)
                            os.unlink(temp_file.name)
                    else:
                        continue

                    if template is None:
                        continue

                    # Convert template to grayscale
                    if len(template.shape) == 3:
                        template_gray = cv2.cvtColor(template, cv2.COLOR_BGR2GRAY)
                    else:
                        template_gray = template.copy()
                    
                    # Resize template to match image size
                    template_resized = cv2.resize(template_gray, (gray.shape[1], gray.shape[0]))
                    
                    for method in methods:
                        # Perform template matching
                        result = cv2.matchTemplate(gray, template_resized, method)
                        min_val, max_val, min_loc, max_loc = cv2.minMaxLoc(result)
                        
                        # For TM_CCOEFF_NORMED and TM_CCORR_NORMED, higher values are better
                        if method in [cv2.TM_CCOEFF_NORMED, cv2.TM_CCORR_NORMED]:
                            match_score = max_val
                        else:
                            match_score = 1 - min_val
                        
                        # Update best match if better
                        if match_score > best_match[1]:
                            best_match = (template_name, match_score)

                except Exception as e:
                    logger.warning(f"Error matching template {template_name}: {str(e)}")
                    continue
            
            logger.info(f"Template matching result: {best_match[0]} with confidence {best_match[1]:.2f}")
            return best_match

        except Exception as e:
            logger.error(f"Error matching template: {str(e)}")
            return ("unknown", 0.0)

    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract information.
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in SUPPORTED_EXTENSIONS:
                raise ValueError(f"Unsupported file extension: {file_ext}")

            # Step 1: Convert document to image for template matching
            if file_ext in ['.jpg', '.jpeg', '.png']:
                image = cv2.imread(file_path)
            elif file_ext == '.pdf':
                images = convert_from_path(file_path, first_page=1, last_page=1)
                image = cv2.cvtColor(np.array(images[0]), cv2.COLOR_RGB2BGR)
            elif file_ext == '.docx':
                doc = Document(file_path)
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                    doc.save(temp_file.name)
                    images = convert_from_path(temp_file.name, first_page=1, last_page=1)
                    image = cv2.cvtColor(np.array(images[0]), cv2.COLOR_RGB2BGR)
                    os.unlink(temp_file.name)
            else:
                raise ValueError(f"Unsupported file type for template matching: {file_ext}")

            # Step 2: Match template first - this is mandatory
            doc_type, template_confidence = self.match_template(image)
            
            if doc_type == "unknown" or template_confidence < MIN_CONFIDENCE_THRESHOLD:
                raise ValueError(f"Could not match document template (confidence: {template_confidence})")

            logger.info(f"Successfully matched template: {doc_type} with confidence {template_confidence}")

            # Step 3: Only proceed with text extraction if template is matched
            text_extractor = self.text_extractor_factory.create_extractor(file_path)
            text = text_extractor.extract_text(file_path)
            
            # Step 4: Verify document genuineness
            genuineness_score = self.verify_document_genuineness(text)
            if genuineness_score < MIN_GENUINENESS_SCORE:
                raise ValueError("Document appears to be non-genuine")
            
            # Step 5: Get appropriate field extractor for the matched template
            field_extractor = self.field_extractor_factory.create_extractor(doc_type)
            
            # Step 6: Extract fields using template-specific patterns
            fields = field_extractor.extract_fields(text)
            
            # Step 7: Verify extracted information
            verification_score = self.verify_extracted_info(fields)
            if verification_score < VERIFICATION_THRESHOLD:
                raise ValueError("Extracted information verification failed")
            
            return {
                'document_type': doc_type,
                'template_confidence': template_confidence,
                'genuineness_score': genuineness_score,
                'verification_score': verification_score,
                'extracted_fields': fields,
                'raw_text': text
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            logger.error(traceback.format_exc())
            raise

    def verify_document_genuineness(self, text: str) -> float:
        """
        Verify if the document appears to be genuine.
        Returns a score between 0 and 1.
        """
        text = text.lower()
        score = 1.0
        
        # Check for non-genuine indicators
        for indicator in NON_GENUINE_INDICATORS:
            if indicator in text:
                score -= 0.2
        
        # Check for document indicator keywords
        keyword_count = sum(1 for keyword in DOCUMENT_INDICATOR_KEYWORDS if keyword in text)
        score += min(keyword_count * 0.1, 0.3)
        
        return max(0.0, min(1.0, score))

    def verify_extracted_info(self, fields: Dict[str, Any]) -> float:
        """
        Verify the extracted information.
        Returns a score between 0 and 1.
        """
        try:
            if not fields:
                return 0.0
            
            score = 0.0
            total_fields = len(fields)
            
            # Check for required fields
            for field_name, value in fields.items():
                if value and isinstance(value, str) and len(value.strip()) > 0:
                    score += 1.0
            
            # Normalize score
            return score / total_fields if total_fields > 0 else 0.0
        except Exception as e:
            logger.error(f"Error verifying extracted info: {str(e)}")
            return 0.0

    def _clean_value(self, value: str) -> str:
        """
        Clean and validate extracted value.
        """
        try:
            if not value:
                return ""
            
            # Remove common OCR errors
            for pattern, replacement in OCR_ERROR_PATTERNS.items():
                value = re.sub(pattern, replacement, value)
            
            # Remove invalid characters
            for pattern in CHARACTER_PATTERNS:
                value = re.sub(pattern, "", value)
            
            # Remove extra whitespace
            value = " ".join(value.split())
            
            return value.strip()
        except Exception as e:
            logger.error(f"Error cleaning value: {str(e)}")
            return ""

    def detect_document_type(self, text: str) -> str:
        """
        Detect the type of document based on its content.
        """
        text = text.lower()
        
        # Check for document type indicators
        for doc_type, indicators in DOCUMENT_INDICATORS.items():
            if any(indicator.lower() in text for indicator in indicators):
                return doc_type
        
        # Check for document type mappings
        for key, value in DOCUMENT_TYPE_MAPPING.items():
            if key in text:
                return value
        
        # Check document categories
        for category, types in DOCUMENT_CATEGORIES.items():
            for doc_type in types:
                if doc_type.replace('_', ' ') in text:
                    return doc_type
        
        return "unknown"

    def extract_fields(self, text: str) -> Dict[str, Any]:
        """
        Extract fields from the document text.
        """
        try:
            fields = {}
            doc = self.nlp(text)
            
            # Extract fields based on patterns
            for field_name, patterns in FIELD_PATTERNS.items():
                for pattern in patterns:
                    matches = re.finditer(pattern, text, re.IGNORECASE)
                    for match in matches:
                        value = match.group(1).strip()
                        if value:
                            fields[field_name] = value
                            break
            
            # Clean and validate extracted values
            cleaned_fields = {}
            for field_name, value in fields.items():
                if isinstance(value, (dict, list)):
                    value = str(value)
                cleaned_value = self._clean_value(value)
                if cleaned_value:
                    cleaned_fields[field_name] = cleaned_value
            
            return cleaned_fields
        except Exception as e:
            logger.error(f"Error extracting fields: {str(e)}")
            return {}

    def _process_text_content(self, text: str, source_file: str, min_confidence: float) -> Optional[Dict[str, Any]]:
        """Process text content directly without creating temporary files"""
        try:
            # First try to match template from the text
            doc_type = "unknown"
            template_confidence = 0.0

            # Check against template patterns
            for template_name, template in self.templates.items():
                # Get template patterns for this document type
                template_patterns = DOCUMENT_FIELD_TEMPLATES.get(template_name, {}).get("template_patterns", [])
                
                # Calculate match score
                match_score = 0.0
                for pattern in template_patterns:
                    if re.search(pattern.lower(), text.lower()):
                        match_score += 1.0
                
                # Normalize score
                if template_patterns:
                    match_score /= len(template_patterns)
                
                # Update best match if better
                if match_score > template_confidence:
                    doc_type = template_name
                    template_confidence = match_score

            if doc_type == "unknown" or template_confidence < min_confidence:
                logger.warning("Could not determine document type from template matching")
                return {
                    "status": "rejected",
                    "document_type": "unknown",
                    "source_file": source_file,
                    "rejection_reason": "Document type could not be determined from template matching",
                    "confidence": template_confidence
                }

            logger.info(f"Matched template from text: {doc_type} with confidence {template_confidence}")

            # Get appropriate field extractor
            field_extractor = self.field_extractor_factory.create_extractor(doc_type)
            
            # Extract fields
            extracted_data = field_extractor.extract_fields(text)
            logger.info(f"Extracted data: {json.dumps(extracted_data, indent=2)}")

            # Verify document
            verification_result = self.verify_document(extracted_data, doc_type)

            if not verification_result["is_genuine"]:
                rejection_reason = verification_result.get("rejection_reason", "Document failed authenticity verification")
                logger.warning(f"Document rejected - Not genuine: {rejection_reason}")
                return {
                    "status": "rejected",
                    "document_type": doc_type,
                    "source_file": source_file,
                    "rejection_reason": rejection_reason,
                    "verification_result": verification_result
                }

            # Verify extracted information
            verification_score = self.verify_extracted_info(extracted_data)
            if verification_score < VERIFICATION_THRESHOLD:
                return {
                    "status": "rejected",
                    "document_type": doc_type,
                    "source_file": source_file,
                    "rejection_reason": "Document failed validation",
                    "verification_result": verification_result
                }

            return {
                "extracted_data": extracted_data,
                "status": "success",
                "confidence": template_confidence,
                "document_type": doc_type,
                "source_file": source_file,
                "validation_level": "strict",
                "verification_result": verification_result
            }

        except Exception as e:
            logger.exception(f"Error processing text content: {str(e)}")
            return {
                "status": "error",
                "document_type": "unknown",
                "source_file": source_file,
                "rejection_reason": f"Error processing document: {str(e)}"
            }

    def verify_document(self, extracted_data: Dict[str, Any], doc_type: str) -> Dict[str, Any]:
        """Verify document authenticity and data validity"""
        try:
            # Check for required fields
            required_fields = DOCUMENT_FIELD_TEMPLATES.get(doc_type, {}).get("required_fields", [])
            missing_fields = [field for field in required_fields if field not in extracted_data]
            
            # Calculate confidence scores
            field_confidence = len(extracted_data) / len(required_fields) if required_fields else 1.0
            data_quality = sum(1 for v in extracted_data.values() if v and len(str(v).strip()) > 0) / len(extracted_data) if extracted_data else 0.0
            
            overall_confidence = (field_confidence + data_quality) / 2

            verification_result = {
                "is_genuine": overall_confidence >= VERIFICATION_THRESHOLD,
                "confidence_score": overall_confidence,
                "verification_checks": {
                    "field_completeness": {
                        "passed": not missing_fields,
                        "details": f"Missing fields: {', '.join(missing_fields)}" if missing_fields else "All required fields present",
                        "confidence": field_confidence
                    },
                    "data_quality": {
                        "passed": data_quality >= 0.7,
                        "details": f"Data quality score: {data_quality:.2f}",
                        "confidence": data_quality
                    }
                }
            }

            if not verification_result["is_genuine"]:
                verification_result["rejection_reason"] = (
                    f"Low confidence score: {overall_confidence:.2f}. "
                    f"Missing fields: {', '.join(missing_fields)}" if missing_fields else ""
                )

            logger.info(f"Document verification results: {json.dumps(verification_result, indent=2)}")

            return verification_result

        except Exception as e:
            logger.exception(f"Error verifying document genuineness: {str(e)}")
            return {
                "is_genuine": False,
                "confidence_score": 0.0,
                "rejection_reason": f"Error during verification: {str(e)}",
                "verification_checks": {
                    "field_completeness": {"passed": False, "details": f"Error during verification: {str(e)}", "confidence": 0.0},
                    "data_quality": {"passed": False, "details": "Verification failed", "confidence": 0.0}
                }
            }

    def _verify_document_genuineness(self, text: str, doc_type: str) -> Dict[str, Any]:
        """Verify if the document is genuine before extraction"""
        verification_result = {
            "is_genuine": False,
            "confidence_score": 0.0,
            "rejection_reason": "",
            "verification_checks": [],
            "security_features_found": []
        }

        try:
            # Check for non-genuine indicators
            if any(indicator in text.lower() for indicator in NON_GENUINE_INDICATORS):
                verification_result["is_genuine"] = False
                verification_result["rejection_reason"] = "Document contains indicators of being non-genuine"
                return verification_result

            # Check for document type indicators
            doc_type_indicators = DOCUMENT_INDICATORS.get(doc_type, [])
            indicator_count = sum(1 for indicator in doc_type_indicators if indicator.lower() in text.lower())
            indicator_confidence = indicator_count / len(doc_type_indicators) if doc_type_indicators else 0.0

            # Check for security features
            security_features = []
            for feature in DOCUMENT_INDICATOR_KEYWORDS:
                if feature.lower() in text.lower():
                    security_features.append(feature)

            security_confidence = len(security_features) / len(DOCUMENT_INDICATOR_KEYWORDS) if DOCUMENT_INDICATOR_KEYWORDS else 0.0

            # Calculate overall confidence
            overall_confidence = (indicator_confidence + security_confidence) / 2

            verification_result.update({
                "is_genuine": overall_confidence >= MIN_GENUINENESS_SCORE,
                "confidence_score": overall_confidence,
                "verification_checks": [
                    {
                        "type": "document_type",
                        "passed": indicator_confidence >= 0.5,
                        "confidence": indicator_confidence
                    },
                    {
                        "type": "security_features",
                        "passed": security_confidence >= 0.3,
                        "confidence": security_confidence
                    }
                ],
                "security_features_found": security_features
            })

            if not verification_result["is_genuine"]:
                verification_result["rejection_reason"] = (
                    f"Low genuineness confidence: {overall_confidence:.2f}. "
                    f"Found {len(security_features)} security features."
                )

            return verification_result

        except Exception as e:
            logger.error(f"Error in document genuineness verification: {str(e)}")
            verification_result["rejection_reason"] = f"Verification error: {str(e)}"
            return verification_result

    def preprocess_image(self, image: Image.Image) -> np.ndarray:
        """Preprocess image for better OCR results"""
        try:
            # Convert PIL Image to OpenCV format
            img = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
            
            # Convert to grayscale
            gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
            
            # Apply adaptive thresholding
            thresh = cv2.adaptiveThreshold(
                gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # Denoise
            denoised = cv2.fastNlMeansDenoising(thresh)
            
            # Increase contrast
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
            enhanced = clahe.apply(gray)
            
            # Combine results
            final = cv2.bitwise_and(enhanced, denoised)
            
            return final
        except Exception as e:
            logger.error(f"Error preprocessing image: {str(e)}")
            return np.array(image)

    def extract_text_from_image(self, image_path: str) -> str:
        """Extract text from an image using optimized OCR"""
        try:
            # Read image
            img = Image.open(image_path)
            
            # Preprocess image
            processed_img = self.preprocess_image(img)
            
            # Try different PSM modes and languages
            psm_modes = [6, 3, 4, 1]  # Different page segmentation modes
            languages = ['eng', 'eng+fra', 'eng+deu']  # Multiple languages
            best_text = ""
            
            for psm in psm_modes:
                for lang in languages:
                    config = f'--oem 3 --psm {psm} -l {lang}'
                    try:
                        text = pytesseract.image_to_string(
                            processed_img,
                            config=config
                        )
                        if len(text.strip()) > len(best_text.strip()):
                            best_text = text
                    except Exception as e:
                        logger.warning(f"OCR failed for PSM {psm} and language {lang}: {str(e)}")
            
            return best_text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            raise

    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using optimized OCR"""
        try:
            text_content = []
            pdf_document = fitz.open(pdf_path)
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                
                # Try normal text extraction first
                text = page.get_text()
                
                # If no text found or text is too short, try OCR
                if not text.strip() or len(text.strip()) < 50:
                    # Convert page to high-resolution image
                    pix = page.get_pixmap(matrix=fitz.Matrix(400/72, 400/72))  # 400 DPI
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    # Preprocess image
                    processed_img = self.preprocess_image(img)
                    
                    # Try different PSM modes and languages
                    psm_modes = [6, 3, 4, 1]
                    languages = ['eng', 'eng+fra', 'eng+deu']
                    best_text = ""
                    
                    for psm in psm_modes:
                        for lang in languages:
                            config = f'--oem 3 --psm {psm} -l {lang}'
                            try:
                                temp_text = pytesseract.image_to_string(
                                    processed_img,
                                    config=config
                                )
                                if len(temp_text.strip()) > len(best_text.strip()):
                                    best_text = temp_text
                            except Exception as e:
                                logger.warning(f"OCR failed for PSM {psm} and language {lang}: {str(e)}")
                    
                    if best_text.strip():
                        text = best_text
                
                # Process images in the page
                if page.get_images():
                    for img_index, img in enumerate(page.get_images(full=True)):
                        try:
                            xref = img[0]
                            base_image = pdf_document.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # Save image temporarily
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                                temp_file.write(image_bytes)
                                temp_path = temp_file.name
                            
                            # Extract text from image
                            try:
                                img_text = self.extract_text_from_image(temp_path)
                                if img_text.strip():
                                    text += "\n" + img_text
                            except Exception as e:
                                logger.warning(f"OCR failed for image {img_index}: {str(e)}")
                            
                            # Clean up temp file
                            os.unlink(temp_path)
                        except Exception as e:
                            logger.warning(f"Failed to process image {img_index}: {str(e)}")
                
                if text.strip():
                    text_content.append(text.strip())
            
            pdf_document.close()
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise

    def extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file using optimized OCR"""
        try:
            doc = Document(docx_path)
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            
            # Extract text from images
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob
                        
                        # Save image temporarily
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                            temp_file.write(image_data)
                            temp_path = temp_file.name
                        
                        # Extract text from image
                        try:
                            img_text = self.extract_text_from_image(temp_path)
                            if img_text.strip():
                                text_content.append(img_text.strip())
                        except Exception as e:
                            logger.warning(f"OCR failed for image in DOCX: {str(e)}")
                        
                        # Clean up temp file
                        os.unlink(temp_path)
                    except Exception as e:
                        logger.warning(f"Failed to process image in DOCX: {str(e)}")
            
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise 