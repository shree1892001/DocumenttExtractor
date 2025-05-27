from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional
import json
from Extractor.ImageExtractor import ImageTextExtractor
from Common.constants import *
import re
import datetime
from docx import Document
import google.generativeai as genai
import fitz  # PyMuPDF
import tempfile
import os
import logging
from PIL import Image
import pytesseract
import io

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TextExtractor:
    """Handles text extraction from different file formats"""
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.image_extractor = ImageTextExtractor(api_key)
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    def extract_text(self, file_path: str) -> str:
        """Extract text from any supported file format"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.docx':
                return self._extract_from_docx(file_path)
            elif file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                return self._extract_from_image(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise

    def process_with_gemini(self, input_data: str, prompt: str) -> str:
        """Process text or image with Gemini"""
        try:
            # Check if input_data is a file path or text
            if os.path.exists(input_data) and os.path.isfile(input_data):
                # It's a file path, process as image
                image = Image.open(input_data)
                response = self.model.generate_content([prompt, image])
            else:
                # It's text, process directly
                response = self.model.generate_content([prompt, input_data])
            return response.text
        except Exception as e:
            logger.error(f"Error processing with Gemini: {str(e)}")
            raise

    def _extract_from_docx(self, docx_path: str) -> str:
        """Extract text from a .docx file"""
        try:
            doc = Document(docx_path)
            text_content = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    def _extract_from_pdf(self, pdf_path: str) -> str:
        """Extract text from a PDF file (both normal and scanned)"""
        try:
            text_content = []
            pdf_document = fitz.open(pdf_path)
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                
                # Try normal text extraction first
                text = page.get_text()
                
                # If no text found, try OCR for scanned PDF
                if not text.strip():
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text = pytesseract.image_to_string(img)
                
                if text.strip():
                    text_content.append(text.strip())
            
            pdf_document.close()
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise

    def _extract_from_image(self, image_path: str) -> str:
        """Extract text from an image file using OCR and Gemini"""
        try:
            # First try OCR
            img = Image.open(image_path)
            ocr_text = pytesseract.image_to_string(img)
            
            if not ocr_text.strip():
                # If OCR fails, use Gemini Vision
                response = self.process_with_gemini(
                    image_path,
                    "Extract all text from this document image. Return only the raw text without any formatting."
                )
                return response.strip()
            
            return ocr_text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            raise

class DocumentExtractor(ABC):
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.text_extractor = TextExtractor(api_key)

    @abstractmethod
    def get_extraction_prompt(self) -> str:
        """Return the specific prompt for this document type"""
        pass

    @abstractmethod
    def validate_fields(self, extracted_data: Dict[str, Any]) -> bool:
        """Validate the extracted fields"""
        pass

    def extract_fields(self, file_path: str) -> Dict[str, Any]:
        """Extract fields using the document-specific prompt"""
        try:
            # Extract text from the file
            text = self.text_extractor.extract_text(file_path)
            
            # Get document-specific prompt
            prompt = self.get_extraction_prompt()
            
            # Process text with Gemini
            response = self.text_extractor.process_with_gemini(text, prompt)
            
            # Clean and parse JSON response
            extracted_data = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
            
            # Validate extracted data
            if not self.validate_fields(extracted_data):
                raise ValueError("Extracted data validation failed")
            
            return extracted_data
        except Exception as e:
            logger.error(f"Error extracting fields from {file_path}: {str(e)}")
            raise

class LicenseExtractor(DocumentExtractor):
    def get_extraction_prompt(self) -> str:
        return LICENSE_EXTRACTION

    def validate_fields(self, extracted_data: Dict[str, Any]) -> bool:
        required_fields = {
            "license_number", "name", "date_of_birth",
            "valid_from", "valid_until"
        }
        return all(field in extracted_data.get("data", {}) for field in required_fields)

class PancardExtractor(DocumentExtractor):
    def get_extraction_prompt(self) -> str:
        return PAN_CARD_EXTRACTION

    def validate_fields(self, extracted_data: Dict[str, Any]) -> bool:
        required_fields = {"pan_number", "name", "fathers_name", "date_of_birth"}
        return all(field in extracted_data.get("data", {}) for field in required_fields)

class AadhaarCardExtractor(DocumentExtractor):
    def get_extraction_prompt(self) -> str:
        return AADHAR_CARD_EXTRACTION

    def validate_fields(self, extracted_data: Dict[str, Any]) -> bool:
        required_fields = {"aadhaar_number", "name", "gender", "date_of_birth"}
        return all(field in extracted_data.get("data", {}) for field in required_fields)

class SsnExtractor(DocumentExtractor):
    def get_extraction_prompt(self) -> str:
        return SSN_EXTRACTION

    def validate_fields(self, extracted_data: Dict[str, Any]) -> bool:
        required_fields = {"ssn", "name", "date_of_birth", "address"}
        return all(field in extracted_data.get("data", {}) for field in required_fields)

class PassportExtractor(DocumentExtractor):
    def get_extraction_prompt(self) -> str:
        return PASSPORT_EXTRACTION

    def validate_fields(self, extracted_data: Dict[str, Any]) -> bool:

        required_fields = {
            "passport_number",
            "surname",
            "given_names",
            "nationality",
            "date_of_birth",
            "gender",
            "date_of_issue",
            "date_of_expiry",
            "place_of_issue"
        }

        if not all(field in extracted_data.get("data", {}) for field in required_fields):
            return False

        data = extracted_data.get("data", {})

        passport_number = data.get("passport_number", "")
        if not re.match(r'^[A-Z][0-9]{7}$', passport_number):
            return False

        date_fields = ["date_of_birth", "date_of_issue", "date_of_expiry"]
        for date_field in date_fields:
            date_value = data.get(date_field, "")
            try:
                datetime.strptime(date_value, '%Y-%m-%d')
            except ValueError:
                return False
        country_code = data.get("country_code", "")
        if not re.match(r'^[A-Z]{3}$', country_code):
            return False

        if data.get("gender", "").upper() not in ["M", "F", "X"]:
            return False

        return True

class TextExtractorFactory:
    """Factory class for creating text extractors based on file type"""
    
    @staticmethod
    def create_extractor(file_path: str, api_key: str) -> 'TextExtractor':
        """Create appropriate text extractor based on file extension"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.docx':
            return DocxTextExtractor(api_key)
        elif file_ext == '.pdf':
            return PdfTextExtractor(api_key)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
            return ImageTextExtractor(api_key)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

class BaseTextExtractor(ABC):
    """Abstract base class for text extractors"""
    def __init__(self, api_key: str):
        self.api_key = api_key
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')

    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """Extract text from file"""
        pass

    def process_with_gemini(self, text: str, prompt: str) -> str:
        """Process text using Gemini"""
        try:
            response = self.model.generate_content([prompt, text])
            return response.text
        except Exception as e:
            raise ValueError(f"Error processing text with Gemini: {str(e)}")

class DocxTextExtractor(BaseTextExtractor):
    """Extractor for DOCX files"""
    def extract_text(self, file_path: str) -> str:
        try:
            # Check if file exists
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"DOCX file not found at path: {file_path}")
            
            # Check if file is readable
            if not os.access(file_path, os.R_OK):
                raise PermissionError(f"Cannot read DOCX file at path: {file_path}")
            
            # Check if file is a valid DOCX
            if not self._is_valid_docx(file_path):
                raise ValueError(f"Invalid or corrupted DOCX file at path: {file_path}")
            
            # First try to extract text directly
            try:
                doc = Document(file_path)
                text_content = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content.append(para.text.strip())
                text = '\n'.join(text_content)
                
                # If no text found, try to extract images
                if not text.strip():
                    return self._extract_text_from_docx_images(file_path)
                return text
                
            except Exception as e:
                logger.warning(f"Error extracting text directly from DOCX: {str(e)}")
                # If direct text extraction fails, try to extract images
                return self._extract_text_from_docx_images(file_path)
                
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise

    def _is_valid_docx(self, file_path: str) -> bool:
        """Check if file is a valid DOCX file"""
        try:
            # Check file extension
            if not file_path.lower().endswith('.docx'):
                return False
            
            # Check file size
            if os.path.getsize(file_path) == 0:
                return False
            
            # Try to open with python-docx
            try:
                doc = Document(file_path)
                # Try to access a basic property to verify it's a valid DOCX
                _ = doc.core_properties
                return True
            except Exception:
                return False
                
        except Exception:
            return False

    def _extract_text_from_docx_images(self, file_path: str) -> str:
        """Extract text from images in DOCX file"""
        try:
            # Verify file still exists and is readable
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"DOCX file not found at path: {file_path}")
            if not os.access(file_path, os.R_OK):
                raise PermissionError(f"Cannot read DOCX file at path: {file_path}")
            
            # Verify it's still a valid DOCX
            if not self._is_valid_docx(file_path):
                raise ValueError(f"Invalid or corrupted DOCX file at path: {file_path}")
            
            with tempfile.TemporaryDirectory() as temp_dir:
                # Extract images from DOCX
                doc = Document(file_path)
                image_texts = []
                
                # Process each paragraph for images
                for para in doc.paragraphs:
                    for run in para.runs:
                        if run._element.xpath('.//w:drawing'):
                            # Found an image, extract it
                            image_path = os.path.join(temp_dir, f"image_{len(image_texts)}.png")
                            try:
                                # Extract image data
                                image_data = run._element.xpath('.//a:blip/@r:embed')[0]
                                image_part = doc.part.related_parts[image_data]
                                
                                # Save image
                                with open(image_path, 'wb') as f:
                                    f.write(image_part.blob)
                                
                                # Extract text from image
                                img = Image.open(image_path)
                                ocr_text = pytesseract.image_to_string(img)
                                
                                if not ocr_text.strip():
                                    # If OCR fails, use Gemini Vision
                                    response = self.process_with_gemini(
                                        image_path,
                                        "Extract all text from this document image. Return only the raw text without any formatting."
                                    )
                                    ocr_text = response.strip()
                                
                                if ocr_text.strip():
                                    image_texts.append(ocr_text.strip())
                            except Exception as e:
                                logger.warning(f"Error processing image in DOCX: {str(e)}")
                                continue
                
                if not image_texts:
                    raise ValueError("No text or images found in DOCX file")
                
                return '\n'.join(image_texts)
                
        except Exception as e:
            logger.error(f"Error extracting text from DOCX images: {str(e)}")
            raise

    def process_with_gemini(self, image_path: str, prompt: str) -> str:
        """Process image with Gemini Vision"""
        try:
            # Verify image file exists and is readable
            if not os.path.exists(image_path):
                raise FileNotFoundError(f"Image file not found at path: {image_path}")
            if not os.access(image_path, os.R_OK):
                raise PermissionError(f"Cannot read image file at path: {image_path}")
            
            image = Image.open(image_path)
            response = self.model.generate_content([prompt, image])
            return response.text
        except Exception as e:
            logger.error(f"Error processing image with Gemini: {str(e)}")
            raise

class PdfTextExtractor(BaseTextExtractor):
    """Extractor for PDF files"""
    def extract_text(self, file_path: str) -> str:
        try:
            text_content = []
            pdf_document = fitz.open(file_path)
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                
                # Try normal text extraction first
                text = page.get_text()
                
                # If no text found, try OCR for scanned PDF
                if not text.strip():
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text = pytesseract.image_to_string(img)
                
                if text.strip():
                    text_content.append(text.strip())
            
            pdf_document.close()
            return '\n'.join(text_content)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise

class ImageTextExtractor(BaseTextExtractor):
    """Extractor for image files"""
    def extract_text(self, file_path: str) -> str:
        try:
            # First try OCR
            img = Image.open(file_path)
            ocr_text = pytesseract.image_to_string(img)
            
            if not ocr_text.strip():
                # If OCR fails, use Gemini Vision
                response = self.process_with_gemini(
                    file_path,
                    "Extract all text from this document image. Return only the raw text without any formatting."
                )
                return response.strip()
            
            return ocr_text.strip()
        except Exception as e:
            logger.error(f"Error extracting text from image: {str(e)}")
            raise

class DocumentExtractorFactory:
    _extractors = {
        "license": LicenseExtractor,
        "pan_card": PancardExtractor,
        "aadhaar_card": AadhaarCardExtractor,
        "aadhaarcard": AadhaarCardExtractor,
        "aadhaar card": AadhaarCardExtractor,
        "Pancard": PancardExtractor,
        "pancard": PancardExtractor,
        "Passport": PassportExtractor,
        "passport": PassportExtractor,
        "ssn": SsnExtractor,
        "Ssn": SsnExtractor
    }

    def __init__(self, api_key: str):
        self.api_key = api_key

    @classmethod
    def get_extractor(cls, document_type: str, api_key: str) -> DocumentExtractor:
        extractor_class = cls._extractors.get(document_type.lower().replace("_", ""))
        if not extractor_class:
            raise ValueError(f"No extractor found for document type: {document_type}")
        return extractor_class(api_key)

    @classmethod
    def register_extractor(cls, doc_type: str, extractor: type):
        if not issubclass(extractor, DocumentExtractor):
            raise ValueError("Extractor must inherit from DocumentExtractor")
        cls._extractors[doc_type.lower()] = extractor

    def extract_text_from_template(self, template_path: str) -> Dict[str, Any]:
        """Extract text and structure from a template file"""
        try:
            text_extractor = TextExtractorFactory.create_extractor(template_path, self.api_key)
            
            text = text_extractor.extract_text(template_path)
            
            filename = os.path.basename(template_path)
            doc_type = filename.replace('sample_', '').replace('.docx', '').replace('.pdf', '').replace('.jpg', '').replace('.png', '')
            doc_type = self._standardize_document_type(doc_type)
            

            fields = self._extract_fields_from_text(text)
            
            return {
                'document_type': doc_type,
                'content': text,
                'fields': fields,
                'structure': text
            }
        except Exception as e:
            logger.error(f"Error extracting text from template {template_path}: {str(e)}")
            raise

    def _standardize_document_type(self, doc_type: str) -> str:
        """Standardize document type name"""
        doc_type = doc_type.lower().strip()
        type_mapping = {
            "aadhaar": "aadhaar_card",
            "aadhaarcard": "aadhaar_card",
            "aadharcard": "aadhaar_card",
            "pan": "pan_card",
            "pancard": "pan_card",
            "license": "license",
            "driving license": "license",
            "dl": "license"
        }
        return type_mapping.get(doc_type, doc_type)

    def _extract_fields_from_text(self, text: str) -> set:
        """Extract potential field names from text"""
        fields = set()
        # Look for text between brackets or after colons
        field_patterns = [
            r'\{([^}]+)\}',  # {field_name}
            r'([^:]+):',     # field_name:
            r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)'  # Capitalized words
        ]
        
        for pattern in field_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                if isinstance(match, tuple):
                    field = match[0] if match[0] else match[1]
                else:
                    field = match
                fields.add(field.strip())
        
        return fields

    def load_templates_from_directory(self, templates_dir: str) -> Dict[str, Dict[str, Any]]:
        """Load all templates from a directory"""
        templates = {}
        try:
            for filename in os.listdir(templates_dir):
                file_path = os.path.join(templates_dir, filename)
                if os.path.isfile(file_path):
                    template_data = self.extract_text_from_template(file_path)
                    if template_data:
                        templates[template_data['document_type']] = template_data
            return templates
        except Exception as e:
            logger.error(f"Error loading templates from directory {templates_dir}: {str(e)}")
            raise

    def match_document_with_templates(self, file_path: str, templates: Dict[str, Dict[str, Any]]) -> Dict[str, Any]:
        """Match a document against templates"""
        try:
            # Extract text from the document
            text = self.text_extractor.extract_text(file_path)
            
            best_match = {
                'document_type': None,
                'confidence': 0.0,
                'matched_fields': {},
                'additional_info': None
            }
            
            # Try to match against each template
            for doc_type, template in templates.items():
                try:
                    prompt = f"""
                    You are a document analysis expert. Analyze the following document text and determine if it matches the {doc_type} template structure.
                    
                    Document Text:
                    {text}
                    
                    Template Structure:
                    {template['content']}
                    
                    Please analyze the document and return a JSON response with:
                    1. Whether this document matches the {doc_type} template (true/false)
                    2. Confidence score (0-1) based on how well the document matches the template structure
                    3. Extracted fields based on the template structure
                    4. Any additional relevant information found
                    
                    Format the response as:
                    {{
                        "matches_template": true/false,
                        "confidence_score": 0.0-1.0,
                        "extracted_fields": {{
                            "field1": "value1",
                            "field2": "value2"
                        }},
                        "additional_info": "any additional relevant information"
                    }}
                    """
                    
                    response = self.text_extractor.process_with_gemini(text, prompt)
                    match_result = json.loads(response.strip().replace("```json", "").replace("```", "").strip())
                    
                    if match_result.get('matches_template', False) and match_result.get('confidence_score', 0) > best_match['confidence']:
                        best_match = {
                            'document_type': doc_type,
                            'confidence': match_result['confidence_score'],
                            'matched_fields': match_result.get('extracted_fields', {}),
                            'additional_info': match_result.get('additional_info')
                        }
                except Exception as e:
                    logger.warning(f"Error matching template for {doc_type}: {str(e)}")
                    continue
            
            if not best_match['document_type']:
                raise ValueError("No matching template found for the document")
            
            return best_match
            
        except Exception as e:
            logger.error(f"Error matching document with templates: {str(e)}")
            raise